#!/usr/bin/env python3
"""Build the competition-report Word document from the current Markdown source.

The template remains untouched.  Pandoc creates native Word tables, images, and
Office Math, then Word fields are added for chapter-local figure/table captions.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


SCRIPT_PATH = Path(__file__).resolve()
REPORT_DIR = SCRIPT_PATH.parent
REPO_ROOT = REPORT_DIR.parents[1]
DEFAULT_SOURCE = REPORT_DIR / "仿真分析报告_正文骨架.md"
DEFAULT_TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
DEFAULT_OUTPUT = REPORT_DIR / "MoSim_仿真分析报告_国赛版.docx"
DEFAULT_WORK_DIR = REPO_ROOT / "Results" / "docx_build" / "competition_report_20260802"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\n]+)\)")
CAPTION_RE = re.compile(
    r"^(图|表)\s*\d+(?:\s*[-—–]\s*\d+)?[　 \t]+(.+?)\s*$"
)
CHAPTER_PREFIX_RE = re.compile(
    r"^\s*(?:[一二三四五六七八九十百千万]+[、.．]|\d+(?:\.\d+)*[、.．]?)\s*"
)
SUBHEADING_PREFIX_RE = re.compile(r"^\s*\d+(?:\.\d+){1,3}[、.．]?\s*")
LATEX_FENCE_RE = re.compile(r"```latex\s*\r?\n(.*?)\r?\n```", re.DOTALL)


def emit(message: str) -> None:
    print(f"[competition-report] {message}", flush=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--work-dir", type=Path, default=DEFAULT_WORK_DIR)
    parser.add_argument(
        "--export-pdf",
        action="store_true",
        help="Also export a PDF beside the build manifest for visual QA.",
    )
    return parser.parse_args()


def read_source(source: Path) -> tuple[str, str]:
    text = source.read_text(encoding="utf-8")
    title_match = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    if not title_match:
        raise ValueError(f"No document title found in {source}")
    return title_match.group(1).strip(), text


def image_paths(markdown: str) -> list[str]:
    paths: list[str] = []
    for _alt, raw_path in IMAGE_RE.findall(markdown):
        path = raw_path.strip()
        if path.startswith("<") and path.endswith(">"):
            path = path[1:-1]
        paths.append(path)
    return paths


def caption_counts(markdown: str) -> Counter[str]:
    result: Counter[str] = Counter()
    for line in markdown.splitlines():
        match = CAPTION_RE.match(line.strip())
        if match:
            result["figure" if match.group(1) == "图" else "table"] += 1
    return result


def latex_block_to_display_math(match: re.Match[str]) -> str:
    body = match.group(1).strip()
    if body.startswith(r"\[") and body.endswith(r"\]"):
        body = body[2:-2].strip()
    return f"\n$$\n{body}\n$$\n"


def shift_heading_levels(markdown: str) -> str:
    """Remove the Markdown title and promote ##/###/#### to H1/H2/H3."""
    lines = markdown.splitlines(keepends=True)
    output: list[str] = []
    first_title_removed = False
    in_fence = False

    for line in lines:
        if line.lstrip().startswith("```"):
            in_fence = not in_fence
            output.append(line)
            continue
        if not in_fence:
            match = re.match(r"^(#{1,6})\s+(.*?)(\r?\n)?$", line)
            if match:
                hashes, heading, newline = match.groups()
                if len(hashes) == 1 and not first_title_removed:
                    first_title_removed = True
                    continue
                if len(hashes) >= 2:
                    output.append("#" * (len(hashes) - 1) + " " + heading + (newline or ""))
                    continue
        output.append(line)
    return "".join(output)


def make_pandoc_markdown(source_markdown: str) -> str:
    prepared = LATEX_FENCE_RE.sub(latex_block_to_display_math, source_markdown)
    prepared = shift_heading_levels(prepared)

    # A lone Markdown image becomes a Pandoc figure using alt text as a second
    # caption.  Keep it as a pure image; the following project caption is the
    # authoritative caption and receives the Word fields later.
    def rewrite_image(match: re.Match[str]) -> str:
        path = match.group(2).strip()
        return f"![]({path}){{width=15cm}}"

    return IMAGE_RE.sub(rewrite_image, prepared)


def find_pandoc() -> str:
    candidates = [
        os.environ.get("PANDOC"),
        shutil.which("pandoc"),
        r"D:\Dev\Anaconda3\Library\bin\pandoc.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    raise FileNotFoundError("Pandoc was not found. Set PANDOC or install pandoc.")


def run_pandoc(
    pandoc: str,
    normalized_markdown: Path,
    template: Path,
    report_dir: Path,
    content_docx: Path,
    title: str,
) -> None:
    command = [
        pandoc,
        str(normalized_markdown),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        "--to=docx",
        "--standalone",
        "--reference-doc",
        str(template),
        "--resource-path",
        str(report_dir),
        "--dpi=300",
        "--metadata",
        f"title={title}",
        "--output",
        str(content_docx),
    ]
    emit("Running Pandoc for native tables, formulas, and images.")
    subprocess.run(command, cwd=report_dir, check=True)


def clear_paragraph(paragraph) -> None:
    paragraph_xml = paragraph._p
    for child in list(paragraph_xml):
        if child.tag != qn("w:pPr"):
            paragraph_xml.remove(child)


def append_field(paragraph, instruction: str, placeholder: str = "0"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction_text, separate, text, end))
    return run


def set_setting(settings_element, tag: str, value: str | None = None) -> None:
    element = settings_element.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        settings_element.append(element)
    if value is not None:
        element.set(qn("w:val"), value)


def mark_first_row_as_header(table) -> None:
    if not table.rows:
        return
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        row_properties.append(OxmlElement("w:tblHeader"))


def postprocess_content_docx(
    content_docx: Path,
    expected_figures: int,
    expected_tables: int,
) -> dict[str, int]:
    document = Document(str(content_docx))
    style_names = {style.name for style in document.styles}
    caption_style = "图表标题" if "图表标题" in style_names else "Caption"
    heading_style = "Heading 1"
    if heading_style not in style_names:
        raise ValueError("The template does not expose the required Heading 1 style.")

    chapter_count = 0
    figure_count = 0
    table_count = 0

    for paragraph in document.paragraphs:
        if "<w:drawing" in paragraph._p.xml:
            paragraph.style = document.styles[caption_style]

        text = paragraph.text.strip()
        if paragraph.style.name == heading_style:
            if text == "摘要":
                paragraph.style = document.styles["Normal"]
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.bold = True
                continue

            chapter_count += 1
            chapter_title = CHAPTER_PREFIX_RE.sub("", text).strip()
            clear_paragraph(paragraph)
            chapter_field = append_field(paragraph, " SEQ Chapter \\* ARABIC ")
            # The template's Heading 1 style already renders the visible
            # Chinese chapter number. Keep this field only for captions.
            chapter_field.font.hidden = True
            paragraph.add_run(" " + chapter_title)
            continue

        if paragraph.style.name in {"Heading 2", "Heading 3"}:
            subheading_title = SUBHEADING_PREFIX_RE.sub("", text).strip()
            if subheading_title != text:
                clear_paragraph(paragraph)
                paragraph.add_run(subheading_title)
            continue

        caption_match = CAPTION_RE.match(text)
        if not caption_match:
            continue

        caption_type, caption_title = caption_match.groups()
        clear_paragraph(paragraph)
        paragraph.style = document.styles[caption_style]
        paragraph.alignment = 1
        if caption_type == "图":
            figure_count += 1
            sequence_name = "Figure"
        else:
            table_count += 1
            sequence_name = "Table"
        paragraph.add_run(caption_type + " ")
        append_field(paragraph, " SEQ Chapter \\c ")
        paragraph.add_run("-")
        append_field(paragraph, f" SEQ {sequence_name} \\* ARABIC \\s 1 ")
        paragraph.add_run("　" + caption_title)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        mark_first_row_as_header(table)

    set_setting(document.settings.element, "w:doNotCompressPictures")
    set_setting(document.settings.element, "w:updateFields", "true")
    document.save(str(content_docx))

    if figure_count != expected_figures or table_count != expected_tables:
        raise ValueError(
            "Caption conversion mismatch: "
            f"expected figures={expected_figures}, tables={expected_tables}; "
            f"converted figures={figure_count}, tables={table_count}."
        )
    return {
        "chapters": chapter_count,
        "figures": figure_count,
        "tables": table_count,
    }


def finalize_with_word(
    content_docx: Path,
    output: Path,
    pdf_output: Path | None,
) -> dict[str, int]:
    """Save the reference-document based file through Word and update fields."""
    import pythoncom
    import win32com.client

    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists():
        output.unlink()
    if pdf_output and pdf_output.exists():
        pdf_output.unlink()

    pythoncom.CoInitialize()
    app = None
    document = None
    original_picture_setting = None
    try:
        app = win32com.client.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        app.ScreenUpdating = False
        try:
            original_picture_setting = app.Options.DoNotCompressPicturesInFile
            app.Options.DoNotCompressPicturesInFile = True
        except Exception:
            original_picture_setting = None

        document = app.Documents.Open(
            str(content_docx), ReadOnly=True, AddToRecentFiles=False, Visible=False
        )
        document.SaveAs2(str(output), FileFormat=16)
        document.Fields.Update()
        for section in document.Sections:
            for collection in (section.Headers, section.Footers):
                for index in (1, 2, 3):
                    try:
                        collection(index).Range.Fields.Update()
                    except Exception:
                        pass
        document.Repaginate()
        pages = int(document.ComputeStatistics(2))
        field_count = int(document.Fields.Count)
        document.Save()
        if pdf_output:
            document.ExportAsFixedFormat(
                OutputFileName=str(pdf_output), ExportFormat=17, OpenAfterExport=False
            )
        return {"word_pages": pages, "word_body_fields": field_count}
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if app is not None:
            if original_picture_setting is not None:
                try:
                    app.Options.DoNotCompressPicturesInFile = original_picture_setting
                except Exception:
                    pass
            try:
                app.Quit()
            except Exception:
                # ExportAsFixedFormat can terminate Word before the COM cleanup
                # call. The output has already been saved at this point.
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def force_document_settings(output: Path) -> None:
    """Keep no-compression and update-fields settings after the Word save."""
    settings_path = "word/settings.xml"
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(output, "r") as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    root = etree.fromstring(files[settings_path])
    do_not_compress = root.find(f"{{{namespace}}}doNotCompressPictures")
    if do_not_compress is None:
        do_not_compress = etree.Element(f"{{{namespace}}}doNotCompressPictures")
        root.append(do_not_compress)
    update_fields = root.find(f"{{{namespace}}}updateFields")
    if update_fields is None:
        update_fields = etree.Element(f"{{{namespace}}}updateFields")
        root.append(update_fields)
    update_fields.set(f"{{{namespace}}}val", "true")
    files[settings_path] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    temporary = output.with_suffix(output.suffix + ".settings-tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    temporary.replace(output)


def inspect_docx(output: Path) -> dict[str, int | bool]:
    with zipfile.ZipFile(output, "r") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        media_files = [info for info in archive.infolist() if info.filename.startswith("word/media/")]
    return {
        "drawing_count": document_xml.count("<w:drawing"),
        "native_table_count": document_xml.count("<w:tbl>"),
        "office_math_count": document_xml.count("<m:oMath"),
        "chapter_field_count": document_xml.count("SEQ Chapter"),
        "figure_field_count": document_xml.count("SEQ Figure"),
        "table_field_count": document_xml.count("SEQ Table"),
        "media_file_count": len(media_files),
        "media_bytes": sum(info.file_size for info in media_files),
        "no_picture_compression": "doNotCompressPictures" in settings_xml,
        "update_fields_on_open": "updateFields" in settings_xml,
    }


def verify_embedded_images(
    source_dir: Path, source_image_paths: Iterable[str], output: Path
) -> dict[str, int]:
    source_hashes = {
        hashlib.sha256((source_dir / image).read_bytes()).hexdigest()
        for image in source_image_paths
    }
    with zipfile.ZipFile(output, "r") as archive:
        embedded_hashes = {
            hashlib.sha256(archive.read(info.filename)).hexdigest()
            for info in archive.infolist()
            if info.filename.startswith("word/media/")
        }
    return {
        "source_hash_count": len(source_hashes),
        "embedded_hash_count": len(embedded_hashes),
        "missing_source_hash_count": len(source_hashes - embedded_hashes),
        "extra_embedded_hash_count": len(embedded_hashes - source_hashes),
    }


def build_manifest(
    work_dir: Path,
    source: Path,
    template: Path,
    output: Path,
    title: str,
    source_image_paths: Iterable[str],
    source_caption_counts: Counter[str],
    source_latex_count: int,
    converted_counts: dict[str, int],
    word_counts: dict[str, int],
    package_counts: dict[str, int | bool],
    image_embedding: dict[str, int],
    pdf_output: Path | None,
) -> Path:
    image_list = list(source_image_paths)
    manifest = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "title": title,
        "source": str(source),
        "template": str(template),
        "output": str(output),
        "output_bytes": output.stat().st_size,
        "pdf_output": str(pdf_output) if pdf_output else None,
        "pdf_bytes": pdf_output.stat().st_size if pdf_output and pdf_output.exists() else None,
        "source_image_references": len(image_list),
        "source_unique_image_references": len(set(image_list)),
        "source_latex_blocks": source_latex_count,
        "source_figure_captions": source_caption_counts["figure"],
        "source_table_captions": source_caption_counts["table"],
        "converted": converted_counts,
        "word": word_counts,
        "package": package_counts,
        "image_embedding": image_embedding,
    }
    manifest_path = work_dir / "build_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest_path


def main() -> int:
    args = parse_args()
    source = args.source.resolve()
    template = args.template.resolve()
    output = args.output.resolve()
    work_dir = args.work_dir.resolve()
    pdf_output = work_dir / (output.stem + ".pdf") if args.export_pdf else None
    normalized_markdown = work_dir / "normalized_report.md"
    content_docx = work_dir / "pandoc_content.docx"

    if not source.is_file():
        raise FileNotFoundError(source)
    if not template.is_file():
        raise FileNotFoundError(template)
    work_dir.mkdir(parents=True, exist_ok=True)

    title, source_markdown = read_source(source)
    source_images = image_paths(source_markdown)
    source_captions = caption_counts(source_markdown)
    missing_images = [
        image for image in source_images if not (source.parent / Path(image)).is_file()
    ]
    if missing_images:
        preview = "; ".join(missing_images[:5])
        raise FileNotFoundError(f"Missing report images ({len(missing_images)}): {preview}")

    prepared_markdown = make_pandoc_markdown(source_markdown)
    normalized_markdown.write_text(prepared_markdown, encoding="utf-8")
    emit(
        "Validated "
        f"{len(source_images)} image references, {source_captions['figure']} figure captions, "
        f"{source_captions['table']} table captions, and {len(LATEX_FENCE_RE.findall(source_markdown))} LaTex blocks."
    )

    run_pandoc(
        find_pandoc(), normalized_markdown, template, source.parent, content_docx, title
    )
    converted_counts = postprocess_content_docx(
        content_docx, source_captions["figure"], source_captions["table"]
    )
    word_counts = finalize_with_word(content_docx, output, pdf_output)
    force_document_settings(output)
    package_counts = inspect_docx(output)
    image_embedding = verify_embedded_images(source.parent, source_images, output)

    errors: list[str] = []
    if package_counts["drawing_count"] != len(source_images):
        errors.append(
            f"drawing count {package_counts['drawing_count']} != source images {len(source_images)}"
        )
    if package_counts["native_table_count"] < source_captions["table"]:
        errors.append(
            f"native table count {package_counts['native_table_count']} < captions {source_captions['table']}"
        )
    if package_counts["office_math_count"] < len(LATEX_FENCE_RE.findall(source_markdown)):
        errors.append(
            f"Office Math count {package_counts['office_math_count']} < source LaTex blocks"
        )
    if package_counts["figure_field_count"] != source_captions["figure"]:
        errors.append("Figure field count does not match figure captions")
    if package_counts["table_field_count"] != source_captions["table"]:
        errors.append("Table field count does not match table captions")
    if package_counts["chapter_field_count"] != converted_counts["chapters"] + source_captions["figure"] + source_captions["table"]:
        errors.append("Chapter field count does not match chapters plus captions")
    if not package_counts["no_picture_compression"]:
        errors.append("Output does not contain the no-picture-compression setting")
    if not package_counts["update_fields_on_open"]:
        errors.append("Output does not request field updates on open")
    if image_embedding["missing_source_hash_count"]:
        errors.append("One or more source images were not embedded byte-for-byte")

    manifest_path = build_manifest(
        work_dir,
        source,
        template,
        output,
        title,
        source_images,
        source_captions,
        len(LATEX_FENCE_RE.findall(source_markdown)),
        converted_counts,
        word_counts,
        package_counts,
        image_embedding,
        pdf_output,
    )
    emit(f"Wrote {output} ({output.stat().st_size} bytes).")
    emit(f"Wrote {manifest_path}.")
    if errors:
        for error in errors:
            emit("CHECK FAILED: " + error)
        return 2
    emit("Structural checks passed.")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:
        emit(f"FAILED: {error}")
        raise
