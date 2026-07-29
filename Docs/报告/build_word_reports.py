from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\HP\Desktop\MoSim")
REPORT_DIR = ROOT / "Docs" / "报告"
FIG_DIR = REPORT_DIR / "图"
TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
OUT_REPORT = REPORT_DIR / "MoSim_仿真分析报告.docx"
OUT_MANUAL = REPORT_DIR / "MoSim_用户手册.docx"


def clear_doc(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_base_style(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def _append_fld_char(run, kind: str) -> None:
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    run._r.append(fld)


def add_seq_field(paragraph, name: str) -> None:
    begin_run = paragraph.add_run()
    _append_fld_char(begin_run, "begin")

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {name} \\\\* ARABIC "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    _append_fld_char(separate_run, "separate")

    result_run = paragraph.add_run("1")

    end_run = paragraph.add_run()
    _append_fld_char(end_run, "end")


def fld_seq(run, name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {name} \\\\* ARABIC "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_plain_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)


def add_para(doc: Document, text: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if text else None
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def add_table_caption(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("表 ")
    add_seq_field(p, "Table")
    p.add_run(" " + title)


def add_figure_caption(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图 ")
    add_seq_field(p, "Figure")
    p.add_run(" " + title)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_table_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
    add_para(doc, "")


def add_image(doc: Document, path: Path, title: str, width_cm: float = 14.8) -> bool:
    if not path.exists():
        add_para(doc, f"图片文件缺失：{path}")
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(path), width=Cm(width_cm))
    except Exception as exc:
        add_para(doc, f"图片插入失败：{path}；原因：{exc}")
        return False
    add_figure_caption(doc, title)
    return True


def title_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    add_para(doc, "本文档由项目现有模型、脚本、图片和结果记录整理生成，文中图表均置于相关章节正文内。")


def latex_blocks(limit: int = 18) -> list[str]:
    src = (REPORT_DIR / "公式与推导.md").read_text(encoding="utf-8")
    blocks = re.findall(r"```latex\\n(.*?)\\n```", src, re.S)
    return [b.strip() for b in blocks[:limit]]


def controller_items() -> list[tuple[str, Path, Path]]:
    base = FIG_DIR / "控制器"
    items = []
    for model in sorted(base.glob("*/*/01_图形模型.png")):
        result = model.parent / "02_仿真结果.png"
        if result.exists():
            name = model.parent.name
            items.append((name, model, result))
    return items


def build_report() -> None:
    doc = Document(str(TEMPLATE))
    clear_doc(doc)
    set_base_style(doc)
    title_page(doc, "MoSim四旋翼无人机位姿控制系统仿真分析报告", "仿真分析报告")

    add_plain_heading(doc, "摘要", 1)
    add_para(doc, "本文围绕四旋翼无人机位姿控制、扰动抑制、故障容错、编队控制和工程部署验证，构建了以MWORKS为建模与控制器设计核心、以ROS1/Sunray/Gazebo/PX4/MAVROS/px4ctrl为部署验证环境、以RViz/QGC/UE为显示和审核界面的仿真分析体系。系统采用模块化控制器接口，将基线PID、改进PID、滑模、MPC、INDI、扰动观测、安全监督、故障检测隔离与容错控制等路线纳入统一实验框架，并通过图形化模型、固定输入响应、MIL/SIL一致性、Gazebo部署和三机编队场景形成分层证据。")
    add_para(doc, "报告重点说明系统总体方案、四旋翼动力学、控制器族设计、实验场景、关键结果与工程边界。对于已经完成验证的路线，正文给出模型、曲线、指标和图像证据；对于未闭合路线，正文直接说明阻塞原因，避免将理论设计冒充实验通过。")
    add_para(doc, "关键词：MWORKS；四旋翼无人机；位姿控制；鲁棒控制；故障容错；代码生成；联合仿真")

    add_plain_heading(doc, "第1章 绪论", 1)
    add_plain_heading(doc, "1.1 研究背景与工程意义", 2)
    add_para(doc, "四旋翼无人机具有结构紧凑、机动性强、任务适配灵活等特点，但其欠驱动、强耦合、非线性和易受扰动影响的动力学特征，使高精度位姿控制和工程部署验证成为系统设计的关键问题。单纯离线仿真容易忽略执行器饱和、坐标系转换、消息时延、故障注入和外部运行环境差异，因此本项目采用模型化设计、自动代码生成和联合仿真相结合的方式，将控制器从MWORKS模型逐步推进到可部署验证链路。")
    add_plain_heading(doc, "1.2 赛题要求与实现内容", 2)
    add_para(doc, "赛题要求可以归纳为位姿控制优化、典型轨迹与鲁棒性验证、编队控制、模块化扩展和可复现交付五类任务。本文的章节组织和图表安排围绕这些任务展开。")
    add_table(doc, "赛题要求与本文实现内容对应关系", ["赛题要求", "对应章节", "核心实现", "主要图表", "边界说明"], [
        ["位姿控制优化", "第5至第7章", "PID族、滑模族、MPC族、INDI/NDI、经典控制补充和智能增强路线", "控制器模型图、响应图、控制族表", "不将阻塞路线写成通过路线"],
        ["典型轨迹与鲁棒性", "第8至第10章", "悬停、阶跃、8字、螺旋、风扰、参数摄动和故障场景", "评价指标表、扰动故障曲线", "按实际通过、阻塞、未运行分别说明"],
        ["编队控制", "第10章", "三机OpenBlocks场景、参考轨迹跟踪和距离/净空检查", "三机模型、动画、闭环曲线", "为离线全局规划参考跟踪，不宣称在线自主探索"],
        ["模块化与扩展", "第4章、第11至第13章", "统一控制输出、Profile配置、代码生成、ROS回灌和多界面显示", "接口表、部署链路图", "显示界面不替代控制闭环证据"],
    ])

    add_plain_heading(doc, "第2章 系统总体方案与闭环开发流程", 1)
    add_para(doc, "系统按照建模、控制、运行、感知规划和显示五个层级组织。MWORKS承担模型建模、控制器组合、MIL/SIL和代码生成的主线职责；ROS1/Sunray/Gazebo/PX4/MAVROS/px4ctrl承担部署验证；RViz、QGC和UE用于点云、轨迹、飞行状态和场景展示。")
    add_image(doc, FIG_DIR / "手绘架构" / "16_MWORKS建模仿真代码生成反馈闭环.png", "MWORKS建模、仿真、代码生成与反馈优化闭环")
    add_para(doc, "上图展示从需求、建模、控制器仿真、代码生成到部署反馈的闭环。该闭环的核心是将外部部署中暴露的问题回灌到MWORKS场景，而不是把不同仿真软件的截图简单拼接为结论。")
    add_image(doc, FIG_DIR / "手绘架构" / "20_实验平台分层与故障反馈链路.png", "实验平台分层与故障反馈链路")
    add_para(doc, "平台分层后，每个工具只承担本层证据：建模证据、运行证据、显示证据和报告证据分开记录，避免显示画面越权代表控制性能。")
    add_image(doc, FIG_DIR / "手绘架构" / "15_ModelStudio_QGC_Gazebo联合仿真数据流.png", "Model Studio、QGC与Gazebo联合仿真数据流")
    add_para(doc, "Model Studio负责实验配置与入口组织，QGC用于飞控状态和任务界面，Gazebo与PX4/MAVROS提供部署运行环境。")

    add_plain_heading(doc, "第3章 四旋翼无人机多领域建模", 1)
    add_para(doc, "四旋翼模型包括机体六自由度运动、旋翼推力与反扭矩、电机电调动态、传感器、机载计算机、飞控接口以及结果查看器动画。")
    add_image(doc, FIG_DIR / "手绘架构" / "18_坐标系与FrameContract转换关系.png", "坐标系与Frame Contract转换关系")
    add_para(doc, "坐标系转换是MWORKS离线模型、ROS消息、PX4状态和RViz显示之间保持一致的基础。")
    add_image(doc, FIG_DIR / "手绘架构" / "19_四旋翼动力学与控制分配模型.png", "四旋翼动力学与控制分配模型")
    add_para(doc, "动力学模型以总推力、三轴力矩和旋翼转速为主要接口，控制分配模块将高层控制量转换为电机侧命令。")
    add_table(doc, "坐标系、状态量与符号定义", ["符号", "含义", "坐标系", "单位"], [
        ["p=[x,y,z]^T", "位置向量", "世界系", "m"],
        ["v", "线速度", "世界系或机体系", "m/s"],
        ["phi, theta, psi", "滚转、俯仰、偏航角", "姿态参数", "rad"],
        ["omega=[p,q,r]^T", "机体系角速度", "机体系", "rad/s"],
        ["T", "总推力", "机体系z轴方向", "N"],
        ["tau", "三轴控制力矩", "机体系", "N·m"],
    ])
    add_plain_heading(doc, "3.1 关键公式", 2)
    add_para(doc, "以下公式用于说明模型与控制律中的主要数学关系，正文采用LaTeX文本直接给出，便于后续在Word中转换为公式对象。")
    for i, block in enumerate(latex_blocks(16), 1):
        add_code(doc, f"公式({i})  {block}")

    add_plain_heading(doc, "第4章 模块化控制平台与统一接口", 1)
    add_para(doc, "控制平台采用统一输入、统一输出和可组合增强模块结构，使不同控制器能够在相同状态、参考、故障和评价接口下运行。")
    add_image(doc, FIG_DIR / "手绘架构" / "02_控制器族分层与统一接口.png", "控制器族分层与统一接口结构")
    add_para(doc, "控制器族被分为基线控制、现代控制、鲁棒控制、扰动补偿、安全监督、故障容错和编队控制等层级。")
    add_image(doc, FIG_DIR / "手绘架构" / "11_控制输出层级与Runner边界.png", "控制输出层级与Runner边界")
    add_table(doc, "控制器输出边界与适用Runner", ["输出类型", "含义", "适用场景", "边界"], [
        ["ATTITUDE_THRUST", "姿态加总推力命令", "PX4姿态内环可用场景", "不直接控制电机"],
        ["BODY_RATE_THRUST", "角速度加总推力命令", "角速度内环验证", "需明确坐标系和饱和"],
        ["WRENCH", "总推力与三轴力矩", "离线动力学和控制分配验证", "需经过分配矩阵"],
        ["ROTOR_COMMAND", "单电机命令", "故障注入和控制分配重构", "必须经过执行器约束检查"],
    ])
    add_image(doc, FIG_DIR / "手绘架构" / "12_Profile配置与状态注入链路.png", "Profile配置、状态源与事件注入链路")
    add_para(doc, "Profile将控制器、场景、无人机数量、地图、故障、风扰和终端状态组合成可复现实验入口。")
    add_image(doc, FIG_DIR / "APP" / "01_在线建模验证.png", "Model Studio在线建模验证界面")
    add_para(doc, "在线建模验证界面用于选择实验配置并打开对应MWORKS模型。")

    add_plain_heading(doc, "第5章 基线控制与PID族改进", 1)
    add_para(doc, "PID族控制器承担基线验证和工程可解释性对照任务。改进路线包括抗饱和、前馈、增益调度、模糊调节和神经网络辅助调参。")
    add_image(doc, FIG_DIR / "手绘架构" / "10_PID_INDI与NMPC_INDI增强结构.png", "PID-INDI与NMPC-INDI增强控制结构")
    add_para(doc, "PID-INDI将传统误差反馈与增量非线性动态逆补偿结合，用于增强扰动下的响应速度和抗模型误差能力。")

    add_plain_heading(doc, "第6章 现代、鲁棒与非线性控制算法", 1)
    add_para(doc, "本章将全部已归档控制器图形模型和仿真结果直接纳入正文。每个控制器先给出图形化模型，再给出固定输入或代表场景响应，用于证明该路线在当前项目中有明确实现和可审阅结果。")
    add_image(doc, FIG_DIR / "手绘架构" / "09_控制算法家族技术路线.png", "控制算法家族技术路线")
    add_image(doc, FIG_DIR / "手绘架构" / "04_滑模控制族分类与适用场景.png", "滑模控制族分类与适用场景")
    add_table(doc, "现代控制器设计特征与证据状态", ["类别", "代表控制器", "设计特征", "正文证据"], [
        ["PID族", "cascade_pid、gain_scheduled_pid等", "结构简单、工程可解释", "模型图和响应图"],
        ["滑模族", "terminal_smc、super_twisting_smc等", "鲁棒性强、需处理抖振", "模型图和响应图"],
        ["MPC族", "linear_mpc、tube_mpc、mppi等", "显式处理约束和预测", "模型图和响应图"],
        ["经典补充", "LQR、LQI、SO(3)、MRAC、NDI等", "覆盖线性、非线性和自适应路线", "模型图和响应图"],
    ])
    items = controller_items()
    add_para(doc, f"当前正文共插入{len(items)}条控制器路线，每条路线包含图形化模型和仿真结果两张图。")
    for idx, (name, model, result) in enumerate(items, 1):
        add_plain_heading(doc, f"6.{idx} {name} 控制器图形模型与仿真结果", 3)
        add_para(doc, f"{name}路线用于覆盖控制器族中的一个具体实现。下图给出该路线在Sysblock/MWORKS中的图形化模型，重点检查输入输出端口、控制律模块和运行接口是否完整。")
        add_image(doc, model, f"{name}控制器图形模型", 13.8)
        add_para(doc, f"从模型结构可以看出，{name}已被纳入统一控制器接口，并通过固定输入或代表场景进行响应检查。")
        add_image(doc, result, f"{name}控制器仿真结果", 13.8)
        add_para(doc, f"{name}结果图用于支撑该路线已形成可审阅仿真输出。若后续用于竞赛结论，还需结合对应指标表判断其是否达到验收门限。")

    add_plain_heading(doc, "第7章 智能增强、安全与故障容错设计", 1)
    add_para(doc, "智能增强、安全和故障容错模块作为控制器外层能力，分别处理扰动估计、增益调度、安全约束、故障检测隔离和控制分配重构。")
    add_image(doc, FIG_DIR / "手绘架构" / "17_DOB_ESO_L1_AWFF增强补偿结构.png", "DOB、ESO、L1与AWFF增强补偿结构")
    add_image(doc, FIG_DIR / "手绘架构" / "14_神经残差与强化学习增益调度流程.png", "神经残差与强化学习增益调度流程")
    add_image(doc, FIG_DIR / "手绘架构" / "05_安全监督器与降级处置流程.png", "安全监督器与降级处置流程")
    add_image(doc, FIG_DIR / "手绘架构" / "01_故障注入与FTC闭环链路.png", "单电机效率下降与FTC闭环链路")
    add_table(doc, "增强、安全与FTC模块状态", ["模块", "功能", "验证内容", "写作边界"], [
        ["DOB/ESO/L1/AWFF", "扰动估计与补偿", "补偿结构和代表响应", "不替代主控制器验收"],
        ["Neural Residual", "残差补偿", "冻结推理和回退机制", "实验路线，不作为默认控制器"],
        ["RL Gain Scheduler", "增益调度", "有界调度和回退机制", "实验路线，不声明全面优越"],
        ["Safety Supervisor", "安全约束与降级", "触发、约束整形和动作状态", "只在触发条件内成立"],
        ["FDI/FTC", "故障隔离与控制重构", "效率下降、隔离掩码和重构动作", "不声明完整电机失效恢复"],
    ])

    add_plain_heading(doc, "第8章 仿真实验设计与评价方法", 1)
    add_table(doc, "实验平台配置", ["项目", "配置", "用途"], [
        ["建模与控制器", "MWORKS/Sysplorer/Sysblock/Syslab", "图形化模型、MIL、SIL、代码生成"],
        ["运行环境", "Ubuntu-20.04、ROS1 Noetic、Gazebo Classic、PX4、MAVROS、px4ctrl", "部署验证和运行日志"],
        ["显示审核", "RViz、QGC、UE、Model Studio", "点云、轨迹、状态和视频展示"],
        ["结果目录", "Results、Docs/报告/图、Config、Models", "指标、图片、配置和模型资产"],
    ])
    add_table(doc, "统一评价指标及计算窗口", ["指标", "含义", "用途", "通过解释"], [
        ["RMSE", "参考与实际轨迹均方根误差", "跟踪精度评价", "越小越好，需结合场景门限"],
        ["最大误差", "窗口内最大偏差", "瞬态安全性", "反映扰动或故障峰值"],
        ["最小机间距", "多机之间最小实际距离", "编队安全", "大于安全距离才可通过"],
        ["净空下界", "与障碍物距离下界", "避障安全", "大于给定门限才可通过"],
        ["落地/解锁状态", "任务终端状态", "工程完成性", "必须与运行日志一致"],
    ])

    add_plain_heading(doc, "第9章 MWORKS整机闭环仿真与性能对比", 1)
    add_para(doc, "整机闭环仿真用于检查控制器在多领域模型中的状态响应、控制输入、约束触发和任务终端状态。")
    add_image(doc, FIG_DIR / "云纵150" / "Sunray150_三维装配与部件审查联系表.png", "云纵150无人机模型与场景图")
    add_image(doc, FIG_DIR / "控制器" / "06_核心对比路线" / "official_pid" / "02_仿真结果.png", "Official PID固定输入下的误差与推力响应")
    add_image(doc, FIG_DIR / "控制器" / "01_PID族" / "gain_scheduled_pid" / "02_仿真结果.png", "增益调度PID代表响应")
    add_table(doc, "控制器族统一性能对比", ["控制器族", "正文证据", "主要价值", "说明"], [
        ["Official PID", "模型图、结果图", "官方基线", "作为对照基准"],
        ["Gain Scheduled PID", "模型图、结果图、A/B矩阵", "调度增强", "工程候选路线"],
        ["SMC/INDI/MPC", "模型图、结果图", "鲁棒与约束处理", "按场景门限判断"],
        ["Learning Routes", "模型图、结果图", "智能补偿探索", "实验路线，需保留回退"],
    ])

    add_plain_heading(doc, "第10章 扰动、故障、安全与编队场景验证", 1)
    add_para(doc, "扰动、故障、安全和编队场景用于验证系统在非理想条件下的可解释性和止损能力。")
    add_image(doc, FIG_DIR / "三机编队" / "01_OpenBlocks三机整机图形模型.png", "OpenBlocks复杂地图三机可重构编队整机模型")
    add_para(doc, "该模型包含三架无人机、复杂障碍场景、参考轨迹和跟踪控制链路。")
    add_image(doc, FIG_DIR / "手绘架构" / "13_OpenBlocks三机编队场景与轨迹.png", "OpenBlocks三机编队场景与轨迹")
    add_para(doc, "该手绘图说明离线全局规划参考、三机跟踪和安全距离检查之间的关系。")
    add_image(doc, FIG_DIR / "三机编队" / "03_OpenBlocks三机原生动画.png", "MWORKS结果查看器中的OpenBlocks三机原生动画")
    add_para(doc, "动画图用于展示三机在场景中的空间运动关系。")
    add_image(doc, FIG_DIR / "三机编队" / "04_OpenBlocks三机闭环曲线.png", "三机跟踪误差、最小机间距离与障碍物净空下界")
    add_table(doc, "扩展场景验证结果", ["场景", "结果", "关键指标", "说明"], [
        ["OpenBlocks三机", "通过当前门限", "运行时长304.8405 s，RMSE 0.1368/0.1162/0.0963 m，最小机间距1.0909 m，净空下界0.0942 m", "离线全局A*与EGO平滑参考跟踪"],
        ["风扰场景", "完成代表验证", "轨迹偏差和恢复过程入图", "用于鲁棒性说明"],
        ["单电机效率下降", "完成有界FTC验证", "效率下降、FDI和重构链路入图", "不声明完整电机失效"],
    ])

    add_plain_heading(doc, "第11章 实时联合仿真与实验交互", 1)
    add_image(doc, FIG_DIR / "APP" / "02_实时联合仿真.png", "Model Studio实时联合仿真界面")
    add_image(doc, FIG_DIR / "QGC" / "2bf9bb6d9a0fe79cfd6c1c48d399af62.png", "QGC、UE场景和飞行状态综合界面")
    add_para(doc, "实时联合仿真界面用于展示控制器、飞控、场景和状态之间的协作关系。")
    add_table(doc, "实时联合仿真关键接口", ["接口", "方向", "内容", "用途"], [
        ["状态输入", "运行环境到控制器", "位置、速度、姿态、角速度", "控制器计算"],
        ["控制输出", "控制器到飞控/执行器", "姿态、角速度、推力或电机命令", "闭环控制"],
        ["事件输入", "配置到运行环境", "风扰、故障、任务状态", "场景复现"],
        ["显示输出", "运行环境到界面", "轨迹、点云、状态、日志", "审核和录制"],
    ])

    add_plain_heading(doc, "第12章 自动代码生成与部署验证", 1)
    add_image(doc, FIG_DIR / "手绘架构" / "08_代码生成与ROS回灌部署流程.png", "MWORKS代码生成与ROS回灌部署流程")
    add_image(doc, FIG_DIR / "APP" / "03_生成代码部署.png", "Model Studio生成代码部署界面")
    add_image(doc, FIG_DIR / "rviz" / "8字.png", "Sunray/MID360与FAST-LIO点云建图结果")
    add_image(doc, FIG_DIR / "rviz" / "diff.png", "点云地图到Grid3D栅格地图转换结果")
    add_image(doc, FIG_DIR / "FUEL" / "01_FUEL点云地图与探索轨迹.png", "FUEL点云地图与探索轨迹")
    add_image(doc, FIG_DIR / "FUEL" / "06_FUEL_Gazebo与RViz联合画面.png", "FUEL Gazebo与RViz联合画面")
    add_table(doc, "代表控制器MIL/SIL一致性结果", ["对象", "检查内容", "结果写法"], [
        ["生成C控制器", "输入输出列、固定步长和差值", "通过零差或记录最大差值"],
        ["px4ctrl部署", "构建、启动、控制器选择", "说明部署状态和终端状态"],
        ["Gazebo/PX4", "起飞、悬停、降落和日志", "不把截图单独作为性能验收"],
    ])

    add_plain_heading(doc, "第13章 部署问题反馈与控制器再优化", 1)
    add_para(doc, "部署环境暴露的主要问题包括坐标系约定不一致、执行器故障注入窗口不稳定、控制器门限过严、显示界面与控制证据边界容易混淆等。项目通过Profile、接口边界和证据分类将这些问题回灌到MWORKS模型与实验配置。")
    add_table(doc, "部署问题与MWORKS等效复现场景", ["问题", "MWORKS复现方式", "优化方向", "文档写法"], [
        ["坐标系错配", "Frame Contract和Adapter转换测试", "冻结字段语义", "说明ENU/NED/FLU/FRD边界"],
        ["故障注入窗口不稳定", "效率因子和FDI/FTC场景", "固定注入时刻和终端状态", "写有界验证"],
        ["控制门限未通过", "同场景重跑和指标记录", "调参或止损", "区分accepted和blocked"],
        ["显示证据越权", "将RViz/QGC/UE仅作为显示层", "绑定运行日志", "不把界面截图写成控制成功"],
    ])

    add_plain_heading(doc, "第14章 总结与展望", 1)
    add_para(doc, "本文完成了MoSim四旋翼无人机位姿控制系统的模型化设计、控制器族构建、扰动故障安全扩展、三机编队验证、代码生成部署链路和报告图表整理。系统的主要价值在于形成了从MWORKS建模到外部部署反馈再回到MWORKS优化的闭环，而不是停留在单一离线仿真。")
    add_para(doc, "后续工作包括进一步闭合未通过控制器路线、增强在线规划与实时避障、完善UE全局态势显示、压实Gazebo/PX4长期运行门限，并将手工审核步骤逐步转化为自动化检查。")
    add_plain_heading(doc, "参考文献", 1)
    add_para(doc, "文献条目由人工最终整理。")
    doc.save(str(OUT_REPORT))


def build_manual() -> None:
    doc = Document(str(TEMPLATE))
    clear_doc(doc)
    set_base_style(doc)
    title_page(doc, "MoSim四旋翼无人机位姿控制与仿真平台用户手册", "用户手册")

    add_plain_heading(doc, "前言", 1)
    add_para(doc, "本手册面向项目评审、复现实验和演示录制人员，说明MoSim系统的安装环境、项目入口、模型运行、结果查看、代码生成、联合仿真和常见问题处理方法。技术推导和大规模性能比较见《仿真分析报告》。")

    add_plain_heading(doc, "第1章 系统概述", 1)
    add_para(doc, "MoSim以MWORKS为建模与控制器验证主线，支持四旋翼模型仿真、控制器组合、扰动故障实验、代码生成和工程部署验证。")
    add_table(doc, "系统功能与使用入口", ["功能", "主要软件", "操作入口", "输出结果", "当前状态"], [
        ["在线建模验证", "MWORKS/Model Studio", "Model Studio在线建模验证页", "Result.msr、曲线、动画", "可用于报告演示"],
        ["实时联合仿真", "ROS1/Gazebo/PX4/QGC", "Model Studio实时联合仿真页", "状态、事件、日志、截图", "用于部署链路展示"],
        ["生成代码部署", "MWORKS/px4ctrl", "生成代码部署页", "C/C++、构建与运行结果", "用于代码生成闭环"],
    ])
    add_image(doc, FIG_DIR / "手绘架构" / "20_实验平台分层与故障反馈链路.png", "MoSim系统总体组成")
    add_para(doc, "上图给出系统各层职责。用户操作时应先确认当前目标属于建模验证、联合仿真还是生成代码部署。")
    add_image(doc, FIG_DIR / "手绘架构" / "16_MWORKS建模仿真代码生成反馈闭环.png", "MoSim典型操作流程")

    add_plain_heading(doc, "第2章 软件环境与安装配置", 1)
    add_table(doc, "推荐软硬件环境", ["类别", "最低要求", "推荐配置", "说明"], [
        ["操作系统", "Windows 10/11 + Ubuntu-20.04", "Windows 11 + Ubuntu-20.04 WSL/独立环境", "ROS1路线使用Ubuntu-20.04"],
        ["CPU/内存/GPU", "4核CPU、16GB内存、独立或集成GPU", "8核以上CPU、32GB内存、NVIDIA GPU", "复杂场景和UE显示更依赖GPU"],
        ["MWORKS", "可打开Sysplorer/Sysblock/Syslab", "授权正常且能运行模型", "登录和授权异常先处理许可"],
        ["ROS环境", "ROS1 Noetic、Gazebo Classic", "含PX4、MAVROS、px4ctrl、RViz", "不要误用ROS2/x500旧路线"],
    ])
    add_para(doc, "如果MWORKS出现登录、许可或授权窗口，应先完成授权检查，再继续模型运行。不要在授权状态未知时解释为模型错误。")
    add_table(doc, "环境检查项目与通过现象", ["检查项", "命令或操作", "正常现象", "异常处理"], [
        ["MWORKS启动", "打开Sysplorer或Model Studio调用MWORKS", "主窗口正常显示且可打开模型", "先处理登录/许可/授权"],
        ["ROS1环境", "wsl -d Ubuntu-20.04 并运行ROS1预检脚本", "ROS/Gazebo/PX4路径正确", "不要切换到ROS2路线"],
        ["结果目录", "检查Results和Docs/报告/图", "能看到日志、截图、图表", "缺图时先补截图再写结论"],
        ["代码生成", "在Model Studio选择生成代码部署", "生成C/C++产物或明确错误", "失败时保留日志"],
    ])

    add_plain_heading(doc, "第3章 项目文件与模型入口", 1)
    add_table(doc, "MWORKS模型目录说明", ["目录", "主要内容", "是否直接修改", "用户入口"], [
        ["Models", "项目自有MWORKS模型和实验模型", "按任务修改", "Sysplorer/Model Studio"],
        ["Config", "Experiment Profile、场景和配置", "按实验修改", "Model Studio配置页"],
        ["Scripts", "质量检查、代码生成、运行和报告脚本", "谨慎修改", "命令行或APP按钮"],
        ["Results", "日志、指标、manifest、图片和审查结果", "不手改权威结果", "报告和复现依据"],
    ])
    add_table(doc, "录制与复现实验的推荐入口", ["实验", "Profile或模型", "场景", "建议时长", "结果位置"], [
        ["Official PID基线", "official_pid", "悬停/8字", "按模型默认", "Results/control_platform与报告图"],
        ["推荐控制器", "gain_scheduled_pid", "悬停/扰动", "按模型默认", "A/B矩阵和控制器图"],
        ["故障容错", "fdi_ftc_family", "单电机效率下降", "按故障脚本默认", "FTC结果目录和报告图"],
        ["三机编队", "OpenBlocks三机模型", "90 m x 60 m地图", "304.8405 s运行记录", "planning/three_uav_open_blocks_mworks_20260720"],
    ])

    add_plain_heading(doc, "第4章 Model Studio快速入门", 1)
    add_para(doc, "在Julia或Syslab环境中启动APP：")
    add_code(doc, r'include(raw"C:\Users\HP\Desktop\MoSim\apps\model_studio\src\app.jl")')
    add_image(doc, FIG_DIR / "APP" / "01_在线建模验证.png", "MoSim Model Studio主界面")
    add_para(doc, "主界面按在线建模验证、实时联合仿真和生成代码部署三类入口组织。")
    add_image(doc, FIG_DIR / "APP" / "02_实时联合仿真.png", "实时联合仿真入口")
    add_image(doc, FIG_DIR / "APP" / "03_生成代码部署.png", "生成代码部署入口")

    add_plain_heading(doc, "第5章 在线建模验证", 1)
    add_para(doc, "在线建模验证用于打开MWORKS模型、运行仿真、查看Result.msr、保存结果曲线和三维动画。")
    add_image(doc, FIG_DIR / "云纵150" / "Sunray150_三维装配与部件审查联系表.png", "从Model Studio打开MWORKS仿真模型")
    add_image(doc, FIG_DIR / "控制器" / "06_核心对比路线" / "official_pid" / "02_仿真结果.png", "MWORKS仿真结果查看示例")
    add_para(doc, "运行后应检查曲线是否非空、时间轴是否完整、终端状态是否合理。仅有窗口打开不等于仿真成功。")

    add_plain_heading(doc, "第6章 基础场景复现", 1)
    add_table(doc, "基础场景复现检查表", ["场景", "操作", "通过现象", "失败处理", "结果位置"], [
        ["起飞悬停", "选择基线或推荐控制器运行", "姿态稳定、误差曲线收敛", "检查参数和坐标系", "Results/control_platform"],
        ["阶跃轨迹", "切换参考Profile", "响应无明显发散", "降低速度或检查饱和", "Results/control_platform"],
        ["8字轨迹", "选择8字场景", "轨迹闭合且误差可解释", "检查参考生成和控制器门限", "Results/control_platform"],
        ["螺旋爬升", "选择螺旋场景", "高度和水平轨迹同时跟踪", "检查推力余量", "Results/control_platform"],
    ])

    add_plain_heading(doc, "第7章 风扰、故障、安全与编队操作", 1)
    add_image(doc, FIG_DIR / "手绘架构" / "01_故障注入与FTC闭环链路.png", "扰动、故障与编队场景操作逻辑")
    add_para(doc, "故障操作应先确认故障类型、注入时刻和恢复策略。单电机效率下降实验只代表有界效率下降，不代表完整电机失效。")
    add_image(doc, FIG_DIR / "三机编队" / "03_OpenBlocks三机原生动画.png", "三机编队仿真动画")

    add_plain_heading(doc, "第8章 实时联合仿真", 1)
    add_table(doc, "实时联合仿真状态说明", ["状态", "含义", "用户动作", "注意事项"], [
        ["未连接", "QGC/PX4/MAVROS未形成有效链路", "检查启动顺序", "不要直接判断控制器失败"],
        ["已连接", "状态和模式可见", "执行起飞或任务", "保留截图和日志"],
        ["运行中", "飞行和状态持续更新", "观察轨迹、点云和事件", "等待必须有时间上限"],
        ["异常/停止", "模式、位置或日志异常", "停止任务并保存证据", "避免无界等待"],
    ])
    add_image(doc, FIG_DIR / "QGC" / "2bf9bb6d9a0fe79cfd6c1c48d399af62.png", "QGC飞行界面")

    add_plain_heading(doc, "第9章 生成代码与部署", 1)
    add_para(doc, "生成代码流程从MWORKS模型出发，导出C/C++控制器产物，经过MIL/SIL一致性检查后再进入px4ctrl部署验证。")
    add_image(doc, FIG_DIR / "手绘架构" / "08_代码生成与ROS回灌部署流程.png", "控制器代码生成操作流程")
    add_table(doc, "代码生成与部署检查表", ["步骤", "检查内容", "正常结果", "失败处理"], [
        ["选择目标", "控制器和输出类型", "目标与Runner匹配", "检查输出边界表"],
        ["生成代码", "C/C++产物和日志", "生成成功且文件完整", "保存错误日志"],
        ["MIL/SIL", "相同输入下输出差值", "差值在设定门限内", "检查步长和列顺序"],
        ["部署运行", "px4ctrl构建和启动", "启动并完成任务", "检查ROS/PX4链路"],
    ])

    add_plain_heading(doc, "第10章 感知、规划与显示", 1)
    add_image(doc, FIG_DIR / "rviz" / "8字.png", "FAST-LIO点云与定位查看")
    add_image(doc, FIG_DIR / "rviz" / "diff.png", "栅格地图与Diff规划显示")
    add_image(doc, FIG_DIR / "FUEL" / "06_FUEL_Gazebo与RViz联合画面.png", "Gazebo与RViz联合显示")
    add_para(doc, "RViz用于点云、轨迹、地图和Frame检查；UE和QGC用于显示和演示，不替代控制闭环指标。")

    add_plain_heading(doc, "第11章 参数与接口说明", 1)
    add_table(doc, "Experiment Profile关键字段", ["字段", "含义", "示例", "说明"], [
        ["controller_id", "控制器编号", "gain_scheduled_pid", "选择控制器路线"],
        ["scenario", "场景", "hover/figure8/wind/fault", "决定参考和扰动"],
        ["uav_count", "无人机数量", "1或3", "编队场景需为3"],
        ["fault_profile", "故障配置", "motor1_efficiency_0_65", "仅在故障实验使用"],
        ["output_type", "控制输出类型", "ATTITUDE_THRUST", "必须匹配Runner"],
    ])
    add_table(doc, "控制器输入输出接口", ["接口", "输入", "输出", "说明"], [
        ["State", "位置、速度、姿态、角速度", "控制器状态量", "坐标系必须明确"],
        ["Reference", "轨迹、速度、加速度、偏航", "目标状态", "用于误差计算"],
        ["Controller", "State和Reference", "控制命令", "按输出类型解释"],
        ["Metrics", "参考、实际、控制量", "RMSE、最大误差等", "用于报告结论"],
    ])
    add_table(doc, "故障与扰动命令字段", ["字段", "作用", "使用建议", "风险"], [
        ["wind_vector", "注入风扰", "按场景给定", "过大可能导致发散"],
        ["motor_efficiency", "设置电机效率", "用于FTC实验", "不能写成完整失效"],
        ["safety_mode", "启用安全策略", "先低风险场景验证", "触发条件需记录"],
        ["stop_condition", "终止条件", "设置时间和状态上限", "避免长时间卡死"],
    ])

    add_plain_heading(doc, "第12章 常见问题与故障排查", 1)
    add_table(doc, "常见问题处理表", ["问题", "可能原因", "处理方法", "不要做"], [
        ["APP无法启动", "Julia依赖或路径错误", "检查启动命令和项目路径", "不要重装无关环境"],
        ["MWORKS未打开", "授权、窗口或路径问题", "先看可见窗口和授权状态", "不要直接判定模型坏"],
        ["飞机翻滚", "坐标系、参数或控制器饱和", "回到基线场景缩小问题", "不要盲目换架构"],
        ["Result.msr不可见", "仿真未完成或路径不对", "检查结果目录和日志", "不要用空截图当结果"],
        ["Gazebo异常", "ROS/PX4/MAVROS链路问题", "按预检脚本检查", "不要切到旧ROS2路线"],
    ])

    add_plain_heading(doc, "第13章 提交与复现检查", 1)
    add_para(doc, "提交前应保证模型、配置、脚本、结果、报告和视频路径清晰，评委能够根据手册复现最小演示流程。")
    add_table(doc, "最终提交内容检查表", ["内容", "建议路径", "是否必须", "检查方式"], [
        ["MWORKS模型和依赖", "Models", "是", "能从Model Studio或Sysplorer打开"],
        ["用户手册DOCX/PDF", "Docs/报告", "是", "步骤可复现"],
        ["仿真分析报告DOCX/PDF", "Docs/报告", "是", "图表和正文一致"],
        ["演示视频", "Results/submission或指定提交目录", "是", "画面与报告结论一致"],
        ["结果与图片", "Results、Docs/报告/图", "建议提交", "能追溯到模型或运行过程"],
    ])
    add_plain_heading(doc, "参考资料", 1)
    add_para(doc, "参考资料由人工最终整理。")
    doc.save(str(OUT_MANUAL))


if __name__ == "__main__":
    build_report()
    build_manual()
    print(OUT_REPORT)
    print(OUT_MANUAL)
