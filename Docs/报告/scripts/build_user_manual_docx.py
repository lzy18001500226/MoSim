"""Build the MoSim user manual from its Markdown source and the retained template.

The builder intentionally keeps the competition template as the document
container.  It converts Markdown tables into native Word tables, embeds each
source image as its original file bytes, and emits Word SEQ fields for every
figure and table caption.
"""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "Docs" / "报告"
SOURCE = REPORT_DIR / "用户手册_正文骨架.md"
TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
OUTPUT = REPORT_DIR / "MoSim_用户手册.docx"
QA_DIR = REPORT_DIR / "审计" / "用户手册_word_20260802"
MANIFEST = QA_DIR / "build_manifest.json"
HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:[一二三四五六七八九十百千万]+[、.．]|\d+(?:\.\d+){0,3}[、.．]?)\s*"
)

@dataclass
class Block:
    kind: str
    data: Any


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def plain_markdown(text: str) -> str:
    """Keep visible Markdown content while removing lightweight delimiters."""
    text = text.strip()
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    return text


def inline_runs(text: str) -> list[tuple[str, str]]:
    """Split code/bold/link fragments so they stay legible in Word."""
    pattern = r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))"
    runs: list[tuple[str, str]] = []
    for token in re.split(pattern, text):
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            runs.append(("code", token[1:-1]))
        elif token.startswith("**") and token.endswith("**"):
            runs.append(("bold", token[2:-2]))
        elif token.startswith("["):
            match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if match:
                runs.append(("link", f"{match.group(1)}（{match.group(2)}）"))
            else:
                runs.append(("text", token))
        else:
            runs.append(("text", token))
    return runs


def parse_caption(line: str) -> tuple[str, str, str] | None:
    match = re.match(r"^(图|表)\s*(\d+(?:-\d+)*)\s*[　 ]+(.+?)\s*$", line)
    if not match:
        return None
    return match.group(1), match.group(2), plain_markdown(match.group(3))


def parse_table_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [plain_markdown(cell.replace("\\|", "|").strip()) for cell in re.split(r"(?<!\\)\|", text)]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_markdown(source: Path) -> tuple[str, list[Block]]:
    lines = source.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    paragraph_lines: list[str] = []
    pending_image: tuple[str, Path] | None = None
    pending_table_caption: tuple[str, str, str] | None = None
    title: str | None = None
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(part.strip() for part in paragraph_lines).strip()
            if text:
                blocks.append(Block("paragraph", text))
        paragraph_lines = []

    def flush_image() -> None:
        nonlocal pending_image
        if pending_image:
            alt, image_path = pending_image
            raise ValueError(f"Figure caption is missing for: {image_path} ({alt})")
        pending_image = None

    def flush_table_caption() -> None:
        nonlocal pending_table_caption
        if pending_table_caption:
            label, number, caption_title = pending_table_caption
            blocks.append(Block("paragraph", f"{label} {number}　{caption_title}"))
        pending_table_caption = None

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if line.startswith("```"):
            flush_paragraph()
            flush_image()
            flush_table_caption()
            language = line[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index == len(lines):
                raise ValueError(f"Unclosed fenced block in {source}")
            blocks.append(Block("code", {"language": language, "text": "\n".join(code_lines)}))
            index += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            flush_image()
            flush_table_caption()
            pending_image = (image_match.group(1), (source.parent / image_match.group(2)).resolve())
            index += 1
            continue

        caption = parse_caption(stripped)
        if caption:
            flush_paragraph()
            label, number, caption_title = caption
            if label == "图":
                if pending_image is None:
                    blocks.append(Block("paragraph", stripped))
                else:
                    alt, image_path = pending_image
                    blocks.append(
                        Block(
                            "figure",
                            {"alt": alt, "path": image_path, "number": number, "title": caption_title},
                        )
                    )
                    pending_image = None
            else:
                flush_image()
                flush_table_caption()
                pending_table_caption = caption
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            flush_paragraph()
            flush_image()
            flush_table_caption()
            level = len(heading.group(1))
            text = plain_markdown(heading.group(2))
            if level == 1 and title is None:
                title = text
            else:
                blocks.append(Block("heading", {"level": level, "text": text}))
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            flush_image()
            if pending_table_caption is None:
                raise ValueError(f"Word table has no preceding table caption near line {index + 1}")
            _, number, caption_title = pending_table_caption
            headers = parse_table_row(line)
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            blocks.append(Block("table", {"number": number, "title": caption_title, "headers": headers, "rows": rows}))
            pending_table_caption = None
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if pending_image:
            flush_image()
        if pending_table_caption:
            flush_table_caption()
        if stripped == "---":
            flush_paragraph()
            index += 1
            continue
        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    flush_image()
    flush_table_caption()
    if title is None:
        raise ValueError(f"Document title is missing from {source}")
    return title, blocks


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size: Pt | None = None, bold: bool | None = None) -> None:
    """Apply only intentional emphasis; ordinary content inherits template styles."""
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def append_runs(paragraph, text: str, size: Pt | None = None) -> None:
    for kind, value in inline_runs(text):
        run = paragraph.add_run(value)
        if kind == "bold":
            set_run_font(run, size=size, bold=True)
        elif size is not None:
            set_run_font(run, size=size)


def append_field(paragraph, instruction: str, cached_text: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    instruction_run._r.append(instruction_element)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    paragraph.add_run(cached_text)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_update_fields_and_no_compression(doc: Document) -> None:
    settings = doc.settings.element
    for tag in ("w:updateFields", "w:doNotCompressPictures"):
        existing = settings.find(qn(tag))
        if existing is not None:
            settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    settings.append(OxmlElement("w:doNotCompressPictures"))


def enforce_output_settings(output: Path) -> None:
    """Keep the field-refresh and picture-preservation flags after saving."""
    temporary = output.with_name(f"{output.stem}.settings-patched.docx")
    with zipfile.ZipFile(output, "r") as source_archive, zipfile.ZipFile(temporary, "w") as target_archive:
        for member in source_archive.infolist():
            content = source_archive.read(member.filename)
            if member.filename == "word/settings.xml":
                settings_xml = content.decode("utf-8")
                settings_xml = re.sub(r"<w:updateFields\b[^>]*/>", "", settings_xml)
                settings_xml = re.sub(r"<w:doNotCompressPictures\b[^>]*/>", "", settings_xml)
                addition = '<w:updateFields w:val="true"/><w:doNotCompressPictures/>'
                if "</w:settings>" not in settings_xml:
                    raise RuntimeError("word/settings.xml is missing its closing settings element")
                content = settings_xml.replace("</w:settings>", addition + "</w:settings>").encode("utf-8")
            target_archive.writestr(member, content)
    temporary.replace(output)


def restore_template_page_chrome(output: Path) -> None:
    """Restore retained template headers and footers byte-for-byte.

    python-docx rewrites untouched header/footer XML during a document save.
    The body references retain the same part names, so replacing only these
    preserve-only parts keeps the template page furniture exactly intact.
    """
    temporary = output.with_name(f"{output.stem}.chrome-patched.docx")
    with zipfile.ZipFile(TEMPLATE, "r") as template_archive, zipfile.ZipFile(output, "r") as output_archive, zipfile.ZipFile(temporary, "w") as target_archive:
        preserve_names = {
            name
            for name in template_archive.namelist()
            if name.startswith(("word/header", "word/footer", "word/_rels/header", "word/_rels/footer"))
        }
        for member in output_archive.infolist():
            content = output_archive.read(member.filename)
            if member.filename in preserve_names:
                content = template_archive.read(member.filename)
            target_archive.writestr(member, content)
    temporary.replace(output)


def add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(72)
    paragraph.paragraph_format.space_after = Pt(24)
    run = paragraph.add_run(title)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(20)
    run.bold = True

    subtitle = doc.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("用户手册")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    run.bold = True
    doc.add_page_break()


def add_heading(doc: Document, text: str, markdown_level: int) -> None:
    level = max(1, min(3, markdown_level - 1))
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(HEADING_PREFIX_RE.sub("", text).strip())


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if re.match(r"^\([0-9a-zA-Z]+\)", text):
        paragraph.paragraph_format.first_line_indent = Cm(0)
    append_runs(paragraph, text)


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    for line_index, line in enumerate(text.splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def add_caption(doc: Document, label: str, source_number: str, title: str) -> None:
    style_name = "图表标题" if "图表标题" in [style.name for style in doc.styles] else "Caption"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    chapter_match = re.fullmatch(r"(\d+)-(\d+)", source_number)
    if chapter_match:
        chapter, serial = chapter_match.groups()
        literal = f"{label} {chapter}-"
        instruction = f"SEQ {label}_{chapter} \\* ARABIC"
        cached_text = serial
    else:
        literal = f"{label} "
        instruction = f"SEQ {label} \\* ARABIC"
        cached_text = source_number
    run = paragraph.add_run(literal)
    append_field(paragraph, instruction, cached_text)
    paragraph.add_run(f"　{title}")


def set_cell_width(cell, width: int) -> None:
    cell.width = Emu(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width / 635)))
    tc_width.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def table_column_widths(doc: Document, headers: list[str], rows: list[list[str]]) -> list[int]:
    section = doc.sections[0]
    usable = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    weights = []
    for index, header in enumerate(headers):
        longest = max([len(header)] + [len(row[index]) for row in rows])
        weights.append(max(7, min(longest, 38)))
    total = sum(weights)
    widths = [max(int(Cm(1.25)), int(usable * weight / total)) for weight in weights]
    widths[-1] += usable - sum(widths)
    return widths


def set_cell_text(cell, text: str, header: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    append_runs(paragraph, text)


def add_native_table(doc: Document, table_data: dict[str, Any]) -> None:
    headers = table_data["headers"]
    rows = table_data["rows"]
    if not headers or not rows:
        raise ValueError("Markdown table cannot be empty")
    if any(len(row) != len(headers) for row in rows):
        raise ValueError(f"Table {table_data['number']} has inconsistent column counts")
    add_caption(doc, "表", table_data["number"], table_data["title"])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for index, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(cell, text, header=True)
    for source_row in rows:
        row = table.add_row()
        for index, (cell, text) in enumerate(zip(row.cells, source_row)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, text)
    doc.add_paragraph(style="Normal")


def image_display_size(doc: Document, path: Path) -> tuple[Emu, Emu]:
    image = DocxImage.from_file(str(path))
    width, height = int(image.width), int(image.height)
    section = doc.sections[0]
    max_width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    max_height = int(Cm(20.0))
    scale = min(1.0, max_width / width, max_height / height)
    return Emu(int(width * scale)), Emu(int(height * scale))


def add_figure(doc: Document, figure: dict[str, Any]) -> None:
    path = Path(figure["path"])
    if not path.is_file():
        raise FileNotFoundError(f"Figure source does not exist: {path}")
    style_name = "图表标题" if "图表标题" in [style.name for style in doc.styles] else "Caption"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    width, height = image_display_size(doc, path)
    paragraph.add_run().add_picture(str(path), width=width, height=height)
    add_caption(doc, "图", figure["number"], figure["title"])


def media_hashes(docx_path: Path) -> Counter[str]:
    with zipfile.ZipFile(docx_path) as archive:
        return Counter(
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        )


def inspect_output(output: Path, source_images: list[Path], expected_figures: int, expected_tables: int) -> dict[str, Any]:
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    output_images = media_hashes(output)
    source_hashes = Counter(sha256_file(path) for path in source_images)
    return {
        "figure_fields": document_xml.count("SEQ 图"),
        "table_fields": document_xml.count("SEQ 表"),
        "figure_fields_expected": expected_figures,
        "table_fields_expected": expected_tables,
        "native_table_count": document_xml.count("<w:tbl>"),
        "source_images": len(source_images),
        "source_image_hashes_preserved": all(output_images[key] >= value for key, value in source_hashes.items()),
        "do_not_compress_pictures": "w:doNotCompressPictures" in settings_xml,
        "update_fields_on_open": "w:updateFields" in settings_xml,
    }


def build() -> dict[str, Any]:
    title, blocks = parse_markdown(SOURCE)
    figures = [block.data for block in blocks if block.kind == "figure"]
    tables = [block.data for block in blocks if block.kind == "table"]
    source_images = [Path(figure["path"]) for figure in figures]
    if len(figures) != len(source_images) or not all(path.is_file() for path in source_images):
        raise ValueError("One or more source images are unavailable")

    doc = Document(str(TEMPLATE))
    clear_body(doc)
    set_update_fields_and_no_compression(doc)
    add_title(doc, title)
    for block in blocks:
        if block.kind == "heading":
            add_heading(doc, block.data["text"], block.data["level"])
        elif block.kind == "paragraph":
            add_paragraph(doc, block.data)
        elif block.kind == "code":
            add_code_block(doc, block.data["text"])
        elif block.kind == "figure":
            add_figure(doc, block.data)
        elif block.kind == "table":
            add_native_table(doc, block.data)
        else:
            raise ValueError(f"Unsupported block kind: {block.kind}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    enforce_output_settings(OUTPUT)
    restore_template_page_chrome(OUTPUT)
    inspection = inspect_output(OUTPUT, source_images, len(figures), len(tables))
    if inspection["figure_fields"] != len(figures):
        raise RuntimeError(f"Expected {len(figures)} figure fields, found {inspection['figure_fields']}")
    if inspection["table_fields"] != len(tables):
        raise RuntimeError(f"Expected {len(tables)} table fields, found {inspection['table_fields']}")
    if not inspection["source_image_hashes_preserved"]:
        raise RuntimeError("At least one embedded image does not match its source bytes")
    if not inspection["do_not_compress_pictures"]:
        raise RuntimeError("Output document is missing doNotCompressPictures")

    manifest = {
        "source_markdown": str(SOURCE),
        "source_sha256": sha256_file(SOURCE),
        "template_docx": str(TEMPLATE),
        "template_sha256": sha256_file(TEMPLATE),
        "output_docx": str(OUTPUT),
        "output_sha256": sha256_file(OUTPUT),
        "title": title,
        "blocks": len(blocks),
        "headings": sum(block.kind == "heading" for block in blocks),
        "figures": len(figures),
        "tables": len(tables),
        "inspection": inspection,
    }
    MANIFEST.write_bytes((json.dumps(manifest, ensure_ascii=False, indent=2) + "\n").encode("utf-8"))
    return manifest


if __name__ == "__main__":
    result = build()
    print(json.dumps(result, ensure_ascii=False, indent=2))
