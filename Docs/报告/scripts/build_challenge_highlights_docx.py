"""Build the competition-highlights Word draft from the Markdown source.

The builder intentionally starts from the retained national-competition
template. It only replaces the sample body, preserves the template section and
styles, emits native Word tables, inserts original image bytes, and uses Word
SEQ fields for continuous figure/table captions.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "Docs" / "报告"
SOURCE = REPORT_DIR / "挑战赛_参加理由及作品亮点.md"
TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
OUTPUT = REPORT_DIR / "挑战赛_参加理由及作品亮点.docx"
QA_DIR = ROOT / "Results" / "document_qa" / "challenge_highlights_word_20260802"
MANIFEST = QA_DIR / "build_manifest.json"
HEADING_PREFIX_RE = re.compile(
    r"^\s*(?:[一二三四五六七八九十百千万]+[、.．]|\d+(?:\.\d+){0,3}[、.．]?)\s*"
)


@dataclass
class Block:
    kind: str
    data: Any


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def strip_inline_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def split_inline_markdown(text: str) -> list[tuple[str, str]]:
    """Return text runs as (kind, value), with code and bold kept editable."""
    tokens = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    runs: list[tuple[str, str]] = []
    for token in tokens:
        if not token:
            continue
        if token.startswith("`") and token.endswith("`"):
            runs.append(("code", token[1:-1]))
        elif token.startswith("**") and token.endswith("**"):
            runs.append(("bold", token[2:-2]))
        else:
            runs.append(("text", token))
    return runs


def parse_caption(line: str) -> tuple[int, str] | None:
    match = re.match(r"^图\s*(\d+)\s*[　 ]+(.+?)\s*$", line)
    if not match:
        return None
    return int(match.group(1)), strip_inline_markdown(match.group(2))


def parse_table_row(line: str) -> list[str]:
    return [strip_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", line.strip()))


def parse_markdown(source: Path) -> list[Block]:
    lines = source.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    paragraph_lines: list[str] = []
    pending_image: tuple[str, Path] | None = None
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(part.strip() for part in paragraph_lines).strip()
            if text:
                blocks.append(Block("paragraph", text))
        paragraph_lines = []

    def flush_image() -> None:
        nonlocal pending_image
        if pending_image:
            alt, image_path = pending_image
            blocks.append(Block("figure", {"alt": alt, "path": image_path, "number": None, "title": alt}))
        pending_image = None

    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("```"):
            flush_paragraph()
            flush_image()
            language = line[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index == len(lines):
                raise ValueError(f"Unclosed fenced block in {source}")
            blocks.append(Block("code", {"language": language, "text": "\n".join(code_lines)}))
            index += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            flush_paragraph()
            flush_image()
            pending_image = (image_match.group(1), (source.parent / image_match.group(2)).resolve())
            index += 1
            continue

        caption = parse_caption(line.strip())
        if caption:
            flush_paragraph()
            if pending_image is None:
                blocks.append(Block("paragraph", line.strip()))
            else:
                alt, image_path = pending_image
                blocks.append(
                    Block(
                        "figure",
                        {"alt": alt, "path": image_path, "number": caption[0], "title": caption[1]},
                    )
                )
                pending_image = None
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            flush_paragraph()
            flush_image()
            blocks.append(Block("heading", {"level": len(heading.group(1)), "text": strip_inline_markdown(heading.group(2))}))
            index += 1
            continue

        if line.strip().startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            flush_paragraph()
            flush_image()
            headers = parse_table_row(line)
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            blocks.append(Block("table", {"headers": headers, "rows": rows}))
            continue

        if not line.strip():
            flush_paragraph()
            index += 1
            continue

        if pending_image:
            flush_image()
        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    flush_image()
    return blocks


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size: Pt | None = None, bold: bool | None = None, code: bool = False) -> None:
    """Apply only intentional direct formatting; ordinary text uses template styles."""
    if code:
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def append_text_runs(paragraph, text: str, size: Pt | None = None) -> None:
    for kind, value in split_inline_markdown(text):
        run = paragraph.add_run(value)
        if kind == "code":
            set_run_font(run, size=Pt(9.5), code=True)
        elif kind == "bold":
            set_run_font(run, size=size, bold=True)
        elif size is not None:
            set_run_font(run, size=size)


def append_field(paragraph, instruction: str, cached_text: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    paragraph.add_run(cached_text)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_update_fields_and_no_compression(doc: Document) -> None:
    settings = doc.settings.element
    for tag in ("w:updateFields", "w:doNotCompressPictures"):
        existing = settings.find(qn(tag))
        if existing is not None:
            settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    settings.append(OxmlElement("w:doNotCompressPictures"))


def enforce_output_settings(output: Path) -> None:
    """Reapply document-level flags after Word has refreshed cached fields.

    Word 2007 evaluates the fields but drops these two newer document settings
    when saving. Patch only ``word/settings.xml`` and preserve every other
    package member unchanged.
    """
    temporary = output.with_name(f"{output.stem}.settings-patched.docx")
    with zipfile.ZipFile(output, "r") as source_archive, zipfile.ZipFile(temporary, "w") as target_archive:
        for member in source_archive.infolist():
            content = source_archive.read(member.filename)
            if member.filename == "word/settings.xml":
                settings_xml = content.decode("utf-8")
                settings_xml = re.sub(r"<w:updateFields\b[^>]*/>", "", settings_xml)
                settings_xml = re.sub(r"<w:doNotCompressPictures\b[^>]*/>", "", settings_xml)
                addition = '<w:updateFields w:val="true"/><w:doNotCompressPictures/>'
                if "</w:settings>" not in settings_xml:
                    raise RuntimeError("word/settings.xml is missing its closing settings element")
                content = settings_xml.replace("</w:settings>", addition + "</w:settings>").encode("utf-8")
            target_archive.writestr(member, content)
    temporary.replace(output)


def replace_footer_with_page_field(doc: Document) -> None:
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            footer.is_linked_to_previous = False
            for paragraph in list(footer.paragraphs):
                paragraph._element.getparent().remove(paragraph._element)
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run("第 ")
            set_run_font(run, size=Pt(9))
            append_field(paragraph, "PAGE", "1")
            run = paragraph.add_run(" 页")
            set_run_font(run, size=Pt(9))


def add_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    set_run_font(run, size=Pt(16), bold=True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=Pt(15), bold=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {max(1, min(3, level))}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(HEADING_PREFIX_RE.sub("", text).strip())


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if re.match(r"^\([0-9a-zA-Z]+\)", text):
        paragraph.paragraph_format.first_line_indent = Cm(0)
    append_text_runs(paragraph, text)


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.right_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    for line_index, line in enumerate(text.splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=Pt(9.5), code=True)


def add_caption(doc: Document, label: str, sequence: str, number: int, title: str) -> None:
    style_name = "图表标题" if "图表标题" in [style.name for style in doc.styles] else "Caption"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(f"{label} ")
    append_field(paragraph, f"SEQ {sequence} \\* ARABIC", str(number))
    paragraph.add_run(f"　{title}")


def set_cell_width(cell, width: int) -> None:
    cell.width = Emu(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width / 635)))
    tc_width.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def table_column_widths(doc: Document, column_count: int) -> list[int]:
    section = doc.sections[0]
    usable = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    if column_count == 2:
        return [int(usable * 0.31), usable - int(usable * 0.31)]
    base = usable // column_count
    return [base] * (column_count - 1) + [usable - base * (column_count - 1)]


def set_cell_text(cell, text: str, header: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    append_text_runs(paragraph, strip_inline_markdown(text))


def add_native_table(doc: Document, headers: list[str], rows: list[list[str]], caption_number: int) -> None:
    if not headers or not rows:
        raise ValueError("Markdown table cannot be empty")
    if any(len(row) != len(headers) for row in rows):
        raise ValueError("Markdown table rows do not match header column count")
    add_caption(doc, "表", "表", caption_number, "项目包内证据索引")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for index, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(cell, text, header=True)
    for source_row in rows:
        row = table.add_row()
        for index, (cell, text) in enumerate(zip(row.cells, source_row)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, text)
    doc.add_paragraph()


def image_display_size(doc: Document, path: Path) -> tuple[Emu, Emu]:
    image = DocxImage.from_file(str(path))
    width, height = int(image.width), int(image.height)
    section = doc.sections[0]
    max_width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    max_height = int(Cm(20.0))
    scale = min(1.0, max_width / width, max_height / height)
    return Emu(int(width * scale)), Emu(int(height * scale))


def add_figure(doc: Document, figure: dict[str, Any]) -> None:
    path = Path(figure["path"])
    if not path.is_file():
        raise FileNotFoundError(f"Figure source does not exist: {path}")
    number = figure["number"]
    if number is None:
        raise ValueError(f"Figure caption is missing for: {path}")
    style_name = "图表标题" if "图表标题" in [style.name for style in doc.styles] else "Caption"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    width, height = image_display_size(doc, path)
    paragraph.add_run().add_picture(str(path), width=width, height=height)
    add_caption(doc, "图", "图", number, figure["title"])


def output_media_hashes(docx_path: Path) -> Counter[str]:
    with zipfile.ZipFile(docx_path) as archive:
        return Counter(
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        )


def inspect_output(output: Path, source_images: list[Path], expected_figures: int, expected_tables: int) -> dict[str, Any]:
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        footer_xml = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    figure_fields = document_xml.count("SEQ 图 \\* ARABIC")
    table_fields = document_xml.count("SEQ 表 \\* ARABIC")
    page_fields = footer_xml.count(" PAGE ")
    source_hashes = Counter(sha256_file(path) for path in source_images)
    embedded_hashes = output_media_hashes(output)
    doc = Document(output)
    results = {
        "output_path": str(output),
        "output_sha256": sha256_file(output),
        "source_figure_count": expected_figures,
        "source_table_count": expected_tables,
        "native_word_table_count": len(doc.tables),
        "figure_seq_field_count": figure_fields,
        "table_seq_field_count": table_fields,
        "page_field_count": page_fields,
        "update_fields_on_open": "w:updateFields w:val=\"true\"" in settings_xml,
        "do_not_compress_pictures": "w:doNotCompressPictures" in settings_xml,
        "embedded_image_binary_hashes_match_source": source_hashes <= embedded_hashes,
        "embedded_media_count": sum(embedded_hashes.values()),
        "template_or_other_media_count": sum((embedded_hashes - source_hashes).values()),
    }
    errors: list[str] = []
    if figure_fields != expected_figures:
        errors.append(f"expected {expected_figures} figure SEQ fields, found {figure_fields}")
    if table_fields != expected_tables:
        errors.append(f"expected {expected_tables} table SEQ fields, found {table_fields}")
    if len(doc.tables) != expected_tables:
        errors.append(f"expected {expected_tables} native Word tables, found {len(doc.tables)}")
    if page_fields < 1:
        errors.append("expected at least one PAGE field in the template footer")
    if not results["update_fields_on_open"]:
        errors.append("w:updateFields was not enabled")
    if not results["do_not_compress_pictures"]:
        errors.append("w:doNotCompressPictures was not enabled")
    if not results["embedded_image_binary_hashes_match_source"]:
        errors.append("embedded media hashes do not match source image hashes")
    results["errors"] = errors
    results["passed"] = not errors
    return results


def build(template: Path, source: Path, output: Path, manifest: Path) -> dict[str, Any]:
    blocks = parse_markdown(source)
    figure_blocks = [block.data for block in blocks if block.kind == "figure"]
    table_blocks = [block.data for block in blocks if block.kind == "table"]
    if not figure_blocks or not table_blocks:
        raise ValueError("The source must contain at least one figure and one Markdown table")
    figure_numbers = [figure["number"] for figure in figure_blocks]
    if figure_numbers != list(range(1, len(figure_numbers) + 1)):
        raise ValueError(f"Figure captions are not continuous: {figure_numbers}")
    if len(table_blocks) != 1:
        raise ValueError(f"Expected exactly one source Markdown table, found {len(table_blocks)}")
    for figure in figure_blocks:
        if not Path(figure["path"]).is_file():
            raise FileNotFoundError(f"Missing source image: {figure['path']}")

    doc = Document(template)
    clear_body(doc)
    set_update_fields_and_no_compression(doc)
    replace_footer_with_page_field(doc)

    title_consumed = False
    subtitle_consumed = False
    table_counter = 0
    for block in blocks:
        if block.kind == "heading":
            level = block.data["level"]
            text = block.data["text"]
            if level == 1 and not title_consumed:
                title = text
                title_consumed = True
                continue
            if level == 2 and title_consumed and not subtitle_consumed:
                add_title(doc, title, text)
                subtitle_consumed = True
                continue
            if not subtitle_consumed:
                add_title(doc, title, "")
                subtitle_consumed = True
            add_heading(doc, text, 1 if level == 2 else 2 if level == 3 else 3)
        elif block.kind == "paragraph":
            add_paragraph(doc, block.data)
        elif block.kind == "code":
            add_code_block(doc, block.data["text"])
        elif block.kind == "figure":
            add_figure(doc, block.data)
        elif block.kind == "table":
            table_counter += 1
            add_native_table(doc, block.data["headers"], block.data["rows"], table_counter)
        else:
            raise ValueError(f"Unsupported Markdown block: {block.kind}")
    if not subtitle_consumed:
        add_title(doc, title if title_consumed else "MoSim", "")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    manifest.parent.mkdir(parents=True, exist_ok=True)
    inspection = inspect_output(
        output,
        [Path(figure["path"]) for figure in figure_blocks],
        expected_figures=len(figure_blocks),
        expected_tables=len(table_blocks),
    )
    manifest_data = {
        "source_markdown": str(source),
        "source_sha256": sha256_file(source),
        "template_docx": str(template),
        "template_sha256": sha256_file(template),
        "output_docx": str(output),
        "parsed_block_counts": dict(Counter(block.kind for block in blocks)),
        "figure_sources": [str(figure["path"]) for figure in figure_blocks],
        "inspection": inspection,
        "render_status": "not_run_libreoffice_unavailable_at_template_probe",
    }
    manifest.write_text(json.dumps(manifest_data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if not inspection["passed"]:
        raise RuntimeError("Structural DOCX inspection failed: " + "; ".join(inspection["errors"]))
    return manifest_data


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=TEMPLATE)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--manifest", type=Path, default=MANIFEST)
    parser.add_argument("--enforce-output-settings", action="store_true")
    args = parser.parse_args()
    if args.enforce_output_settings:
        enforce_output_settings(args.output.resolve())
        print(args.output.resolve())
        return
    result = build(args.template.resolve(), args.source.resolve(), args.output.resolve(), args.manifest.resolve())
    print(result["output_docx"])
    print(result["inspection"]["output_sha256"])


if __name__ == "__main__":
    main()
