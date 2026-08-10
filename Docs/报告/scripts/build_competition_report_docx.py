#!/usr/bin/env python3
"""Build the competition-report Word document from the current Markdown source.

Pandoc starts from the untouched national-competition template. A separate
formatted report is read only for heading, caption, formula-table, and
three-line-table presentation contracts; its body is never reused.
"""

from __future__ import annotations

import argparse
import base64
import copy
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


SCRIPT_PATH = Path(__file__).resolve()
REPORT_DIR = SCRIPT_PATH.parent.parent
REPO_ROOT = REPORT_DIR.parents[1]
DEFAULT_SOURCE = REPORT_DIR / "草稿" / "仿真分析报告_正文骨架.md"
DEFAULT_TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
DEFAULT_LAYOUT_REFERENCE = REPORT_DIR / "MoSim_仿真分析报告.docx"
DEFAULT_OUTPUT = REPORT_DIR / "MoSim_仿真分析报告_国赛版.docx"
DEFAULT_WORK_DIR = REPO_ROOT / "Results" / "docx_build" / "competition_report"
FORMULA_MANIFEST_BUILDER = REPO_ROOT / "Scripts" / "report" / "prepare_mathtype_formula_manifest.py"

CORE_PROPERTIES_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\n]+)\)")
CAPTION_RE = re.compile(
    r"^(?P<kind>图|表)\s*(?P<chapter>\d+)\s*[-—–]\s*"
    r"(?P<sequence>\d+)[　 \t]+(?P<title>.+?)\s*$"
)
FORMULA_TAG_RE = re.compile(
    r"\\tag\{(?P<chapter>\d+)-(?P<sequence>\d+)(?P<suffix>[a-z]?)\}"
)
CHAPTER_PREFIX_RE = re.compile(
    r"^\s*(?:[一二三四五六七八九十百千万]+[、.．]|\d+(?:\.\d+)*[、.．]?)\s*"
)
SUBHEADING_PREFIX_RE = re.compile(r"^\s*\d+(?:\.\d+){1,3}[、.．]?\s*")
LATEX_FENCE_RE = re.compile(r"```latex\s*\r?\n(.*?)\r?\n```", re.DOTALL)


@dataclass(frozen=True)
class CaptionRecord:
    kind: str
    chapter: int
    sequence: int
    title: str
    line_number: int

    @property
    def field_name(self) -> str:
        return "Figure" if self.kind == "图" else "Table"

    @property
    def label(self) -> str:
        return f"{self.kind} {self.chapter}-{self.sequence}"


@dataclass(frozen=True)
class FormulaTableLayout:
    table_style_name: str
    total_width: int
    left_width: int
    right_width: int
    table_look: dict[str, str]
    left_paragraph_properties: Any
    right_paragraph_properties: Any
    left_paragraph_contract: list[dict[str, Any]]
    right_paragraph_contract: list[dict[str, Any]]


@dataclass(frozen=True)
class TableCellTextLayout:
    table_style_name: str
    paragraph_properties: Any
    normal_size: int
    normal_size_cs: int
    table_size: int
    table_size_cs: int


@dataclass(frozen=True)
class HeadingPresentation:
    paragraph_properties: Any
    run_properties: Any


def emit(message: str) -> None:
    print(f"[competition-report] {message}", flush=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help="Untouched national-competition DOCX used as Pandoc's clean generation base.",
    )
    parser.add_argument(
        "--layout-reference",
        type=Path,
        default=DEFAULT_LAYOUT_REFERENCE,
        help="Formatted report read only for heading, table, and formula presentation contracts.",
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--work-dir", type=Path, default=DEFAULT_WORK_DIR)
    parser.add_argument(
        "--author",
        default=None,
        help="Fixed DOCX author; defaults to the clean template or layout-reference metadata instead of the local Office identity.",
    )
    parser.add_argument(
        "--export-pdf",
        action="store_true",
        help="Also export a PDF beside the build manifest for visual QA.",
    )
    parser.add_argument(
        "--formula-engine",
        choices=("omml",),
        default="omml",
        help=(
            "Create native Office Math in the one-row, two-column, borderless table "
            "geometry read from the reference MathType table."
        ),
    )
    return parser.parse_args()


def read_source(source: Path) -> tuple[str, str]:
    text = source.read_text(encoding="utf-8")
    title_match = re.search(r"(?m)^#\s+(.+?)\s*$", text)
    if not title_match:
        raise ValueError(f"No document title found in {source}")
    return title_match.group(1).strip(), text


def template_author(template: Path) -> str:
    """Read the approved template's creator for deterministic output metadata."""
    with zipfile.ZipFile(template, "r") as archive:
        root = etree.fromstring(archive.read("docProps/core.xml"))
    creator = root.find(f"{{{DC_NS}}}creator")
    return creator.text.strip() if creator is not None and creator.text else ""


def image_paths(markdown: str) -> list[str]:
    paths: list[str] = []
    for _alt, raw_path in IMAGE_RE.findall(markdown):
        path = raw_path.strip()
        if path.startswith("<") and path.endswith(">"):
            path = path[1:-1]
        paths.append(path)
    return paths


def resolve_report_image(source: Path, image: str) -> Path:
    """Resolve report assets from the source draft or the report root."""
    relative = Path(image)
    if relative.is_absolute():
        candidates = [relative]
    else:
        candidates = [source.parent / relative, REPORT_DIR / relative]

    for candidate in candidates:
        resolved = candidate.resolve()
        try:
            resolved.relative_to(REPO_ROOT)
        except ValueError:
            continue
        if resolved.is_file():
            return resolved
    raise FileNotFoundError(image)


def resolve_report_images(source: Path, images: Iterable[str]) -> list[Path]:
    missing: list[str] = []
    resolved: list[Path] = []
    for image in images:
        try:
            resolved.append(resolve_report_image(source, image))
        except FileNotFoundError:
            missing.append(image)
    if missing:
        preview = "; ".join(missing[:5])
        raise FileNotFoundError(f"Missing report images ({len(missing)}): {preview}")
    return resolved


def caption_records(markdown: str) -> list[CaptionRecord]:
    records: list[CaptionRecord] = []
    in_fence = False
    for line_number, line in enumerate(markdown.splitlines(), start=1):
        if line.lstrip().startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        match = CAPTION_RE.match(line.strip())
        if not match:
            continue
        records.append(
            CaptionRecord(
                kind=match.group("kind"),
                chapter=int(match.group("chapter")),
                sequence=int(match.group("sequence")),
                title=match.group("title"),
                line_number=line_number,
            )
        )
    return records


def caption_counts(markdown: str) -> Counter[str]:
    result: Counter[str] = Counter()
    for record in caption_records(markdown):
        result["figure" if record.kind == "图" else "table"] += 1
    return result


def source_primary_chapter_titles(markdown: str) -> list[str]:
    """Return numbered chapter titles after the document title and abstract."""
    titles: list[str] = []
    in_fence = False
    for line in markdown.splitlines():
        if line.lstrip().startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        match = re.match(r"^##\s+(.+?)\s*$", line)
        if not match:
            continue
        title = match.group(1).strip()
        if title == "摘要":
            continue
        titles.append(CHAPTER_PREFIX_RE.sub("", title).strip())
    return titles


def latex_block_to_display_math(match: re.Match[str]) -> str:
    body = match.group(1).strip()
    if body.startswith(r"\[") and body.endswith(r"\]"):
        body = body[2:-2].strip()
    # Equation labels are emitted as Word fields beside the formula, never as
    # literal TeX inside the math object.
    body = FORMULA_TAG_RE.sub("", body).strip()
    return f"\n$$\n{body}\n$$\n"


def shift_heading_levels(markdown: str) -> str:
    """Remove the Markdown title and promote ##/###/#### to H1/H2/H3."""
    lines = markdown.splitlines(keepends=True)
    output: list[str] = []
    first_title_removed = False
    in_fence = False

    for line in lines:
        if line.lstrip().startswith("```"):
            in_fence = not in_fence
            output.append(line)
            continue
        if not in_fence:
            match = re.match(r"^(#{1,6})\s+(.*?)(\r?\n)?$", line)
            if match:
                hashes, heading, newline = match.groups()
                if len(hashes) == 1 and not first_title_removed:
                    first_title_removed = True
                    continue
                if len(hashes) >= 2:
                    output.append("#" * (len(hashes) - 1) + " " + heading + (newline or ""))
                    continue
        output.append(line)
    return "".join(output)


def protect_parenthesized_markers(markdown: str) -> str:
    """Keep source markers such as ``(1)`` as plain text, not Pandoc lists."""
    output: list[str] = []
    in_fenced_block = False
    for line in markdown.splitlines(keepends=True):
        if line.lstrip().startswith("```"):
            in_fenced_block = not in_fenced_block
        elif not in_fenced_block:
            line = re.sub(r"^(\s*)\((\d+)\)(?=\s)", r"\1&#40;\2&#41;", line)
        output.append(line)
    return "".join(output)


def make_pandoc_markdown(source_markdown: str) -> str:
    prepared = LATEX_FENCE_RE.sub(latex_block_to_display_math, source_markdown)
    prepared = shift_heading_levels(prepared)
    prepared = protect_parenthesized_markers(prepared)

    # A lone Markdown image becomes a Pandoc figure using alt text as a second
    # caption.  Keep it as a pure image; the following project caption is the
    # authoritative caption and receives the Word fields later.
    def rewrite_image(match: re.Match[str]) -> str:
        path = match.group(2).strip()
        return f"![]({path}){{width=15cm}}"

    return IMAGE_RE.sub(rewrite_image, prepared)


def find_pandoc() -> str:
    candidates = [
        os.environ.get("PANDOC"),
        shutil.which("pandoc"),
        r"D:\Dev\Anaconda3\Library\bin\pandoc.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    raise FileNotFoundError("Pandoc was not found. Set PANDOC or install pandoc.")


def prepare_formula_manifest(source: Path, work_dir: Path) -> tuple[Path, list[dict[str, Any]]]:
    """Reuse the report-owned formula parser so source tags drive Word fields."""
    if not FORMULA_MANIFEST_BUILDER.is_file():
        raise FileNotFoundError(FORMULA_MANIFEST_BUILDER)

    manifest_path = work_dir / "mathtype_formula_manifest.json"
    command = [
        sys.executable,
        str(FORMULA_MANIFEST_BUILDER),
        "--source",
        str(source),
        "--output",
        str(manifest_path),
    ]
    emit("Preparing the tagged display-formula manifest.")
    subprocess.run(command, cwd=REPO_ROOT, check=True)
    payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    records = payload.get("formulas")
    if not isinstance(records, list) or not records:
        raise ValueError("Formula manifest contains no display-formula records.")
    return manifest_path, records


def run_pandoc(
    pandoc: str,
    normalized_markdown: Path,
    template: Path,
    resource_roots: Iterable[Path],
    content_docx: Path,
    title: str,
) -> None:
    resource_path = os.pathsep.join(
        str(path) for path in dict.fromkeys(path.resolve() for path in resource_roots)
    )
    command = [
        pandoc,
        str(normalized_markdown),
        "--from=markdown+tex_math_dollars+tex_math_single_backslash",
        "--to=docx",
        "--standalone",
        "--reference-doc",
        str(template),
        "--resource-path",
        resource_path,
        "--dpi=300",
        "--metadata",
        f"title={title}",
        "--output",
        str(content_docx),
    ]
    emit("Running Pandoc for native tables, formulas, and images.")
    subprocess.run(command, cwd=REPORT_DIR, check=True)


def clear_paragraph(paragraph) -> None:
    paragraph_xml = paragraph._p
    for child in list(paragraph_xml):
        if child.tag != qn("w:pPr"):
            paragraph_xml.remove(child)


def append_field(paragraph, instruction: str, placeholder: str = "0"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction_text, separate, text, end))
    return run


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
XML_NS = "http://www.w3.org/XML/1998/namespace"


def math_tag(local_name: str) -> str:
    return f"{{{MATH_NS}}}{local_name}"


def linearize_inline_math(
    node, vertical: str | None = None
) -> list[tuple[str, str | None]]:
    """Turn simple inline OMML into text chunks with sub/superscript marks."""
    tag = node.tag
    if tag == math_tag("t"):
        return [(node.text or "", vertical)]

    def child(local_name: str):
        return node.find(math_tag(local_name))

    if tag == math_tag("sSub"):
        base = child("e")
        subscript = child("sub")
        if base is not None and subscript is not None:
            return linearize_inline_math(base, vertical) + linearize_inline_math(
                subscript, "subscript"
            )
    elif tag == math_tag("sSup"):
        base = child("e")
        superscript = child("sup")
        if base is not None and superscript is not None:
            return linearize_inline_math(base, vertical) + linearize_inline_math(
                superscript, "superscript"
            )
    elif tag == math_tag("sSubSup"):
        base = child("e")
        subscript = child("sub")
        superscript = child("sup")
        if base is not None and subscript is not None and superscript is not None:
            return (
                linearize_inline_math(base, vertical)
                + linearize_inline_math(subscript, "subscript")
                + linearize_inline_math(superscript, "superscript")
            )
    elif tag == math_tag("f"):
        numerator = child("num")
        denominator = child("den")
        if numerator is not None and denominator is not None:
            return (
                linearize_inline_math(numerator, vertical)
                + [("/", vertical)]
                + linearize_inline_math(denominator, vertical)
            )
    elif tag == math_tag("d"):
        delimiter_properties = child("dPr")
        begin = ""
        end = ""
        if delimiter_properties is not None:
            begin_node = delimiter_properties.find(math_tag("begChr"))
            end_node = delimiter_properties.find(math_tag("endChr"))
            if begin_node is not None:
                begin = begin_node.get(math_tag("val"), "")
            if end_node is not None:
                end = end_node.get(math_tag("val"), "")
        chunks: list[tuple[str, str | None]] = []
        if begin:
            chunks.append((begin, vertical))
        for child_node in node:
            if child_node is not delimiter_properties:
                chunks.extend(linearize_inline_math(child_node, vertical))
        if end:
            chunks.append((end, vertical))
        return chunks

    # OMML properties do not contain visible text. For the remaining inline
    # constructs, preserve their text-bearing descendants in document order.
    chunks: list[tuple[str, str | None]] = []
    for child_node in node:
        chunks.extend(linearize_inline_math(child_node, vertical))
    return chunks


def merge_math_chunks(
    chunks: Iterable[tuple[str, str | None]]
) -> list[tuple[str, str | None]]:
    merged: list[tuple[str, str | None]] = []
    for text, vertical in chunks:
        if not text:
            continue
        if merged and merged[-1][1] == vertical:
            merged[-1] = (merged[-1][0] + text, vertical)
        else:
            merged.append((text, vertical))
    return merged


def make_body_text_run(text: str, vertical: str | None):
    run = OxmlElement("w:r")
    if vertical is not None:
        run_properties = OxmlElement("w:rPr")
        vertical_align = OxmlElement("w:vertAlign")
        vertical_align.set(qn("w:val"), vertical)
        run_properties.append(vertical_align)
        run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(f"{{{XML_NS}}}space", "preserve")
    run.append(text_element)
    return run


def convert_inline_math_to_body_text(document: Document) -> int:
    """Replace inline OMML with ordinary Word runs; keep display OMML intact."""
    converted = 0
    math_nodes = list(document.element.body.iter(math_tag("oMath")))
    for math_node in math_nodes:
        ancestor = math_node.getparent()
        if any(
            parent.tag == math_tag("oMathPara")
            for parent in iter_ancestors(ancestor)
        ):
            continue
        parent = math_node.getparent()
        if parent is None or parent.tag == qn("w:r"):
            continue
        chunks = merge_math_chunks(linearize_inline_math(math_node))
        if not chunks:
            continue
        index = parent.index(math_node)
        parent.remove(math_node)
        for text, vertical in reversed(chunks):
            parent.insert(index, make_body_text_run(text, vertical))
        converted += 1
    return converted


def iter_ancestors(node):
    while node is not None:
        yield node
        node = node.getparent()


def set_setting(settings_element, tag: str, value: str | None = None) -> None:
    element = settings_element.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        settings_element.append(element)
    if value is not None:
        element.set(qn("w:val"), value)


def mark_first_row_as_header(table) -> None:
    if not table.rows:
        return
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    if row_properties.find(qn("w:tblHeader")) is None:
        row_properties.append(OxmlElement("w:tblHeader"))


def find_style(document: Document, names: Iterable[str], style_type=None):
    wanted = {name.casefold() for name in names}
    for style in document.styles:
        if style.name.casefold() in wanted and (style_type is None or style.type == style_type):
            return style
    names_text = ", ".join(names)
    raise ValueError(f"The reference document does not expose the required style: {names_text}")


def xml_element_contract(element) -> dict[str, Any]:
    return {
        "tag": etree.QName(element).localname,
        "attributes": {
            etree.QName(key).localname: value for key, value in element.attrib.items()
        },
        "children": [xml_element_contract(child) for child in element],
    }


def paragraph_properties_contract(properties) -> list[dict[str, Any]]:
    if properties is None:
        return []
    return [
        xml_element_contract(child)
        for child in properties
        if child.tag not in {qn("w:pStyle"), qn("w:sectPr")}
    ]


def replace_paragraph_properties(paragraph, source_properties) -> None:
    paragraph_xml = paragraph._p
    current = paragraph_xml.pPr
    if current is not None:
        paragraph_xml.remove(current)
    if source_properties is None:
        return

    replacement = copy.deepcopy(source_properties)
    for child in list(replacement):
        if child.tag in {qn("w:pStyle"), qn("w:sectPr")}:
            replacement.remove(child)
    paragraph_xml.insert(0, replacement)


def replace_run_properties(run, source_properties) -> None:
    run_xml = run._r
    current = run_xml.rPr
    if current is not None:
        run_xml.remove(current)
    if source_properties is not None:
        run_xml.insert(0, copy.deepcopy(source_properties))


def heading_presentation(paragraph) -> HeadingPresentation:
    source_run = next((run for run in paragraph.runs if run.text), None)
    return HeadingPresentation(
        paragraph_properties=copy.deepcopy(paragraph._p.pPr)
        if paragraph._p.pPr is not None
        else None,
        run_properties=copy.deepcopy(source_run._r.rPr)
        if source_run is not None and source_run._r.rPr is not None
        else None,
    )


def reference_heading_presentations(template: Path) -> tuple[HeadingPresentation, HeadingPresentation]:
    """Read the title and abstract presentation from the formatted reference."""
    reference = Document(str(template))
    nonempty = [paragraph for paragraph in reference.paragraphs if paragraph.text.strip()]
    if not nonempty:
        raise ValueError("The layout reference has no main-story title paragraph.")
    title = nonempty[0]
    abstracts = [paragraph for paragraph in nonempty if paragraph.text.strip() == "摘要"]
    if len(abstracts) != 1:
        raise ValueError(
            "The layout reference must contain exactly one main-story 摘要 heading; "
            f"found {len(abstracts)}."
        )
    abstract = abstracts[0]
    if paragraph_outline_level(title) != 0 or paragraph_outline_level(abstract) != 0:
        raise ValueError("The reference title and 摘要 must use first-level heading presentation.")
    return heading_presentation(title), heading_presentation(abstract)


def apply_heading_presentation(paragraph, text: str, heading_style, presentation: HeadingPresentation) -> None:
    clear_paragraph(paragraph)
    replace_paragraph_properties(paragraph, presentation.paragraph_properties)
    paragraph.style = heading_style
    run = paragraph.add_run(text)
    replace_run_properties(run, presentation.run_properties)


def reference_formula_table_layout(template: Path) -> FormulaTableLayout:
    """Read the accepted MathType wrapper geometry from the layout reference."""
    reference = Document(str(template))
    candidates = [
        table
        for table in reference.tables
        if "Equation.DSMT4" in table._tbl.xml
    ]
    if len(candidates) != 1:
        raise ValueError(
            "The layout reference must contain exactly one MathType formula table; "
            f"found {len(candidates)}."
        )
    table = candidates[0]
    if len(table.rows) != 1 or len(table.columns) != 2:
        raise ValueError("The reference MathType formula table must be one row by two columns.")

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        raise ValueError("The reference MathType formula table has no table grid.")
    widths = [int(column.get(qn("w:w"), "0")) for column in grid.findall(qn("w:gridCol"))]
    if len(widths) != 2 or any(width <= 0 for width in widths):
        raise ValueError("The reference MathType formula table must define two positive grid widths.")

    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    total_width = int(table_width.get(qn("w:w"), "0")) if table_width is not None else 0
    if total_width <= 0:
        total_width = sum(widths)
    if total_width != sum(widths):
        raise ValueError("The reference MathType formula table width does not match its grid.")

    borders = table_properties.find(qn("w:tblBorders"))
    expected_sides = {"top", "left", "bottom", "right", "insideH", "insideV"}
    found_borders = (
        {
            etree.QName(border).localname: border.get(qn("w:val"), "")
            for border in borders
        }
        if borders is not None
        else {}
    )
    if set(found_borders) != expected_sides or any(
        value != "none" for value in found_borders.values()
    ):
        raise ValueError("The reference MathType formula table must be explicitly borderless.")

    look = table_properties.find(qn("w:tblLook"))
    if look is None:
        raise ValueError("The reference MathType formula table has no table look settings.")
    table_look = {etree.QName(key).localname: value for key, value in look.attrib.items()}
    left_paragraph = table.rows[0].cells[0].paragraphs[0]
    right_paragraph = table.rows[0].cells[1].paragraphs[0]
    if left_paragraph._p.pPr is None or right_paragraph._p.pPr is None:
        raise ValueError("The reference MathType formula table must define cell paragraph formatting.")
    return FormulaTableLayout(
        table_style_name=table.style.name,
        total_width=total_width,
        left_width=widths[0],
        right_width=widths[1],
        table_look=table_look,
        left_paragraph_properties=copy.deepcopy(left_paragraph._p.pPr),
        right_paragraph_properties=copy.deepcopy(right_paragraph._p.pPr),
        left_paragraph_contract=paragraph_properties_contract(left_paragraph._p.pPr),
        right_paragraph_contract=paragraph_properties_contract(right_paragraph._p.pPr),
    )


def reference_table_cell_text_layout(template: Path) -> TableCellTextLayout:
    """Read the approved three-line table text contract from the template."""
    reference = Document(str(template))
    three_line_style = find_style(reference, ("三线表",), WD_STYLE_TYPE.TABLE)
    sample_paragraph = None
    for table in reference.tables:
        if table.style.name.casefold() != three_line_style.name.casefold():
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    ppr = paragraph._p.pPr
                    paragraph_rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
                    if paragraph_rpr is None:
                        continue
                    size = paragraph_rpr.find(qn("w:sz"))
                    size_cs = paragraph_rpr.find(qn("w:szCs"))
                    if size is not None and size_cs is not None:
                        sample_paragraph = paragraph
                        break
                if sample_paragraph is not None:
                    break
            if sample_paragraph is not None:
                break
        if sample_paragraph is not None:
            break
    if sample_paragraph is None or sample_paragraph._p.pPr is None:
        raise ValueError("The template has no three-line table cell text contract.")

    paragraph_properties = copy.deepcopy(sample_paragraph._p.pPr)
    paragraph_rpr = paragraph_properties.find(qn("w:rPr"))
    size = paragraph_rpr.find(qn("w:sz")) if paragraph_rpr is not None else None
    size_cs = paragraph_rpr.find(qn("w:szCs")) if paragraph_rpr is not None else None
    indent = paragraph_properties.find(qn("w:ind"))
    if size is None or size_cs is None or indent is None:
        raise ValueError("The template three-line table cell lacks size or indent settings.")
    if indent.get(qn("w:firstLine"), "0") != "0" or indent.get(qn("w:firstLineChars"), "0") != "0":
        raise ValueError("The template three-line table cell must have no first-line indent.")

    normal_style = find_style(reference, ("Normal",), WD_STYLE_TYPE.PARAGRAPH)
    normal_rpr = normal_style.element.rPr
    normal_size = normal_rpr.find(qn("w:sz")) if normal_rpr is not None else None
    normal_size_cs = normal_rpr.find(qn("w:szCs")) if normal_rpr is not None else None
    if normal_size is None or normal_size_cs is None:
        raise ValueError("The template Normal style lacks explicit font sizes.")

    table_size = int(size.get(qn("w:val"), "0"))
    table_size_cs = int(size_cs.get(qn("w:val"), "0"))
    normal_size_value = int(normal_size.get(qn("w:val"), "0"))
    normal_size_cs_value = int(normal_size_cs.get(qn("w:val"), "0"))
    if not (0 < table_size < normal_size_value and 0 < table_size_cs < normal_size_cs_value):
        raise ValueError(
            "The template three-line table font must be smaller than the Normal font: "
            f"table=({table_size},{table_size_cs}), normal=({normal_size_value},{normal_size_cs_value})."
        )
    return TableCellTextLayout(
        table_style_name=three_line_style.name,
        paragraph_properties=paragraph_properties,
        normal_size=normal_size_value,
        normal_size_cs=normal_size_cs_value,
        table_size=table_size,
        table_size_cs=table_size_cs,
    )


def apply_table_cell_text_layout(paragraph, layout: TableCellTextLayout) -> None:
    """Apply the template's table text size and zero first-line indent."""
    current_properties = paragraph._p.pPr
    current_alignment = (
        copy.deepcopy(current_properties.find(qn("w:jc")))
        if current_properties is not None and current_properties.find(qn("w:jc")) is not None
        else None
    )
    replace_paragraph_properties(paragraph, layout.paragraph_properties)
    properties = paragraph._p.pPr
    if properties is None:
        raise ValueError("A table-cell paragraph has no paragraph properties after template application.")

    # Pandoc may assign right/center alignment directly to a cell. Preserve it
    # while replacing the typography and indentation contract from the template.
    template_alignment = properties.find(qn("w:jc"))
    if template_alignment is not None:
        properties.remove(template_alignment)
    if current_alignment is not None:
        run_properties = properties.find(qn("w:rPr"))
        if run_properties is None:
            properties.append(current_alignment)
        else:
            properties.insert(properties.index(run_properties), current_alignment)

    # Prevent a direct run size from overriding the paragraph-level table size.
    for run in paragraph.runs:
        run_properties = run._r.rPr
        if run_properties is None:
            continue
        for tag in (qn("w:sz"), qn("w:szCs")):
            value = run_properties.find(tag)
            if value is not None:
                run_properties.remove(value)
        if len(run_properties) == 0:
            run._r.remove(run_properties)


def paragraph_outline_level(paragraph) -> int | None:
    """Resolve the outline level from direct formatting or the inherited style."""
    ppr = paragraph._p.pPr
    if ppr is not None:
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is not None:
            return int(outline.get(qn("w:val"), "0"))

    style = paragraph.style
    while style is not None:
        style_ppr = style.element.pPr
        if style_ppr is not None:
            outline = style_ppr.find(qn("w:outlineLvl"))
            if outline is not None:
                return int(outline.get(qn("w:val"), "0"))
        style = style.base_style
    return None


def append_hidden_field(paragraph, instruction: str, placeholder: str) -> None:
    field_run = append_field(paragraph, instruction, placeholder)
    field_run.font.hidden = True


def append_simple_field(paragraph, instruction: str, placeholder: str) -> None:
    """Append a simple field whose cached text makes a headless inspection useful."""
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_table_look(
    table, *, first_row: bool | None = None, attributes: dict[str, str] | None = None
) -> None:
    table_properties = table._tbl.tblPr
    look = table_properties.find(qn("w:tblLook"))
    if look is None:
        look = OxmlElement("w:tblLook")
        table_properties.append(look)
    if attributes is not None:
        for key, value in attributes.items():
            look.set(qn(f"w:{key}"), value)
        return
    if first_row is None:
        raise ValueError("A table look requires first-row or reference attributes.")
    look.set(qn("w:val"), "04A0")
    look.set(qn("w:firstRow"), "1" if first_row else "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "1")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0")
    look.set(qn("w:noVBand"), "1")


def remove_direct_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    direct_borders = table_properties.find(qn("w:tblBorders"))
    if direct_borders is not None:
        table_properties.remove(direct_borders)
    for row in table.rows:
        for cell in row.cells:
            cell_properties = cell._tc.get_or_add_tcPr()
            direct_borders = cell_properties.find(qn("w:tcBorders"))
            if direct_borders is not None:
                cell_properties.remove(direct_borders)


def set_borderless_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    current = table_properties.find(qn("w:tblBorders"))
    if current is not None:
        table_properties.remove(current)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    table_properties.append(borders)


def set_dxa_width(properties, tag: str, width: int) -> None:
    value = properties.find(qn(tag))
    if value is None:
        value = OxmlElement(tag)
        properties.append(value)
    value.set(qn("w:w"), str(width))
    value.set(qn("w:type"), "dxa")


def apply_three_line_table_style(
    table, three_line_style, table_cell_layout: TableCellTextLayout
) -> None:
    table.style = three_line_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    remove_direct_table_borders(table)
    set_table_look(table, first_row=True)
    mark_first_row_as_header(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                apply_table_cell_text_layout(paragraph, table_cell_layout)


def configure_formula_table(
    table,
    formula_table_style,
    formula_layout: FormulaTableLayout,
    record: dict[str, Any],
) -> None:
    """Match the baseline 1x2 MathType table geometry around a display formula."""
    table.style = formula_table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_borderless_table_borders(table)
    set_table_look(table, attributes=formula_layout.table_look)
    set_dxa_width(table._tbl.tblPr, "w:tblW", formula_layout.total_width)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in (formula_layout.left_width, formula_layout.right_width):
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    left_cell, right_cell = table.rows[0].cells
    for cell, width in (
        (left_cell, formula_layout.left_width),
        (right_cell, formula_layout.right_width),
    ):
        cell_properties = cell._tc.get_or_add_tcPr()
        set_dxa_width(cell_properties, "w:tcW", width)
        vertical = cell_properties.find(qn("w:vAlign"))
        if vertical is None:
            vertical = OxmlElement("w:vAlign")
            cell_properties.append(vertical)
        vertical.set(qn("w:val"), "center")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    left_paragraph = left_cell.paragraphs[0]
    right_paragraph = right_cell.paragraphs[0]
    replace_paragraph_properties(left_paragraph, formula_layout.left_paragraph_properties)
    replace_paragraph_properties(right_paragraph, formula_layout.right_paragraph_properties)
    right_paragraph.add_run("(")
    append_field(right_paragraph, " SEQ Chapter \\c ", str(record["chapter"]))
    right_paragraph.add_run("-")
    append_field(
        right_paragraph,
        f" SEQ Equation \\r {int(record['sequence'])} \\* ARABIC ",
        str(record["sequence"]),
    )
    right_paragraph.add_run(str(record.get("suffix", "")) + ")")


def replace_display_math_with_formula_tables(
    document: Document,
    formula_records: Iterable[dict[str, Any]],
    formula_layout: FormulaTableLayout,
) -> int:
    records = list(formula_records)
    candidates: list[tuple[Any, Any]] = []
    for paragraph in list(document.paragraphs):
        display_nodes = [
            child for child in paragraph._p if child.tag == math_tag("oMathPara")
        ]
        if display_nodes:
            if len(display_nodes) != 1:
                raise ValueError("A display-equation paragraph contains multiple oMathPara nodes.")
            candidates.append((paragraph, display_nodes[0]))
    if len(candidates) != len(records):
        raise ValueError(
            "Display-formula conversion mismatch: "
            f"document has {len(candidates)} oMathPara nodes, manifest has {len(records)} records."
        )

    formula_table_style = find_style(
        document, (formula_layout.table_style_name,), WD_STYLE_TYPE.TABLE
    )
    for (paragraph, math_node), record in zip(candidates, records, strict=True):
        table = document.add_table(rows=1, cols=2)
        configure_formula_table(table, formula_table_style, formula_layout, record)
        left_paragraph = table.rows[0].cells[0].paragraphs[0]
        left_paragraph._p.append(copy.deepcopy(math_node))
        paragraph._p.addprevious(table._tbl)
        paragraph._p.getparent().remove(paragraph._p)
    return len(records)


def iter_all_paragraphs(document: Document):
    yield from document.paragraphs

    def visit_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from visit_tables(cell.tables)

    yield from visit_tables(document.tables)


def postprocess_content_docx(
    content_docx: Path,
    document_title: str,
    expected_chapter_titles: Iterable[str],
    expected_caption_records: Iterable[CaptionRecord],
    formula_records: Iterable[dict[str, Any]],
    formula_layout: FormulaTableLayout,
    table_cell_layout: TableCellTextLayout,
    title_presentation: HeadingPresentation,
    abstract_presentation: HeadingPresentation,
) -> dict[str, Any]:
    document = Document(str(content_docx))
    caption_style = find_style(document, ("图表标题", "Caption"), WD_STYLE_TYPE.PARAGRAPH)
    normal_style = find_style(document, ("Normal",), WD_STYLE_TYPE.PARAGRAPH)
    heading_one_style = find_style(document, ("heading 1", "Heading 1"), WD_STYLE_TYPE.PARAGRAPH)
    three_line_style = find_style(document, ("三线表",), WD_STYLE_TYPE.TABLE)
    chapters = list(expected_chapter_titles)
    captions = list(expected_caption_records)
    expected_figures = sum(record.kind == "图" for record in captions)
    expected_tables = sum(record.kind == "表" for record in captions)

    chapter_count = 0
    figure_count = 0
    table_count = 0
    three_line_table_count = 0
    expected_caption_labels: dict[str, list[str]] = {"figure": [], "table": []}
    caption_index = 0
    title_seen = False
    abstract_seen = False

    for paragraph in document.paragraphs:
        if "<w:drawing" in paragraph._p.xml:
            paragraph.style = caption_style

        text = paragraph.text.strip()
        outline_level = paragraph_outline_level(paragraph)
        if text == document_title:
            if title_seen:
                raise ValueError(f"Duplicate document title in output: {document_title!r}")
            apply_heading_presentation(
                paragraph, document_title, heading_one_style, title_presentation
            )
            title_seen = True
            continue
        if text == "摘要":
            if abstract_seen:
                raise ValueError("Duplicate 摘要 heading in output.")
            apply_heading_presentation(
                paragraph, "摘要", heading_one_style, abstract_presentation
            )
            abstract_seen = True
            continue
        if outline_level == 0:
            chapter_count += 1
            if chapter_count > len(chapters):
                raise ValueError(f"Unexpected first-level heading in output: {text!r}")
            chapter_title = CHAPTER_PREFIX_RE.sub("", text).strip()
            if chapter_title != chapters[chapter_count - 1]:
                raise ValueError(
                    "First-level heading mismatch: "
                    f"expected {chapters[chapter_count - 1]!r}, got {chapter_title!r}."
                )
            clear_paragraph(paragraph)
            paragraph.style = heading_one_style
            chapter_instruction = (
                " SEQ Chapter \\r 1 " if chapter_count == 1 else " SEQ Chapter "
            )
            # The heading style renders the visible Chinese chapter number.
            # These hidden fields drive captions and reset equation numbering.
            append_hidden_field(paragraph, chapter_instruction, str(chapter_count))
            append_hidden_field(paragraph, " SEQ Equation \\r 0 ", "0")
            paragraph.add_run(chapter_title)
            continue

        if outline_level in {1, 2}:
            subheading_title = SUBHEADING_PREFIX_RE.sub("", text).strip()
            if subheading_title != text:
                clear_paragraph(paragraph)
                paragraph.add_run(subheading_title)
            continue

        caption_match = CAPTION_RE.match(text)
        if not caption_match:
            continue

        if caption_index >= len(captions):
            raise ValueError(f"Unexpected caption in output: {text!r}")
        expected = captions[caption_index]
        caption_index += 1
        caption_type = caption_match.group("kind")
        caption_chapter = int(caption_match.group("chapter"))
        caption_sequence = int(caption_match.group("sequence"))
        if (
            caption_type != expected.kind
            or caption_chapter != expected.chapter
            or caption_sequence != expected.sequence
        ):
            raise ValueError(
                "Caption order or label mismatch: "
                f"expected {expected.label} at source line {expected.line_number}, got {text!r}."
            )
        if expected.chapter != chapter_count:
            raise ValueError(
                f"Caption {expected.label} follows chapter {chapter_count}, not chapter {expected.chapter}."
            )
        clear_paragraph(paragraph)
        paragraph.style = caption_style
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption_type == "图":
            figure_count += 1
            sequence = expected.sequence
            sequence_name = "Figure"
            expected_caption_labels["figure"].append(expected.label)
        else:
            table_count += 1
            sequence = expected.sequence
            sequence_name = "Table"
            expected_caption_labels["table"].append(expected.label)
        paragraph.add_run(caption_type + " ")
        append_field(paragraph, " SEQ Chapter \\c ", str(chapter_count))
        paragraph.add_run("-")
        # Source captions are authoritative. Some chapters intentionally
        # continue an existing Figure/Table label, so reset the field to the
        # explicit source sequence rather than inferring from nearby captions.
        sequence_instruction = f" SEQ {sequence_name} \\r {sequence} \\* ARABIC "
        append_field(paragraph, sequence_instruction, str(sequence))
        paragraph.add_run("　" + expected.title)

    if not title_seen or not abstract_seen:
        raise ValueError(
            "Title/abstract conversion mismatch: "
            f"title_seen={title_seen}, abstract_seen={abstract_seen}."
        )
    if chapter_count != len(chapters):
        raise ValueError(
            f"First-level heading conversion mismatch: expected {len(chapters)}, got {chapter_count}."
        )
    if caption_index != len(captions):
        raise ValueError(
            f"Caption conversion mismatch: expected {len(captions)}, got {caption_index}."
        )

    for table in document.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            apply_three_line_table_style(table, three_line_style, table_cell_layout)
            three_line_table_count += 1

    source_code_to_body = 0
    for paragraph in iter_all_paragraphs(document):
        if paragraph.style.name == "Source Code":
            paragraph.style = normal_style
            source_code_to_body += 1

    inline_math_count = convert_inline_math_to_body_text(document)
    formula_table_count = replace_display_math_with_formula_tables(
        document, formula_records, formula_layout
    )
    set_setting(document.settings.element, "w:doNotCompressPictures")
    set_setting(document.settings.element, "w:updateFields", "true")
    document.save(str(content_docx))

    if figure_count != expected_figures or table_count != expected_tables:
        raise ValueError(
            "Caption conversion mismatch: "
            f"expected figures={expected_figures}, tables={expected_tables}; "
            f"converted figures={figure_count}, tables={table_count}."
        )
    if three_line_table_count != expected_tables:
        raise ValueError(
            "Three-line table conversion mismatch: "
            f"expected tables={expected_tables}; styled tables={three_line_table_count}."
        )
    return {
        "chapters": chapter_count,
        "figures": figure_count,
        "tables": table_count,
        "three_line_tables": three_line_table_count,
        "formula_tables": formula_table_count,
        "expected_caption_labels": expected_caption_labels,
        "inline_math_to_text": inline_math_count,
        "source_code_to_body": source_code_to_body,
    }


def finalize_with_word(
    content_docx: Path,
    output: Path,
    pdf_output: Path | None,
) -> dict[str, int]:
    """Save the generated file through a hidden Word instance and update fields."""

    content_docx = content_docx.resolve()
    output = output.resolve()
    pdf_output = pdf_output.resolve() if pdf_output else None
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists():
        output.unlink()
    if pdf_output and pdf_output.exists():
        pdf_output.unlink()

    def powershell_quote(value: str) -> str:
        return "'" + value.replace("'", "''") + "'"

    source_literal = powershell_quote(str(content_docx))
    output_literal = powershell_quote(str(output))
    pdf_literal = powershell_quote(str(pdf_output)) if pdf_output else "$null"
    export_pdf = "$true" if pdf_output else "$false"
    script = f"""
$ErrorActionPreference = 'Stop'
$sourcePath = {source_literal}
$outputPath = {output_literal}
$pdfPath = {pdf_literal}
$exportPdf = {export_pdf}
$word = $null
$document = $null
$restorePictureSetting = $false
$originalPictureSetting = $false
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $word.ScreenUpdating = $false
    try {{
        $originalPictureSetting = $word.Options.DoNotCompressPicturesInFile
        $word.Options.DoNotCompressPicturesInFile = $true
        $restorePictureSetting = $true
    }} catch {{}}
    $document = $word.Documents.Open($sourcePath, $false, $true, $false)
    $document.SaveAs2($outputPath, 16)
    $document.Fields.Update() | Out-Null
    foreach ($section in $document.Sections) {{
        foreach ($collection in @($section.Headers, $section.Footers)) {{
            foreach ($index in 1, 2, 3) {{
                try {{ $collection.Item($index).Range.Fields.Update() | Out-Null }} catch {{}}
            }}
        }}
    }}
    $document.Repaginate()
    $pages = [int]$document.ComputeStatistics(2)
    $fieldCount = [int]$document.Fields.Count
    $document.Save()
    if ($exportPdf) {{ $document.ExportAsFixedFormat($pdfPath, 17, $false) }}
    [pscustomobject]@{{word_pages=$pages; word_body_fields=$fieldCount}} | ConvertTo-Json -Compress
}} finally {{
    if ($null -ne $document) {{ try {{ $document.Close($false) }} catch {{}} }}
    if ($null -ne $word) {{
        if ($restorePictureSetting) {{ try {{ $word.Options.DoNotCompressPicturesInFile = $originalPictureSetting }} catch {{}} }}
        try {{ $word.Quit() }} catch {{}}
    }}
}}
"""
    encoded = base64.b64encode(script.encode("utf-16le")).decode("ascii")
    # Use host PowerShell COM because the bundled Python runtime does not ship pywin32.
    completed = subprocess.run(
        ["powershell.exe", "-NoProfile", "-NonInteractive", "-EncodedCommand", encoded],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if completed.returncode:
        raise RuntimeError(
            "Word finalization failed: " + (completed.stderr.strip() or completed.stdout.strip())
        )
    try:
        result = json.loads(completed.stdout.strip().splitlines()[-1])
        return {
            "word_pages": int(result["word_pages"]),
            "word_body_fields": int(result["word_body_fields"]),
        }
    except (IndexError, KeyError, TypeError, ValueError, json.JSONDecodeError) as error:
        raise RuntimeError(
            f"Word finalization returned no usable result: {completed.stdout!r}"
        ) from error


def force_document_settings(output: Path) -> None:
    """Keep no-compression and update-fields settings after the Word save."""
    settings_path = "word/settings.xml"
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(output, "r") as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    root = etree.fromstring(files[settings_path])
    do_not_compress = root.find(f"{{{namespace}}}doNotCompressPictures")
    if do_not_compress is None:
        do_not_compress = etree.Element(f"{{{namespace}}}doNotCompressPictures")
        root.append(do_not_compress)
    update_fields = root.find(f"{{{namespace}}}updateFields")
    if update_fields is None:
        update_fields = etree.Element(f"{{{namespace}}}updateFields")
        root.append(update_fields)
    update_fields.set(f"{{{namespace}}}val", "true")
    files[settings_path] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    temporary = output.with_suffix(output.suffix + ".settings-tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    temporary.replace(output)


def force_document_core_properties(output: Path, title: str, author: str) -> None:
    """Pin metadata after Word COM saves so Office account identity cannot leak in."""
    core_path = "docProps/core.xml"
    with zipfile.ZipFile(output, "r") as archive:
        files = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    root = etree.fromstring(files[core_path])

    def set_property(namespace: str, local_name: str, value: str) -> None:
        node = root.find(f"{{{namespace}}}{local_name}")
        if node is None:
            node = etree.Element(f"{{{namespace}}}{local_name}")
            root.append(node)
        node.text = value

    set_property(DC_NS, "title", title)
    set_property(DC_NS, "creator", author)
    set_property(CORE_PROPERTIES_NS, "lastModifiedBy", author)
    files[core_path] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    temporary = output.with_suffix(output.suffix + ".metadata-tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    temporary.replace(output)


def inspect_docx(output: Path) -> dict[str, Any]:
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    office_namespace = "urn:schemas-microsoft-com:office:office"
    namespaces = {"w": word_namespace, "m": MATH_NS, "o": office_namespace}

    with zipfile.ZipFile(output, "r") as archive:
        document_root = etree.fromstring(archive.read("word/document.xml"))
        styles_root = etree.fromstring(archive.read("word/styles.xml"))
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        core_xml = archive.read("docProps/core.xml")
        media_files = [info for info in archive.infolist() if info.filename.startswith("word/media/")]

    style_names = {
        style.get(qn("w:styleId")): style.find(qn("w:name")).get(qn("w:val"), "")
        for style in styles_root.findall(qn("w:style"))
        if style.find(qn("w:name")) is not None
    }

    def normalized_fields(element) -> list[str]:
        return [
            " ".join(value.split())
            for value in element.xpath(
                ".//w:instrText/text() | .//w:fldSimple/@w:instr", namespaces=namespaces
            )
        ]

    def visible_text_without_hidden_runs(paragraph) -> str:
        text: list[str] = []
        for run in paragraph.findall(qn("w:r")):
            properties = run.find(qn("w:rPr"))
            if properties is not None and properties.find(qn("w:vanish")) is not None:
                continue
            text.extend(node.text or "" for node in run.findall(qn("w:t")))
        return "".join(text).strip()

    def direct_text_layout(paragraph) -> dict[str, Any]:
        properties = paragraph.find("w:pPr", namespaces)
        indent = properties.find("w:ind", namespaces) if properties is not None else None
        run_properties = properties.find("w:rPr", namespaces) if properties is not None else None
        size = run_properties.find("w:sz", namespaces) if run_properties is not None else None
        size_cs = run_properties.find("w:szCs", namespaces) if run_properties is not None else None
        return {
            "first_line": indent.get(qn("w:firstLine")) if indent is not None else None,
            "first_line_chars": indent.get(qn("w:firstLineChars")) if indent is not None else None,
            "size": size.get(qn("w:val")) if size is not None else None,
            "size_cs": size_cs.get(qn("w:val")) if size_cs is not None else None,
        }

    instructions = normalized_fields(document_root)
    formula_table_contracts: list[dict[str, Any]] = []
    three_line_table_contracts: list[dict[str, Any]] = []
    tables = document_root.xpath(".//w:body//w:tbl", namespaces=namespaces)
    for table in tables:
        style = table.find("w:tblPr/w:tblStyle", namespaces)
        style_id = style.get(qn("w:val"), "") if style is not None else ""
        style_name = style_names.get(style_id, style_id)
        rows = table.xpath("./w:tr", namespaces=namespaces)
        if style_name == "三线表":
            look = table.find("w:tblPr/w:tblLook", namespaces)
            three_line_table_contracts.append(
                {
                    "style": style_name,
                    "first_row_is_header": bool(
                        rows
                        and rows[0].xpath("./w:trPr/w:tblHeader", namespaces=namespaces)
                    ),
                    "direct_borders": {
                        etree.QName(border.tag).localname: border.get(qn("w:val"), "")
                        for border in table.xpath(
                            "./w:tblPr/w:tblBorders/*", namespaces=namespaces
                        )
                    },
                    "table_look": {
                        etree.QName(key).localname: value
                        for key, value in look.attrib.items()
                    }
                    if look is not None
                    else {},
                    "cell_paragraph_layouts": [
                        direct_text_layout(paragraph)
                        for paragraph in table.xpath(".//w:tc/w:p", namespaces=namespaces)
                    ],
                }
            )
        has_formula = bool(
            table.xpath(
                ".//m:oMath | .//o:OLEObject[@ProgID='Equation.DSMT4']",
                namespaces=namespaces,
            )
        )
        if not has_formula:
            continue
        cells = rows[0].xpath("./w:tc", namespaces=namespaces) if rows else []
        grid_widths = [
            int(column.get(qn("w:w"), "0"))
            for column in table.xpath("./w:tblGrid/w:gridCol", namespaces=namespaces)
        ]
        borders = {
            etree.QName(border.tag).localname: border.get(qn("w:val"), "")
            for border in table.xpath("./w:tblPr/w:tblBorders/*", namespaces=namespaces)
        }
        formula_table_contracts.append(
            {
                "rows": len(rows),
                "columns": len(cells),
                "style": style_name,
                "grid_widths": grid_widths,
                "borders": borders,
                "left_paragraph_contract": paragraph_properties_contract(
                    cells[0].find("w:p/w:pPr", namespaces) if len(cells) == 2 else None
                ),
                "right_paragraph_contract": paragraph_properties_contract(
                    cells[1].find("w:p/w:pPr", namespaces) if len(cells) == 2 else None
                ),
                "field_instructions": normalized_fields(table),
                "right_cell_text": "".join(
                    cells[1].xpath(".//w:t/text()", namespaces=namespaces)
                )
                if len(cells) == 2
                else "",
                "omath_objects": len(table.xpath(".//m:oMath", namespaces=namespaces)),
                "mathtype_objects": len(
                    table.xpath(
                        ".//o:OLEObject[@ProgID='Equation.DSMT4']", namespaces=namespaces
                    )
                ),
            }
        )

    heading_style_ids = {
        style_id for style_id, name in style_names.items() if name.casefold() == "heading 1"
    }
    heading_contracts: list[dict[str, Any]] = []
    for paragraph in document_root.xpath(".//w:body/w:p", namespaces=namespaces):
        style = paragraph.find("w:pPr/w:pStyle", namespaces)
        style_id = style.get(qn("w:val"), "") if style is not None else ""
        if style_id not in heading_style_ids:
            continue
        heading_contracts.append(
            {
                "style": style_names.get(style_id, style_id),
                "visible_text": visible_text_without_hidden_runs(paragraph),
                "field_instructions": normalized_fields(paragraph),
            }
        )

    caption_field_contracts: list[dict[str, Any]] = []
    for paragraph in document_root.xpath(".//w:body/w:p", namespaces=namespaces):
        fields = normalized_fields(paragraph)
        kind = (
            "图"
            if any(field.startswith("SEQ Figure") for field in fields)
            else "表"
            if any(field.startswith("SEQ Table") for field in fields)
            else None
        )
        if kind is None:
            continue
        visible_text = "".join(paragraph.xpath(".//w:t/text()", namespaces=namespaces)).strip()
        match = CAPTION_RE.match(visible_text)
        caption_field_contracts.append(
            {
                "kind": kind,
                "label": (
                    f"{match.group('kind')} {match.group('chapter')}-{match.group('sequence')}"
                    if match
                    else None
                ),
                "visible_text": visible_text,
                "field_instructions": fields,
            }
        )

    core_root = etree.fromstring(core_xml)
    creator = core_root.find(f"{{{DC_NS}}}creator")
    last_modified_by = core_root.find(f"{{{CORE_PROPERTIES_NS}}}lastModifiedBy")
    exported_document = Document(str(output))
    source_code_paragraph_count = sum(
        paragraph.style.name == "Source Code" for paragraph in iter_all_paragraphs(exported_document)
    )
    return {
        "drawing_count": len(document_root.xpath(".//w:drawing", namespaces=namespaces)),
        "native_table_count": len(tables),
        "office_math_count": len(document_root.xpath(".//m:oMath", namespaces=namespaces)),
        "chapter_field_count": sum(field.startswith("SEQ Chapter") for field in instructions),
        "figure_field_count": sum(field.startswith("SEQ Figure") for field in instructions),
        "table_field_count": sum(field.startswith("SEQ Table") for field in instructions),
        "media_file_count": len(media_files),
        "media_bytes": sum(info.file_size for info in media_files),
        "no_picture_compression": "doNotCompressPictures" in settings_xml,
        "update_fields_on_open": "updateFields" in settings_xml,
        "creator": creator.text if creator is not None and creator.text else "",
        "last_modified_by": last_modified_by.text if last_modified_by is not None and last_modified_by.text else "",
        "source_code_paragraph_count": source_code_paragraph_count,
        "three_line_table_count": len(three_line_table_contracts),
        "formula_table_contracts": formula_table_contracts,
        "three_line_table_contracts": three_line_table_contracts,
        "heading_contracts": heading_contracts,
        "caption_field_contracts": caption_field_contracts,
    }


def validate_caption_field_contracts(
    contracts: Iterable[dict[str, Any]], records: Iterable[CaptionRecord]
) -> list[str]:
    errors: list[str] = []
    actual = list(contracts)
    expected = list(records)
    if len(actual) != len(expected):
        return [f"caption field contract count {len(actual)} != expected {len(expected)}"]

    for index, (contract, record) in enumerate(zip(actual, expected, strict=True), start=1):
        expected_sequence = f"SEQ {record.field_name} \\r {record.sequence} \\* ARABIC"
        expected_fields = ["SEQ Chapter \\c", expected_sequence]
        if contract["kind"] != record.kind or contract["label"] != record.label:
            errors.append(
                f"caption {index} label is {contract['label']!r}, expected {record.label!r}"
            )
        if contract["field_instructions"] != expected_fields:
            errors.append(
                f"caption {record.label} fields are {contract['field_instructions']!r}, "
                f"expected {expected_fields!r}"
            )
    return errors


def validate_formula_table_contracts(
    contracts: Iterable[dict[str, Any]],
    records: Iterable[dict[str, Any]],
    formula_layout: FormulaTableLayout,
) -> list[str]:
    errors: list[str] = []
    actual = list(contracts)
    expected = list(records)
    if len(actual) != len(expected):
        return [f"formula table count {len(actual)} != expected {len(expected)}"]

    expected_borders = {
        "top": "none",
        "left": "none",
        "bottom": "none",
        "right": "none",
        "insideH": "none",
        "insideV": "none",
    }
    expected_grid = [formula_layout.left_width, formula_layout.right_width]
    for index, (contract, record) in enumerate(zip(actual, expected, strict=True), start=1):
        expected_fields = [
            "SEQ Chapter \\c",
            f"SEQ Equation \\r {int(record['sequence'])} \\* ARABIC",
        ]
        if contract["rows"] != 1 or contract["columns"] != 2:
            errors.append(f"formula {index} is not a one-row two-column table")
        if contract["style"] != formula_layout.table_style_name:
            errors.append(
                f"formula {index} style is {contract['style']!r}, "
                f"expected {formula_layout.table_style_name!r}"
            )
        if contract["grid_widths"] != expected_grid:
            errors.append(
                f"formula {index} grid is {contract['grid_widths']!r}, expected {expected_grid!r}"
            )
        if contract["borders"] != expected_borders:
            errors.append(f"formula {index} borders are not explicitly borderless")
        if contract["left_paragraph_contract"] != formula_layout.left_paragraph_contract:
            errors.append(f"formula {index} body paragraph style/indent differs from the template")
        if contract["right_paragraph_contract"] != formula_layout.right_paragraph_contract:
            errors.append(f"formula {index} number paragraph style/indent differs from the template")
        if contract["field_instructions"] != expected_fields:
            errors.append(
                f"formula {index} fields are {contract['field_instructions']!r}, "
                f"expected {expected_fields!r}"
            )
        expected_label = f"({record['expected_number']})"
        if "".join(str(contract["right_cell_text"]).split()) != expected_label:
            errors.append(
                f"formula {index} visible number is {contract['right_cell_text']!r}, "
                f"expected {expected_label!r}"
            )
        if contract["omath_objects"] != 1 or contract["mathtype_objects"] != 0:
            errors.append(
                f"formula {index} expected one Office Math object and no MathType OLE object"
            )
    return errors


def validate_three_line_table_contracts(
    contracts: Iterable[dict[str, Any]],
    expected_count: int,
    table_cell_layout: TableCellTextLayout,
) -> list[str]:
    errors: list[str] = []
    actual = list(contracts)
    if len(actual) != expected_count:
        return [f"three-line table count {len(actual)} != expected {expected_count}"]

    for table_index, contract in enumerate(actual, start=1):
        if contract["style"] != table_cell_layout.table_style_name:
            errors.append(
                f"table {table_index} style is {contract['style']!r}, "
                f"expected {table_cell_layout.table_style_name!r}"
            )
        for paragraph_index, layout in enumerate(
            contract["cell_paragraph_layouts"], start=1
        ):
            if layout["first_line"] != "0" or layout["first_line_chars"] != "0":
                errors.append(
                    f"table {table_index} cell paragraph {paragraph_index} has a first-line indent"
                )
            if layout["size"] != str(table_cell_layout.table_size):
                errors.append(
                    f"table {table_index} cell paragraph {paragraph_index} size is "
                    f"{layout['size']!r}, expected {table_cell_layout.table_size!r}"
                )
            if layout["size_cs"] != str(table_cell_layout.table_size_cs):
                errors.append(
                    f"table {table_index} cell paragraph {paragraph_index} complex-script size is "
                    f"{layout['size_cs']!r}, expected {table_cell_layout.table_size_cs!r}"
                )
    return errors


def verify_embedded_images(
    source_image_files: Iterable[Path], output: Path
) -> dict[str, int]:
    source_hashes = {
        hashlib.sha256(image.read_bytes()).hexdigest() for image in source_image_files
    }
    with zipfile.ZipFile(output, "r") as archive:
        embedded_hashes = {
            hashlib.sha256(archive.read(info.filename)).hexdigest()
            for info in archive.infolist()
            if info.filename.startswith("word/media/")
        }
    return {
        "source_hash_count": len(source_hashes),
        "embedded_hash_count": len(embedded_hashes),
        "missing_source_hash_count": len(source_hashes - embedded_hashes),
        "extra_embedded_hash_count": len(embedded_hashes - source_hashes),
    }


def build_manifest(
    work_dir: Path,
    source: Path,
    template: Path,
    layout_reference: Path,
    output: Path,
    title: str,
    author: str,
    source_image_paths: Iterable[str],
    source_caption_counts: Counter[str],
    source_latex_count: int,
    source_chapter_count: int,
    formula_manifest_path: Path,
    formula_layout: FormulaTableLayout,
    table_cell_layout: TableCellTextLayout,
    converted_counts: dict[str, int],
    word_counts: dict[str, int],
    package_counts: dict[str, Any],
    image_embedding: dict[str, int],
    pdf_output: Path | None,
) -> Path:
    image_list = list(source_image_paths)
    manifest = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "title": title,
        "author": author,
        "source": str(source),
        "template": str(template),
        "template_sha256": hashlib.sha256(template.read_bytes()).hexdigest(),
        "layout_reference": str(layout_reference),
        "layout_reference_sha256": hashlib.sha256(layout_reference.read_bytes()).hexdigest(),
        "output": str(output),
        "output_bytes": output.stat().st_size,
        "pdf_output": str(pdf_output) if pdf_output else None,
        "pdf_bytes": pdf_output.stat().st_size if pdf_output and pdf_output.exists() else None,
        "source_image_references": len(image_list),
        "source_unique_image_references": len(set(image_list)),
        "source_latex_blocks": source_latex_count,
        "source_chapter_headings": source_chapter_count,
        "source_figure_captions": source_caption_counts["figure"],
        "source_table_captions": source_caption_counts["table"],
        "formula_manifest": str(formula_manifest_path),
        "formula_layout": {
            "table_style": formula_layout.table_style_name,
            "total_width": formula_layout.total_width,
            "left_width": formula_layout.left_width,
            "right_width": formula_layout.right_width,
            "table_look": formula_layout.table_look,
        },
        "table_cell_layout": {
            "table_style": table_cell_layout.table_style_name,
            "normal_size": table_cell_layout.normal_size,
            "normal_size_cs": table_cell_layout.normal_size_cs,
            "table_size": table_cell_layout.table_size,
            "table_size_cs": table_cell_layout.table_size_cs,
        },
        "converted": converted_counts,
        "word": word_counts,
        "package": package_counts,
        "image_embedding": image_embedding,
    }
    manifest_path = work_dir / "build_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest_path


def main() -> int:
    args = parse_args()
    source = args.source.resolve()
    template = args.template.resolve()
    layout_reference = args.layout_reference.resolve()
    output = args.output.resolve()
    work_dir = args.work_dir.resolve()
    pdf_output = work_dir / (output.stem + ".pdf") if args.export_pdf else None
    normalized_markdown = work_dir / "normalized_report.md"
    content_docx = work_dir / "pandoc_content.docx"

    if not source.is_file():
        raise FileNotFoundError(source)
    if not template.is_file():
        raise FileNotFoundError(template)
    if not layout_reference.is_file():
        raise FileNotFoundError(layout_reference)
    work_dir.mkdir(parents=True, exist_ok=True)

    title, source_markdown = read_source(source)
    author = (args.author or template_author(template) or template_author(layout_reference)).strip()
    if not author:
        raise ValueError(
            "No output author was supplied and neither the clean template nor the layout reference has creator metadata."
        )
    source_images = image_paths(source_markdown)
    source_image_files = resolve_report_images(source, source_images)
    source_caption_records = caption_records(source_markdown)
    source_captions = caption_counts(source_markdown)
    source_chapters = source_primary_chapter_titles(source_markdown)
    formula_manifest_path, formula_records = prepare_formula_manifest(source, work_dir)
    formula_layout = reference_formula_table_layout(layout_reference)
    table_cell_layout = reference_table_cell_text_layout(layout_reference)
    title_presentation, abstract_presentation = reference_heading_presentations(layout_reference)

    prepared_markdown = make_pandoc_markdown(source_markdown)
    normalized_markdown.write_text(prepared_markdown, encoding="utf-8")
    emit(
        "Validated "
        f"{len(source_images)} image references, {source_captions['figure']} figure captions, "
        f"{source_captions['table']} table captions, {len(source_chapters)} chapter headings, and "
        f"{len(formula_records)} tagged display formulas."
    )

    run_pandoc(
        find_pandoc(),
        normalized_markdown,
        template,
        (source.parent, REPORT_DIR),
        content_docx,
        title,
    )
    converted_counts = postprocess_content_docx(
        content_docx,
        title,
        source_chapters,
        source_caption_records,
        formula_records,
        formula_layout,
        table_cell_layout,
        title_presentation,
        abstract_presentation,
    )
    word_counts = finalize_with_word(content_docx, output, pdf_output)
    force_document_settings(output)
    force_document_core_properties(output, title, author)
    package_counts = inspect_docx(output)
    image_embedding = verify_embedded_images(source_image_files, output)

    errors: list[str] = []
    if package_counts["drawing_count"] != len(source_images):
        errors.append(
            f"drawing count {package_counts['drawing_count']} != source images {len(source_images)}"
        )
    if package_counts["native_table_count"] < source_captions["table"]:
        errors.append(
            f"native table count {package_counts['native_table_count']} < captions {source_captions['table']}"
        )
    if package_counts["office_math_count"] != len(formula_records):
        errors.append(
            f"Office Math count {package_counts['office_math_count']} != formula manifest {len(formula_records)}"
        )
    if package_counts["figure_field_count"] != source_captions["figure"]:
        errors.append("Figure field count does not match figure captions")
    if package_counts["table_field_count"] != source_captions["table"]:
        errors.append("Table field count does not match table captions")
    expected_chapter_fields = (
        converted_counts["chapters"]
        + source_captions["figure"]
        + source_captions["table"]
        + len(formula_records)
    )
    if package_counts["chapter_field_count"] != expected_chapter_fields:
        errors.append("Chapter field count does not match headings, captions, and formulas")
    if package_counts["three_line_table_count"] != source_captions["table"]:
        errors.append("Three-line table count does not match source table captions")
    errors.extend(
        validate_caption_field_contracts(
            package_counts["caption_field_contracts"], source_caption_records
        )
    )
    errors.extend(
        validate_formula_table_contracts(
            package_counts["formula_table_contracts"], formula_records, formula_layout
        )
    )
    errors.extend(
        validate_three_line_table_contracts(
            package_counts["three_line_table_contracts"],
            source_captions["table"],
            table_cell_layout,
        )
    )
    if not package_counts["no_picture_compression"]:
        errors.append("Output does not contain the no-picture-compression setting")
    if not package_counts["update_fields_on_open"]:
        errors.append("Output does not request field updates on open")
    if package_counts["creator"] != author or package_counts["last_modified_by"] != author:
        errors.append("Output metadata author does not match the requested author")
    if package_counts["source_code_paragraph_count"]:
        errors.append("Output still contains Source Code paragraphs instead of the body style")
    if image_embedding["missing_source_hash_count"]:
        errors.append("One or more source images were not embedded byte-for-byte")

    manifest_path = build_manifest(
        work_dir,
        source,
        template,
        layout_reference,
        output,
        title,
        author,
        source_images,
        source_captions,
        len(LATEX_FENCE_RE.findall(source_markdown)),
        len(source_chapters),
        formula_manifest_path,
        formula_layout,
        table_cell_layout,
        converted_counts,
        word_counts,
        package_counts,
        image_embedding,
        pdf_output,
    )
    emit(f"Wrote {output} ({output.stat().st_size} bytes).")
    emit(f"Wrote {manifest_path}.")
    if errors:
        for error in errors:
            emit("CHECK FAILED: " + error)
        return 2
    emit("Structural checks passed.")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:
        emit(f"FAILED: {error}")
        raise
