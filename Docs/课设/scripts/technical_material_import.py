"""Copy bounded technical-report sections into independent course reports."""

from __future__ import annotations

import copy
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


@dataclass(frozen=True)
class ImportStats:
    paragraphs: int
    tables: int
    image_references: int
    unique_source_images: int


CAPTION_PATTERN = re.compile(
    r"^(?P<label>[图表])\s*-?\s*(?P<number>\d+(?:[-.]\d+)*(?:[A-Za-z])?)\s*(?P<title>.+?)\s*$"
)
CAPTION_EXPLANATION_PATTERN = re.compile(
    r"^(?P<title>.+?)[。；]\s*(?P<explanation>(?:图中|该图|表中|该表|该界面|该流程|其中).+)$"
)


def _set_run_font(run, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def _add_seq_field(paragraph, identifier: str, result: int) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" SEQ {identifier} \\* ARABIC "
    begin_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(str(result))
    _set_run_font(result_run, "宋体", 10.5)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _ensure_caption_style(document: Document) -> None:
    if "图表标题" in document.styles:
        style = document.styles["图表标题"]
    else:
        style = document.styles.add_style("图表标题", WD_STYLE_TYPE.PARAGRAPH)
    normal = document.styles["Normal"]
    style.base_style = normal
    style.next_paragraph_style = normal
    style.font.bold = True
    style.font.size = Pt(10.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = None
    style.paragraph_format.space_after = None
    style.paragraph_format.line_spacing = None

    # Match the source report's quick-style flag so Word exposes it consistently.
    if style.element.find(qn("w:qFormat")) is None:
        style.element.append(OxmlElement("w:qFormat"))


def _set_border(border, value: str, size: int | None = None) -> None:
    border.set(qn("w:val"), value)
    if size is not None:
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
    else:
        for attribute in (qn("w:sz"), qn("w:space"), qn("w:color")):
            border.attrib.pop(attribute, None)


def _set_three_line_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge_name in ("top", "bottom", "left", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        _set_border(edge, "single", 12) if edge_name in {"top", "bottom"} else _set_border(edge, "nil")

    for cell in table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        cell_borders = cell_properties.find(qn("w:tcBorders"))
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_properties.append(cell_borders)
        header_border = cell_borders.find(qn("w:bottom"))
        if header_border is None:
            header_border = OxmlElement("w:bottom")
            cell_borders.append(header_border)
        _set_border(header_border, "single", 6)


def _format_body_table(table, three_line_style) -> None:
    is_data_table = len(table.rows) > 1
    if is_data_table:
        # Keep the user-visible style name consistent with the supplied report.
        _set_table_style_id(table, three_line_style.style_id)
        _set_three_line_borders(table)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_properties = table.rows[0]._tr.get_or_add_trPr()
        table_header = row_properties.find(qn("w:tblHeader"))
        if table_header is None:
            table_header = OxmlElement("w:tblHeader")
            row_properties.append(table_header)
        table_header.set(qn("w:val"), "true")

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            shading = cell_properties.find(qn("w:shd"))
            if shading is not None:
                cell_properties.remove(shading)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.right_indent = Cm(0)
                if is_data_table:
                    # The source three-line-table style centers table content.
                    paragraph.alignment = None
                for run in paragraph.runs:
                    _set_run_font(run, "宋体", 12)
                    if is_data_table and row_index == 0:
                        run.bold = True


def _display_width(text: str) -> float:
    width = 0.0
    for character in text.replace("\n", ""):
        width += 0.55 if character.isascii() else 1.0
    return width


def _set_cell_margins(cell, horizontal: int = 72, vertical: int = 36) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for edge_name, amount in (("top", vertical), ("left", horizontal), ("bottom", vertical), ("right", horizontal)):
        edge = margins.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            margins.append(edge)
        edge.set(qn("w:w"), str(amount))
        edge.set(qn("w:type"), "dxa")


def _optimise_generated_table_geometry(table) -> None:
    """Improve readability of markdown-generated data tables without touching source tables."""
    if len(table.rows) <= 1 or any(
        len({id(cell._tc) for cell in row.cells}) != len(row.cells) for row in table.rows
    ):
        return

    columns = len(table.rows[0].cells)
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    scores = [max(_display_width(row.cells[index].text) for row in table.rows) for index in range(columns)]
    widths = []
    for header, score in zip(headers, scores):
        if any(token in header for token in ("编号", "序号", "ID", "类别", "状态")):
            widths.append(2.0)
        else:
            widths.append(max(3.2, min(7.0, 1.4 + score * 0.19)))
    scale = 16.0 / sum(widths)
    widths = [width * scale for width in widths]

    grid = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        grid[index].set(qn("w:w"), str(round(width * 567)))
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = Cm(widths[column_index])
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                if row_index == 0 or scores[column_index] < 14:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _append_explanation_after(document: Document, caption, text: str) -> None:
    explanation = document.add_paragraph()
    explanation.style = document.styles["Normal"]
    explanation.paragraph_format.first_line_indent = Cm(0.74)
    run = explanation.add_run(text)
    _set_run_font(run, "宋体", 10.5)
    caption._p.addnext(explanation._p)


def _table_caption_title(table, paragraph_by_element: dict) -> str:
    element = table._tbl.getprevious()
    while element is not None:
        paragraph = paragraph_by_element.get(element)
        if paragraph is not None and paragraph.style.name.startswith(("MoSim Heading", "Heading")):
            return re.sub(r"^(?:\d+(?:\.\d+)*\.?|[（(]\d+[）)])\s*", "", paragraph.text).strip()
        element = element.getprevious()

    headers = [cell.text.strip().replace("\n", "、") for cell in table.rows[0].cells]
    return "、".join(header for header in headers if header)[:40] or "数据项"


def _has_table_caption(table, paragraph_by_element: dict) -> bool:
    previous = table._tbl.getprevious()
    paragraph = paragraph_by_element.get(previous)
    if paragraph is None:
        return False
    match = CAPTION_PATTERN.match(paragraph.text.strip())
    return match is not None and match.group("label") == "表"


def _insert_missing_table_captions(document: Document, tables) -> list:
    paragraph_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    inserted = []
    for table in tables:
        if len(table.rows) <= 1 or _has_table_caption(table, paragraph_by_element):
            continue
        caption = document.add_paragraph()
        caption.style = document.styles["Normal"]
        caption.add_run(f"表0 {_table_caption_title(table, paragraph_by_element)}")
        table._tbl.addprevious(caption._p)
        paragraph_by_element[caption._p] = caption
        inserted.append(table)
    return inserted


def apply_report_layout(document: Document, body_table_start: int = 3) -> dict[str, int]:
    """Apply the course-report table and caption contract after all content is imported."""
    _ensure_caption_style(document)
    three_line_style = document.styles["三线表"]

    tables = document.tables[body_table_start:]
    for table in tables:
        _format_body_table(table, three_line_style)
    inserted_tables = _insert_missing_table_captions(document, tables)
    for table in inserted_tables:
        _optimise_generated_table_geometry(table)

    body_paragraphs = list(document.paragraphs)
    for table in tables:
        body_paragraphs.extend(
            paragraph for row in table.rows for cell in row.cells for paragraph in cell.paragraphs
        )
    for paragraph in body_paragraphs:
        for run in paragraph.runs:
            if "我们选择" in run.text or "我们构建" in run.text:
                run.text = run.text.replace("我们选择", "技术选型采用").replace("我们构建", "本项目构建")

    figure_count = 0
    table_count = 0
    caption_count = 0
    explanations = 0
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        match = CAPTION_PATTERN.match(text)
        if match is None:
            if list(paragraph._p.iter(qn("w:drawing"))):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
            continue

        label = match.group("label")
        title = match.group("title").strip()
        explanation = None
        split = CAPTION_EXPLANATION_PATTERN.match(title)
        if split is not None:
            title = split.group("title").strip()
            explanation = split.group("explanation").strip()

        if label == "图":
            figure_count += 1
            number = figure_count
        else:
            table_count += 1
            number = table_count

        _clear_paragraph(paragraph)
        paragraph.style = document.styles["图表标题"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        paragraph.paragraph_format.line_spacing = None
        paragraph.paragraph_format.keep_with_next = label == "表"
        prefix = paragraph.add_run(label)
        _set_run_font(prefix, "宋体", 10.5)
        _add_seq_field(paragraph, label, number)
        title_run = paragraph.add_run(f" {title}")
        _set_run_font(title_run, "宋体", 10.5)
        caption_count += 1

        if explanation:
            _append_explanation_after(document, paragraph, explanation)
            explanations += 1

    return {
        "tables": len(tables),
        "captions": caption_count,
        "figures": figure_count,
        "table_captions": table_count,
        "inserted_table_captions": len(inserted_tables),
        "caption_explanations": explanations,
    }


def validate_report_layout(output: Path, body_table_start: int = 3) -> dict[str, int]:
    """Check the body table and caption rules after the DOCX is saved."""
    document = Document(output)
    data_tables = document.tables[body_table_start:]
    bad_style = []
    bad_indent = []
    bad_font_size = []
    for table_index, table in enumerate(data_tables, start=body_table_start):
        if len(table.rows) > 1 and table.style.name != "三线表":
            bad_style.append((table_index, table.style.name))
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    indent = paragraph.paragraph_format.first_line_indent
                    if indent is not None and abs(indent.cm) > 0.001:
                        bad_indent.append((table_index, row_index, cell_index, paragraph_index))
                    for run in paragraph.runs:
                        if run.font.size is not None and abs(run.font.size.pt - 12) > 0.01:
                            bad_font_size.append((table_index, row_index, cell_index, paragraph_index))

    caption_style = document.styles["图表标题"]

    # Word may collapse a 10.5-point direct font setting into the Normal-style
    # inheritance on save. Zero twips are likewise exposed as None by
    # python-docx, so inspect the effective size and the stored XML zero value.
    effective_caption_size = caption_style.font.size
    size_style = caption_style.base_style
    while effective_caption_size is None and size_style is not None:
        effective_caption_size = size_style.font.size
        size_style = size_style.base_style
    caption_indentation = caption_style.element.pPr
    zero_first_line = False
    if caption_indentation is not None:
        indent = caption_indentation.find(qn("w:ind"))
        zero_first_line = indent is not None and (
            indent.get(qn("w:firstLine")) == "0"
            or indent.get(qn("w:firstLineChars")) == "0"
        )
    caption_style_bad = (
        not caption_style.font.bold
        or effective_caption_size is None
        or abs(effective_caption_size.pt - 10.5) > 0.01
        or caption_style.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER
        or not zero_first_line
        or caption_style.next_paragraph_style.name != "Normal"
    )

    bad_captions = []
    caption_count = 0
    for paragraph in document.paragraphs:
        instruction = "".join(
            instruction_text.text or "" for instruction_text in paragraph._p.iter(qn("w:instrText"))
        )
        if "SEQ" not in instruction:
            continue
        caption_count += 1
        text = paragraph.text
        if (
            paragraph.style.name != "图表标题"
            or not text.startswith(("图", "表"))
            or len(text) < 2
            or text[1].isspace()
        ):
            bad_captions.append(text)

    if bad_style or bad_indent or bad_font_size or caption_style_bad or bad_captions:
        raise RuntimeError(
            "Course-report layout validation failed: "
            f"styles={bad_style[:3]}, indents={bad_indent[:3]}, "
            f"font_sizes={bad_font_size[:3]}, caption_style={caption_style_bad}, "
            f"captions={bad_captions[:3]}"
        )
    return {"body_tables": len(data_tables), "captions": caption_count}


def retain_front_matter(document: Document, paragraph_ranges: tuple[tuple[int, int], ...]) -> None:
    """Keep one cover, task sheet, response sheet, and the template TOC."""
    body = document._element.body
    children = list(body)
    positions = {child: index for index, child in enumerate(children)}
    keep = {child for child in children if child.tag in {qn("w:sdt"), qn("w:sectPr")}}
    for start, end in paragraph_ranges:
        start_element = document.paragraphs[start]._element
        end_element = document.paragraphs[end]._element
        for child in children[positions[start_element]:positions[end_element]]:
            keep.add(child)
    for child in children:
        if child not in keep:
            body.remove(child)


def _normalise_heading(text: str) -> str:
    return text.rstrip("0123456789 ").strip()


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"MoSim Heading {level}"]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 16, 2: 14, 3: 12}[level])
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "黑体")


def _set_normal_style(element) -> None:
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_style.set(qn("w:val"), "Normal")


def _centre_caption(element) -> None:
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        return
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        from docx.oxml import OxmlElement

        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "center")


def _strip_section_breaks(element) -> None:
    for section_properties in list(element.iter(qn("w:sectPr"))):
        section_properties.getparent().remove(section_properties)


def _strip_edit_session_ids(element) -> None:
    removable = (
        qn("w14:paraId"),
        qn("w14:textId"),
        qn("w:rsidR"),
        qn("w:rsidRDefault"),
        qn("w:rsidP"),
        qn("w:rsidRPr"),
    )
    for node in element.iter():
        for attribute in removable:
            node.attrib.pop(attribute, None)


def _append_element(document: Document, element) -> None:
    body = document._element.body
    body.insert(len(body) - 1, element)


def _append_text_paragraph(document: Document, text: str, centred: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    if centred:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")


def _cell_text(cell) -> str:
    text_parts: list[str] = []
    for node in cell._tc.iter():
        if node.tag == qn("w:p"):
            # python-docx exposes cell paragraphs as newlines. Keep that
            # boundary when a source table is reconstructed for the template.
            if text_parts and text_parts[-1] != "\n":
                text_parts.append("\n")
        elif node.tag in {qn("w:t"), qn("m:t")} and node.text:
            text_parts.append(node.text)
        elif node.tag == qn("w:br"):
            text_parts.append("\n")
    return "".join(text_parts).strip()


def _install_table_style(
    document: Document,
    source: Document,
    source_name: str,
    imported_style_id: str,
    imported_name: str,
) -> tuple[str, str]:
    """Map a source table style to an existing or collision-free target style."""
    styles = document.styles.element
    source_style = next(
        (
            style
            for style in source.styles.element.findall(qn("w:style"))
            if style.get(qn("w:type")) == "table"
            and style.find(qn("w:name")) is not None
            and style.find(qn("w:name")).get(qn("w:val")) == source_name
        ),
        None,
    )
    if source_style is None:
        raise RuntimeError(f"Technical report is missing table style: {source_name}")

    source_style_id = source_style.get(qn("w:styleId"))
    for style in styles.findall(qn("w:style")):
        if style.get(qn("w:type")) != "table":
            continue
        name = style.find(qn("w:name"))
        if style.get(qn("w:styleId")) == imported_style_id:
            return source_style_id, imported_style_id
        if name is not None and name.get(qn("w:val")) == source_name:
            return source_style_id, style.get(qn("w:styleId"))

    cloned_style = copy.deepcopy(source_style)
    cloned_style.set(qn("w:styleId"), imported_style_id)
    cloned_style.find(qn("w:name")).set(qn("w:val"), imported_name)
    based_on = cloned_style.find(qn("w:basedOn"))
    if based_on is not None:
        cloned_style.remove(based_on)
    styles.append(cloned_style)
    return source_style_id, imported_style_id


def _set_table_style_id(table, style_id: str) -> None:
    table_properties = table._tbl.tblPr
    table_style = table_properties.find(qn("w:tblStyle"))
    if table_style is None:
        from docx.oxml import OxmlElement

        table_style = OxmlElement("w:tblStyle")
        table_properties.insert(0, table_style)
    table_style.set(qn("w:val"), style_id)


def _copy_math_cell_paragraphs(target_cell, source_cell) -> None:
    """Keep OMML while using a new target table structure."""
    target_tc = target_cell._tc
    for paragraph in list(target_tc.findall(qn("w:p"))):
        target_tc.remove(paragraph)

    source_paragraphs = list(source_cell._tc.findall(qn("w:p")))
    for source_paragraph in source_paragraphs:
        clone = copy.deepcopy(source_paragraph)
        _strip_section_breaks(clone)
        _strip_edit_session_ids(clone)
        target_tc.append(clone)

    if not source_paragraphs:
        from docx.oxml import OxmlElement

        target_tc.append(OxmlElement("w:p"))


def _append_table(
    document: Document,
    source_table,
    style_map: dict[str, str],
    preserve_table_xml: bool,
) -> None:
    has_math = any(
        True
        for row in source_table.rows
        for cell in row.cells
        for _ in cell._tc.iter(qn("m:oMath"))
    )
    has_images = any(
        True
        for row in source_table.rows
        for cell in row.cells
        for _ in cell._tc.iter(qn("a:blip"))
    )
    if has_images:
        raise RuntimeError("Technical material contains an image inside a table")

    source_properties = source_table._tbl.tblPr
    source_style = source_properties.find(qn("w:tblStyle"))
    source_style_id = source_style.get(qn("w:val")) if source_style is not None else None
    target_style_id = style_map.get(source_style_id) if source_style_id else None

    if preserve_table_xml:
        clone = copy.deepcopy(source_table._tbl)
        _strip_section_breaks(clone)
        _strip_edit_session_ids(clone)
        table_properties = clone.find(qn("w:tblPr"))
        table_style = table_properties.find(qn("w:tblStyle")) if table_properties is not None else None
        if table_style is not None and target_style_id:
            table_style.set(qn("w:val"), target_style_id)
        _append_element(document, clone)
        return

    rows = [[_cell_text(cell) for cell in row.cells] for row in source_table.rows]
    column_count = max((len(row) for row in rows), default=1)
    table = document.add_table(rows=len(rows), cols=column_count)
    if target_style_id:
        _set_table_style_id(table, target_style_id)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            source_cell = source_table.rows[row_index].cells[column_index]
            target_cell = table.cell(row_index, column_index)
            if any(True for _ in source_cell._tc.iter(qn("m:oMath"))):
                _copy_math_cell_paragraphs(target_cell, source_cell)
            else:
                target_cell.text = value


def _renumber_body_drawing_ids(document: Document) -> None:
    for number, drawing in enumerate(document._element.body.iter(qn("wp:docPr")), start=1):
        drawing.set("id", str(number))
        drawing.set("name", f"MoSimCourseFigure{number}")


def import_paragraph_range(
    document: Document,
    technical_report: Path,
    start_paragraph: int,
    end_paragraph: int | None,
    section_heading: str,
    preserve_table_xml: bool = True,
) -> ImportStats:
    """Import a half-open body range while rebuilding image relationships."""
    source = Document(technical_report)
    source_three_line_style, target_three_line_style = _install_table_style(
        document,
        source,
        source_name="三线表",
        imported_style_id="MoSimThreeLineTable",
        imported_name="三线表",
    )
    source_grid_style, target_grid_style = _install_table_style(
        document,
        source,
        source_name="Table Grid",
        imported_style_id="MoSimTableGrid",
        imported_name="MoSim 网格表",
    )
    paragraphs = source.paragraphs
    if not 0 <= start_paragraph < len(paragraphs):
        raise ValueError(f"Invalid start paragraph: {start_paragraph}")
    if end_paragraph is not None and not start_paragraph < end_paragraph <= len(paragraphs):
        raise ValueError(f"Invalid end paragraph: {end_paragraph}")

    paragraph_positions = {paragraph._element: index for index, paragraph in enumerate(paragraphs)}
    paragraph_by_element = {paragraph._element: paragraph for paragraph in paragraphs}
    tables_by_element = {table._tbl: table for table in source.tables}
    in_range = False
    copied_paragraphs = 0
    copied_tables = 0
    copied_images = 0
    source_image_ids: set[str] = set()
    target_relationships: dict[str, str] = {}
    style_map = {
        source_three_line_style: target_three_line_style,
        source_grid_style: target_grid_style,
    }

    _add_heading(document, section_heading, 1)
    with tempfile.TemporaryDirectory(prefix="mosim-course-media-") as temporary_directory:
        media_directory = Path(temporary_directory)
        for child in source._element.body.iterchildren():
            if child.tag == qn("w:p"):
                position = paragraph_positions.get(child)
                if position == start_paragraph:
                    in_range = True
                if position == end_paragraph:
                    break
            if not in_range or child.tag == qn("w:sectPr"):
                continue

            clone = copy.deepcopy(child)
            _strip_section_breaks(clone)
            _strip_edit_session_ids(clone)
            image_refs = list(clone.iter(qn("a:blip")))
            for blip in image_refs:
                old_relationship = blip.get(qn("r:embed"))
                if not old_relationship:
                    continue
                source_image_ids.add(old_relationship)
                if old_relationship not in target_relationships:
                    source_part = source.part.related_parts[old_relationship]
                    filename = Path(str(source_part.partname)).name
                    media_path = media_directory / filename
                    media_path.write_bytes(source_part.blob)
                    new_relationship, _ = document.part.get_or_add_image(str(media_path))
                    target_relationships[old_relationship] = new_relationship
                blip.set(qn("r:embed"), target_relationships[old_relationship])
            copied_images += len(image_refs)

            if child.tag == qn("w:p"):
                original = paragraph_by_element[child]
                if original.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
                    _add_heading(
                        document,
                        _normalise_heading(original.text),
                        {"Heading 1": 2, "Heading 2": 3, "Heading 3": 3}[original.style.name],
                    )
                elif image_refs or list(clone.iter(qn("m:oMath"))):
                    _set_normal_style(clone)
                    _append_element(document, clone)
                else:
                    _append_text_paragraph(
                        document,
                        original.text,
                        centred=original.style.name == "图表标题",
                    )
                copied_paragraphs += 1
            elif child.tag == qn("w:tbl"):
                _append_table(
                    document,
                    tables_by_element[child],
                    style_map,
                    preserve_table_xml,
                )
                copied_tables += 1

    _renumber_body_drawing_ids(document)
    return ImportStats(
        paragraphs=copied_paragraphs,
        tables=copied_tables,
        image_references=copied_images,
        unique_source_images=len(source_image_ids),
    )


def image_reference_count(document_path: Path) -> int:
    document = Document(document_path)
    return sum(1 for _ in document._element.body.iter(qn("a:blip")))
