#!/usr/bin/env python3
"""Build the two course reports from Markdown with the national-report template.

This adapter intentionally reuses the simulation-report export primitives. It
does not rewrite either Markdown source: source images, captions, tables, and
formula blocks remain authoritative.
"""

from __future__ import annotations

import argparse
from hashlib import sha256
import json
import re
import sys
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[3]
COURSE_DIR = ROOT / "Docs" / "课设"
REPORT_DIR = ROOT / "Docs" / "报告"
TEMPLATE = REPORT_DIR / "国赛论文模版.docx"
LAYOUT_REFERENCE = REPORT_DIR / "MoSim_仿真分析报告.docx"
SIM_EXPORT_DIR = REPORT_DIR / "scripts"

if str(SIM_EXPORT_DIR) not in sys.path:
    sys.path.insert(0, str(SIM_EXPORT_DIR))

import build_competition_report_docx as sim_export  # noqa: E402


@dataclass(frozen=True)
class ReportSpec:
    key: str
    source: Path
    output: Path
    work_dir: Path
    trailing_body_line_spacing: float | None = None


REPORTS = {
    "software-construction": ReportSpec(
        key="software-construction",
        source=COURSE_DIR / "软件构造课程设计_正文.md",
        output=COURSE_DIR / "MoSim_软件构造课程设计_国赛模板版.docx",
        work_dir=ROOT / "Results" / "docx_build" / "course_reports_national_template" / "software_construction",
    ),
    "project-practice-iii": ReportSpec(
        key="project-practice-iii",
        source=COURSE_DIR / "项目综合实践III_正文.md",
        output=COURSE_DIR / "MoSim_项目综合实践III_国赛模板版.docx",
        work_dir=ROOT / "Results" / "docx_build" / "course_reports_national_template" / "project_practice_iii",
        trailing_body_line_spacing=1.1,
    ),
}

PARENTHESIZED_ITEM_RE = re.compile(r"^\((?P<number>\d+)\)\s+(?P<text>\S.+)$")
MARKDOWN_HEADING_RE = re.compile(r"^(?P<markers>#{1,6})\s+(?P<text>\S.*?)\s*$")
FENCE_RE = re.compile(r"^\s*(?P<marker>`{3,}|~{3,})(?P<info>.*?)\s*$")


@dataclass(frozen=True)
class MarkdownHeading:
    word_level: int
    text: str


@dataclass(frozen=True)
class ImagePresentation:
    width: int
    height: int
    crop: tuple[tuple[str, str], ...]


FORMULA_TAG_RE = re.compile(
    r"\\tag\{(?P<chapter>\d+)-(?P<sequence>\d+)(?P<suffix>[a-z]?)\}"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--report", choices=(*REPORTS, "all"), default="all")
    parser.add_argument("--export-pdf", action="store_true")
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def item_texts(markdown: str) -> list[str]:
    items: list[str] = []
    in_fence = False
    for line in markdown.splitlines():
        if FENCE_RE.fullmatch(line):
            in_fence = not in_fence
            continue
        if not in_fence and PARENTHESIZED_ITEM_RE.fullmatch(line.strip()):
            items.append(line.strip())
    return items


def visible_item_text(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def course_pandoc_markdown(markdown: str) -> str:
    """Preserve headings and emit parenthesized items as separate body paragraphs."""
    lines = markdown.splitlines(keepends=True)
    output: list[str] = []
    in_fence = False
    for index, line in enumerate(lines):
        if FENCE_RE.fullmatch(line.rstrip("\r\n")):
            in_fence = not in_fence
        output.append(line)
        if in_fence or not PARENTHESIZED_ITEM_RE.fullmatch(line.strip()):
            continue
        if index + 1 >= len(lines) or lines[index + 1].strip():
            output.append("\r\n" if line.endswith("\r\n") else "\n")
    return sim_export.make_pandoc_markdown("".join(output))


def markdown_headings(markdown: str) -> list[MarkdownHeading]:
    """Map source headings below the document title to their Word heading level."""
    headings: list[MarkdownHeading] = []
    in_fence = False
    title_seen = False
    for line in markdown.splitlines():
        if FENCE_RE.fullmatch(line):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        match = MARKDOWN_HEADING_RE.fullmatch(line)
        if match is None:
            continue
        source_level = len(match.group("markers"))
        if source_level == 1 and not title_seen:
            title_seen = True
            continue
        if source_level == 1:
            raise ValueError("Course Markdown contains more than one level-one heading.")
        headings.append(MarkdownHeading(source_level - 1, match.group("text")))
    if not title_seen:
        raise ValueError("Course Markdown has no document title heading.")
    return headings


def code_fence_count(markdown: str) -> int:
    """Count non-LaTeX fenced blocks, which must become one-cell Word tables."""
    count = 0
    in_fence = False
    fence_info = ""
    for line in markdown.splitlines():
        match = FENCE_RE.fullmatch(line)
        if match is None:
            continue
        if not in_fence:
            in_fence = True
            fence_info = match.group("info").strip().casefold()
            continue
        if fence_info != "latex":
            count += 1
        in_fence = False
        fence_info = ""
    if in_fence:
        raise ValueError("Course Markdown has an unterminated fenced block.")
    return count


def prepare_course_formula_manifest(
    source: Path, work_dir: Path
) -> tuple[Path, list[dict[str, Any]]]:
    """Build a course-local formula manifest without the simulation report's 108-formula gate."""
    markdown = source.read_text(encoding="utf-8")
    formulas: list[dict[str, Any]] = []
    previous_chapter = -1
    seen_numbers: set[str] = set()
    chapter_sequences: dict[int, list[int]] = {}

    for formula_id, match in enumerate(sim_export.LATEX_FENCE_RE.finditer(markdown), start=1):
        source_tex = match.group(1).strip()
        if not (source_tex.startswith(r"\[") and source_tex.endswith(r"\]")):
            raise ValueError(f"Display formula {formula_id:03d} lacks \\[...\\] delimiters")
        tag_matches = list(FORMULA_TAG_RE.finditer(source_tex))
        if len(tag_matches) != 1:
            raise ValueError(
                f"Display formula {formula_id:03d} must contain exactly one \\tag{{chapter-sequence}}"
            )
        tag = tag_matches[0]
        chapter = int(tag.group("chapter"))
        sequence = int(tag.group("sequence"))
        suffix = tag.group("suffix")
        expected_number = f"{chapter}-{sequence}{suffix}"
        if chapter < previous_chapter:
            raise ValueError(f"Equation chapter order regressed at {expected_number}")
        if expected_number in seen_numbers:
            raise ValueError(f"Duplicate equation number {expected_number}")

        previous_chapter = chapter
        seen_numbers.add(expected_number)
        chapter_sequences.setdefault(chapter, []).append(sequence)
        source_start_line = markdown[: match.start(1)].count("\n") + 1
        source_end_line = markdown[: match.end(1)].count("\n") + 1
        body = source_tex[2:-2]
        math_body = FORMULA_TAG_RE.sub("", body).strip()
        formulas.append(
            {
                "formula_id": formula_id,
                "source_start_line": source_start_line,
                "source_end_line": source_end_line,
                "chapter": chapter,
                "sequence": sequence,
                "suffix": suffix,
                "expected_number": expected_number,
                "equation_field_instruction": f"SEQ Equation \\r {sequence} \\* ARABIC",
                "source_tex": source_tex,
                "math_body": math_body,
                "source_tex_sha256": sha256(source_tex.encode("utf-8")).hexdigest(),
            }
        )

    if not formulas:
        raise ValueError("Course source contains no display formulas.")
    for chapter, sequences in chapter_sequences.items():
        expected = list(range(1, max(sequences) + 1))
        if sorted(set(sequences)) != expected:
            raise ValueError(
                f"Chapter {chapter} equation numbers are {sorted(set(sequences))}, expected {expected}"
            )

    manifest_path = work_dir / "mathtype_formula_manifest.json"
    manifest_path.write_text(
        json.dumps(
            {
                "schema": "mosim.course.formula_manifest.v1",
                "source": str(source),
                "display_formula_count": len(formulas),
                "formulas": formulas,
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    return manifest_path, formulas


def heading_style(document: Document, level: int):
    """Use existing heading styles, adding semantic deeper levels when needed."""
    name = f"Heading {level}"
    try:
        return sim_export.find_style(document, (name,), WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        fallback = sim_export.find_style(document, ("Heading 3",), WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = fallback
        properties = OxmlElement("w:pPr")
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level - 1))
        properties.append(outline)
        style.element.append(properties)
        return style


def _caption_counts(records: list[sim_export.CaptionRecord]) -> Counter[str]:
    return Counter("figure" if record.kind == "图" else "table" for record in records)


def _caption_key(kind: str, chapter: int | str, sequence: int | str, title: str) -> tuple[str, int, int, str]:
    return (kind, int(chapter), int(sequence), " ".join(title.split()))


def _inline_shape_image_hash(document: Document, shape) -> str | None:
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    relationship_id = blip.get(qn("r:embed"))
    if relationship_id is None:
        return None
    return sha256(document.part.related_parts[relationship_id].blob).hexdigest()


def _inline_shape_crop(shape) -> tuple[tuple[str, str], ...]:
    blip_fill = shape._inline.graphic.graphicData.pic.blipFill
    source_rect = blip_fill.find(qn("a:srcRect"))
    return tuple(sorted(source_rect.attrib.items())) if source_rect is not None else ()


def reference_image_presentations(reference: Path) -> dict[str, list[ImagePresentation]]:
    """Index source-report image geometry so course documents retain reviewed framing."""
    document = Document(reference)
    presentations: dict[str, list[ImagePresentation]] = {}
    for shape in document.inline_shapes:
        image_hash = _inline_shape_image_hash(document, shape)
        if image_hash is None:
            continue
        presentations.setdefault(image_hash, []).append(
            ImagePresentation(shape.width, shape.height, _inline_shape_crop(shape))
        )
    return presentations


def _set_inline_shape_crop(shape, crop: tuple[tuple[str, str], ...]) -> None:
    blip_fill = shape._inline.graphic.graphicData.pic.blipFill
    for source_rect in blip_fill.findall(qn("a:srcRect")):
        blip_fill.remove(source_rect)
    if not crop:
        return
    source_rect = OxmlElement("a:srcRect")
    for key, value in crop:
        source_rect.set(key, value)
    blip = blip_fill.find(qn("a:blip"))
    insertion_index = list(blip_fill).index(blip) + 1 if blip is not None else 0
    blip_fill.insert(insertion_index, source_rect)


def apply_reference_image_presentations(
    document: Document,
    presentations: dict[str, list[ImagePresentation]],
) -> dict[str, int]:
    """Restore source-report image dimensions and crop for images reused by a course report."""
    uses: dict[str, int] = {}
    stats = {"matched": 0, "unmatched": 0}
    for shape in document.inline_shapes:
        image_hash = _inline_shape_image_hash(document, shape)
        candidates = presentations.get(image_hash or "", [])
        if not candidates:
            stats["unmatched"] += 1
            continue
        occurrence = uses.get(image_hash, 0)
        presentation = candidates[min(occurrence, len(candidates) - 1)]
        uses[image_hash] = occurrence + 1
        shape.width = presentation.width
        shape.height = presentation.height
        _set_inline_shape_crop(shape, presentation.crop)
        stats["matched"] += 1
    return stats


def validate_reference_image_presentations(
    document: Document,
    presentations: dict[str, list[ImagePresentation]],
) -> tuple[dict[str, int], list[str]]:
    uses: dict[str, int] = {}
    stats = {"matched": 0, "unmatched": 0}
    errors: list[str] = []
    for index, shape in enumerate(document.inline_shapes, start=1):
        image_hash = _inline_shape_image_hash(document, shape)
        candidates = presentations.get(image_hash or "", [])
        if not candidates:
            stats["unmatched"] += 1
            continue
        occurrence = uses.get(image_hash, 0)
        expected = candidates[min(occurrence, len(candidates) - 1)]
        uses[image_hash] = occurrence + 1
        if (shape.width, shape.height, _inline_shape_crop(shape)) != (
            expected.width,
            expected.height,
            expected.crop,
        ):
            errors.append(f"Image {index} does not retain the reference-report presentation")
        stats["matched"] += 1
    return stats, errors


def remove_page_number_fields(document: Document) -> int:
    removed = 0
    for section in document.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in footer.paragraphs:
                if "PAGE" in paragraph._p.xml.upper():
                    sim_export.clear_paragraph(paragraph)
                    removed += 1
    return removed


def _minimum_table_column_width(total_width: int, column_count: int) -> int:
    preferred = {2: 1800, 3: 1000, 4: 900, 5: 750, 6: 600}
    return min(preferred.get(column_count, total_width // (column_count * 2)), total_width // column_count)


def apply_course_table_geometry(table) -> list[int]:
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        raise ValueError("Course table has no grid.")
    columns = grid.findall(qn("w:gridCol"))
    column_count = len(columns)
    if column_count != len(table.columns):
        raise ValueError("Course table grid does not match its column count.")
    total_width = sum(int(column.get(qn("w:w"), "0")) for column in columns)
    if total_width <= 0:
        raise ValueError("Course table has no positive grid width.")

    scores: list[float] = []
    for index in range(column_count):
        lengths = [
            min(80, max(1, len(re.sub(r"\\s+", "", row.cells[index].text))))
            for row in table.rows
        ]
        scores.append(max(10.0, sum(lengths) / len(lengths)))

    minimum = _minimum_table_column_width(total_width, column_count)
    remaining = total_width - minimum * column_count
    if remaining < 0:
        raise ValueError("Course table minimum widths exceed the table width.")
    score_total = sum(scores)
    widths = [minimum + round(remaining * score / score_total) for score in scores]
    widths[-1] += total_width - sum(widths)

    table.autofit = False
    sim_export.set_dxa_width(table._tbl.tblPr, "w:tblW", total_width)
    for column, width in zip(columns, widths, strict=True):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            sim_export.set_dxa_width(cell._tc.get_or_add_tcPr(), "w:tcW", width)
    return widths


def _table_total_width(table) -> int:
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return 0
    return sum(
        int(column.get(qn("w:w"), "0"))
        for column in grid.findall(qn("w:gridCol"))
    )


def _course_body_table_width(document: Document) -> int:
    for table in document.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            width = _table_total_width(table)
            if width > 0:
                return width
    section = document.sections[0]
    return section.page_width.twips - section.left_margin.twips - section.right_margin.twips


def replace_code_fences_with_one_cell_tables(
    document: Document,
    normal_style,
    table_cell_layout,
) -> int:
    """Move each Pandoc code paragraph into a borderless one-cell body table."""
    code_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.name == "Source Code"
    ]
    if not code_paragraphs:
        return 0

    total_width = _course_body_table_width(document)
    if total_width <= 0:
        raise ValueError("Course code table has no positive body width.")

    for paragraph in code_paragraphs:
        table = document.add_table(rows=1, cols=1)
        table.style = document.styles["Normal Table"]
        table.alignment = 1
        table.autofit = False
        sim_export.set_borderless_table_borders(table)
        sim_export.set_dxa_width(table._tbl.tblPr, "w:tblW", total_width)
        grid = table._tbl.find(qn("w:tblGrid"))
        if grid is None:
            raise ValueError("New course code table has no column grid.")
        columns = grid.findall(qn("w:gridCol"))
        if len(columns) != 1:
            raise ValueError("New course code table is not one-column.")
        columns[0].set(qn("w:w"), str(total_width))
        cell = table.cell(0, 0)
        sim_export.set_dxa_width(cell._tc.get_or_add_tcPr(), "w:tcW", total_width)

        paragraph.style = normal_style
        sim_export.apply_table_cell_text_layout(paragraph, table_cell_layout)
        paragraph._p.addprevious(table._tbl)
        placeholder = cell.paragraphs[0]._p
        placeholder.getparent().replace(placeholder, paragraph._p)
    return len(code_paragraphs)


def validate_course_table_geometry(document: Document) -> list[str]:
    errors: list[str] = []
    for index, table in enumerate(document.tables, 1):
        if len(table.rows) < 2 or len(table.columns) < 2:
            continue
        grid = table._tbl.find(qn("w:tblGrid"))
        widths = [] if grid is None else [
            int(column.get(qn("w:w"), "0")) for column in grid.findall(qn("w:gridCol"))
        ]
        if len(widths) != len(table.columns) or not widths:
            errors.append(f"Table {index} has an invalid column grid")
            continue
        minimum = _minimum_table_column_width(sum(widths), len(widths))
        if table.autofit or min(widths) < minimum:
            errors.append(f"Table {index} has unstable column geometry")
    return errors


def separate_adjacent_formula_tables(document: Document, normal_style) -> int:
    """Keep Word from coalescing adjacent borderless equation tables on save."""
    inserted = 0
    for table in list(document.tables):
        if "<m:oMath" not in table._tbl.xml:
            continue
        next_sibling = table._tbl.getnext()
        if next_sibling is None or next_sibling.tag != qn("w:tbl"):
            continue
        if "<m:oMath" not in next_sibling.xml:
            continue
        spacer = document.add_paragraph()
        spacer.style = normal_style
        spacer.paragraph_format.first_line_indent = 0
        spacer.paragraph_format.space_before = 0
        spacer.paragraph_format.space_after = 0
        spacer.paragraph_format.line_spacing = Pt(1)
        table._tbl.addnext(spacer._p)
        inserted += 1
    return inserted


def compact_trailing_body_paragraph(
    document: Document, line_spacing: float | None
) -> dict[str, Any] | None:
    """Eliminate a one-line trailing page without changing document styles."""
    if line_spacing is None:
        return None
    for index in range(len(document.paragraphs) - 1, -1, -1):
        paragraph = document.paragraphs[index]
        if not paragraph.text.strip():
            continue
        paragraph.paragraph_format.line_spacing = line_spacing
        return {"paragraph_index": index + 1, "line_spacing": line_spacing}
    raise ValueError("Cannot compact a document with no trailing body paragraph")


def validate_code_fence_tables(document: Document, expected_count: int, table_cell_layout) -> list[str]:
    code_tables = [
        table
        for table in document.tables
        if len(table.rows) == 1 and len(table.columns) == 1
    ]
    errors: list[str] = []
    if len(code_tables) != expected_count:
        return [f"One-cell code table count {len(code_tables)} does not match fenced code {expected_count}"]

    for table_index, table in enumerate(code_tables, start=1):
        if table.style.name != "Normal Table":
            errors.append(f"Code table {table_index} does not use the Normal Table convention")
        if table.autofit:
            errors.append(f"Code table {table_index} has unstable autofit geometry")
        paragraph = table.cell(0, 0).paragraphs[0]
        if paragraph.style.name != "Normal":
            errors.append(f"Code table {table_index} does not contain a Normal body paragraph")
        properties = paragraph._p.pPr
        run_properties = properties.find(qn("w:rPr")) if properties is not None else None
        indent = properties.find(qn("w:ind")) if properties is not None else None
        size = run_properties.find(qn("w:sz")) if run_properties is not None else None
        size_cs = run_properties.find(qn("w:szCs")) if run_properties is not None else None
        if indent is None or indent.get(qn("w:firstLine")) != "0" or indent.get(qn("w:firstLineChars")) != "0":
            errors.append(f"Code table {table_index} cell has a first-line indent")
        if size is None or size.get(qn("w:val")) != str(table_cell_layout.table_size):
            errors.append(f"Code table {table_index} cell size does not match the five-point table contract")
        if size_cs is None or size_cs.get(qn("w:val")) != str(table_cell_layout.table_size_cs):
            errors.append(f"Code table {table_index} cell complex-script size does not match the table contract")
    return errors


def course_caption_records(markdown: str) -> list[sim_export.CaptionRecord]:
    """Read captions using the course source's documented media ordering."""
    lines = markdown.splitlines()
    source_records = sim_export.caption_records(markdown)
    records_by_line = {record.line_number: record for record in source_records}
    records: list[sim_export.CaptionRecord] = []

    for index, line in enumerate(lines):
        if not line.strip().startswith("!["):
            continue
        following_index = next(
            (candidate for candidate in range(index + 1, len(lines)) if lines[candidate].strip()),
            None,
        )
        following = (
            records_by_line.get(following_index + 1)
            if following_index is not None
            else None
        )
        if following is None or following.kind != "图":
            raise ValueError(
                f"Image at source line {index + 1} is not followed by a figure caption."
            )
        records.append(following)

    for record in source_records:
        index = record.line_number - 1
        following = next(
            (line.strip() for line in lines[index + 1 :] if line.strip()), ""
        )
        if record.kind == "表" and following.startswith("|"):
            records.append(record)
    figure_labels = [record.label for record in records if record.kind == "图"]
    duplicate_labels = sorted({label for label in figure_labels if figure_labels.count(label) > 1})
    if duplicate_labels:
        raise ValueError(
            "Duplicate figure captions in course source: " + ", ".join(duplicate_labels)
        )
    return sorted(records, key=lambda record: record.line_number)


def _paragraph_num_id(paragraph) -> str | None:
    properties = paragraph._p.pPr
    if properties is None:
        return None
    numbering = properties.find(qn("w:numPr"))
    if numbering is None:
        return None
    number_id = numbering.find(qn("w:numId"))
    return number_id.get(qn("w:val")) if number_id is not None else None


def postprocess_course_docx(
    content_docx: Path,
    document_title: str,
    expected_chapters: list[str],
    expected_headings: list[MarkdownHeading],
    captions: list[sim_export.CaptionRecord],
    formula_records: list[dict[str, Any]],
    formula_layout,
    table_cell_layout,
    title_presentation,
    image_presentations: dict[str, list[ImagePresentation]],
    trailing_body_line_spacing: float | None,
) -> dict[str, Any]:
    document = Document(content_docx)
    caption_style = sim_export.find_style(
        document, ("图表标题", "Caption"), WD_STYLE_TYPE.PARAGRAPH
    )
    normal_style = sim_export.find_style(document, ("Normal",), WD_STYLE_TYPE.PARAGRAPH)
    heading_one = sim_export.find_style(
        document, ("Heading 1", "heading 1"), WD_STYLE_TYPE.PARAGRAPH
    )
    heading_styles = {
        level: heading_style(document, level)
        for level in {record.word_level for record in expected_headings} | {1}
    }
    three_line = sim_export.find_style(document, ("三线表",), WD_STYLE_TYPE.TABLE)

    caption_counts = _caption_counts(captions)
    caption_index = 0
    chapter_index = 0
    heading_index = 0
    title_seen = False
    expected_caption_labels: dict[str, list[str]] = {"figure": [], "table": []}
    valid_caption_keys = {
        _caption_key(record.kind, record.chapter, record.sequence, record.title)
        for record in captions
    }

    for paragraph in document.paragraphs:
        if "<w:drawing" in paragraph._p.xml:
            paragraph.style = normal_style
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.alignment = 1

        text = paragraph.text.strip()
        if text == document_title:
            if title_seen:
                raise ValueError(f"Duplicate title in generated document: {document_title!r}")
            sim_export.apply_heading_presentation(
                paragraph, document_title, heading_one, title_presentation
            )
            title_seen = True
            continue

        if heading_index < len(expected_headings) and text == expected_headings[heading_index].text:
            heading = expected_headings[heading_index]
            heading_index += 1
            if heading.word_level == 1:
                if paragraph.style.name.casefold() != heading_one.name.casefold():
                    raise ValueError(f"Markdown chapter is not a Word Heading 1: {text!r}")
                if chapter_index >= len(expected_chapters):
                    raise ValueError(f"Unexpected first-level heading: {text!r}")
                title = sim_export.CHAPTER_PREFIX_RE.sub("", text).strip()
                expected = expected_chapters[chapter_index]
                if title != expected:
                    raise ValueError(
                        f"Chapter {chapter_index + 1} is {title!r}, expected {expected!r}."
                    )
                sim_export.clear_paragraph(paragraph)
                paragraph.style = heading_one
                chapter_instruction = " SEQ Chapter \\r 1 " if chapter_index == 0 else " SEQ Chapter "
                sim_export.append_hidden_field(paragraph, chapter_instruction, str(chapter_index + 1))
                sim_export.append_hidden_field(paragraph, " SEQ Equation \\r 0 ", "0")
                paragraph.add_run(title)
                chapter_index += 1
                continue

            paragraph.style = heading_styles[heading.word_level]
            title = sim_export.SUBHEADING_PREFIX_RE.sub("", text).strip()
            if title != text:
                sim_export.clear_paragraph(paragraph)
                paragraph.add_run(title)
            continue

        caption_match = sim_export.CAPTION_RE.match(text)
        if not caption_match:
            continue
        candidate_key = _caption_key(
            caption_match.group("kind"),
            caption_match.group("chapter"),
            caption_match.group("sequence"),
            caption_match.group("title"),
        )
        if candidate_key not in valid_caption_keys:
            continue
        if caption_index >= len(captions):
            raise ValueError(f"Unexpected caption: {text!r}")

        expected = captions[caption_index]
        caption_index += 1
        kind = caption_match.group("kind")
        chapter = int(caption_match.group("chapter"))
        sequence = int(caption_match.group("sequence"))
        if (kind, chapter, sequence) != (expected.kind, expected.chapter, expected.sequence):
            raise ValueError(
                f"Caption order mismatch at source line {expected.line_number}: {text!r}"
            )
        if chapter != chapter_index:
            raise ValueError(
                f"Caption {expected.label} follows chapter {chapter_index}, not chapter {chapter}."
            )

        sim_export.clear_paragraph(paragraph)
        paragraph.style = caption_style
        paragraph.alignment = 1
        field_name = "Figure" if kind == "图" else "Table"
        field_key = "figure" if kind == "图" else "table"
        expected_caption_labels[field_key].append(expected.label)
        paragraph.add_run(kind + " ")
        sim_export.append_field(paragraph, " SEQ Chapter \\c ", str(chapter_index))
        paragraph.add_run("-")
        sim_export.append_field(
            paragraph,
            f" SEQ {field_name} \\r {sequence} \\* ARABIC ",
            str(sequence),
        )
        paragraph.add_run("　" + expected.title)

    if not title_seen:
        raise ValueError("Pandoc did not create the document title paragraph.")
    if chapter_index != len(expected_chapters):
        raise ValueError(
            f"Chapter count {chapter_index} does not match source {len(expected_chapters)}."
        )
    if heading_index != len(expected_headings):
        missing = expected_headings[heading_index]
        raise ValueError(f"Markdown heading was not found in Pandoc output: {missing.text!r}")
    if caption_index != len(captions):
        raise ValueError(
            f"Caption count {caption_index} does not match source {len(captions)}."
        )

    three_line_count = 0
    for table in document.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            sim_export.apply_three_line_table_style(table, three_line, table_cell_layout)
            apply_course_table_geometry(table)
            three_line_count += 1

    code_fence_tables = replace_code_fences_with_one_cell_tables(
        document, normal_style, table_cell_layout
    )

    inline_math_to_text = sim_export.convert_inline_math_to_body_text(document)
    formula_table_count = 0
    if formula_records:
        formula_table_count = sim_export.replace_display_math_with_formula_tables(
            document, formula_records, formula_layout
        )
    formula_table_separators = separate_adjacent_formula_tables(document, normal_style)
    image_presentation = apply_reference_image_presentations(document, image_presentations)
    trailing_body_compaction = compact_trailing_body_paragraph(
        document, trailing_body_line_spacing
    )
    page_number_fields_removed = remove_page_number_fields(document)

    sim_export.set_setting(document.settings.element, "w:doNotCompressPictures")
    sim_export.set_setting(document.settings.element, "w:updateFields", "true")
    document.save(content_docx)

    if three_line_count != caption_counts["table"]:
        raise ValueError(
            f"Three-line table count {three_line_count} does not match captions {caption_counts['table']}."
        )

    return {
        "chapters": chapter_index,
        "figures": caption_counts["figure"],
        "tables": caption_counts["table"],
        "three_line_tables": three_line_count,
        "formula_tables": formula_table_count,
        "formula_table_separators": formula_table_separators,
        "markdown_headings": len(expected_headings),
        "expected_caption_labels": expected_caption_labels,
        "inline_math_to_text": inline_math_to_text,
        "code_fence_tables": code_fence_tables,
        "image_presentation": image_presentation,
        "trailing_body_compaction": trailing_body_compaction,
        "page_number_fields_removed": page_number_fields_removed,
    }


def field_instruction_count(output: Path, field_name: str) -> int:
    pattern = re.compile(rf"\\b{re.escape(field_name)}\\b", re.IGNORECASE)
    count = 0
    with zipfile.ZipFile(output) as archive:
        for name in archive.namelist():
            if not (
                name == "word/document.xml"
                or name.startswith("word/header")
                or name.startswith("word/footer")
            ):
                continue
            text = archive.read(name).decode("utf-8")
            count += sum(
                bool(pattern.search(instruction))
                for instruction in re.findall(r"<w:instrText[^>]*>(.*?)</w:instrText>", text)
            )
            count += sum(
                bool(pattern.search(instruction))
                for instruction in re.findall(r'<w:fldSimple[^>]*w:instr="([^"]*)"', text)
            )
    return count


def validate_output(
    spec: ReportSpec,
    output: Path,
    source_markdown: str,
    captions: list[sim_export.CaptionRecord],
    formula_records: list[dict[str, Any]],
    formula_layout,
    table_cell_layout,
    converted: dict[str, Any],
    author: str,
    source_images: list[Path],
    image_presentations: dict[str, list[ImagePresentation]],
) -> dict[str, Any]:
    caption_counts = _caption_counts(captions)
    package = sim_export.inspect_docx(output)
    image_embedding = sim_export.verify_embedded_images(source_images, output)
    expected_items = item_texts(source_markdown)
    expected_headings = markdown_headings(source_markdown)
    expected_code_fences = code_fence_count(source_markdown)
    document = Document(output)
    all_paragraphs = list(sim_export.iter_all_paragraphs(document))
    actual_items = [
        paragraph
        for paragraph in document.paragraphs
        if PARENTHESIZED_ITEM_RE.fullmatch(paragraph.text.strip())
    ]
    expected_heading_counts = Counter({0: 1})
    expected_heading_counts.update(heading.word_level - 1 for heading in expected_headings)
    actual_heading_counts: Counter[int] = Counter()
    for paragraph in all_paragraphs:
        outline_level = sim_export.paragraph_outline_level(paragraph)
        if outline_level is not None:
            actual_heading_counts[outline_level] += 1

    errors: list[str] = []
    image_presentation, presentation_errors = validate_reference_image_presentations(
        document, image_presentations
    )
    if package["drawing_count"] != len(source_images):
        errors.append("Drawing count does not match source image references")
    if package["figure_field_count"] != caption_counts["figure"]:
        errors.append("Figure SEQ field count does not match source captions")
    if package["table_field_count"] != caption_counts["table"]:
        errors.append("Table SEQ field count does not match source captions")
    expected_chapter_fields = (
        converted["chapters"]
        + caption_counts["figure"]
        + caption_counts["table"]
        + len(formula_records)
    )
    if package["chapter_field_count"] != expected_chapter_fields:
        errors.append("Chapter field count does not match headings and captions")
    if package["three_line_table_count"] != caption_counts["table"]:
        errors.append("Three-line table count does not match source table captions")
    if package["office_math_count"] != len(formula_records):
        errors.append("Office Math count does not match source display formulas")
    errors.extend(presentation_errors)
    if package["source_code_paragraph_count"]:
        errors.append("Output retains Source Code paragraphs instead of the body style")
    errors.extend(validate_course_table_geometry(document))
    errors.extend(validate_code_fence_tables(document, expected_code_fences, table_cell_layout))
    if converted["code_fence_tables"] != expected_code_fences:
        errors.append("Converted code-fence table count does not match Markdown source")
    if not package["no_picture_compression"]:
        errors.append("Output does not disable picture compression")
    if not package["update_fields_on_open"]:
        errors.append("Output does not request field updates")
    if field_instruction_count(output, "PAGE"):
        errors.append("Output retains PAGE fields")
    if field_instruction_count(output, "TOC"):
        errors.append("Output retains TOC fields")
    if package["creator"] != author or package["last_modified_by"] != author:
        errors.append("Output metadata author differs from the approved template author")
    if image_embedding["missing_source_hash_count"]:
        errors.append("One or more source images were not embedded byte-for-byte")
    if [paragraph.text.strip() for paragraph in actual_items] != [
        visible_item_text(item) for item in expected_items
    ]:
        errors.append("Parenthesized body item text differs from Markdown source")
    for paragraph in actual_items:
        if paragraph.style.name != "Normal" or _paragraph_num_id(paragraph) is not None:
            errors.append("Parenthesized Markdown item is not an unnumbered Normal body paragraph")
            break
    if actual_heading_counts != expected_heading_counts:
        errors.append(
            "Word heading outline levels do not exactly match Markdown headings: "
            f"actual={dict(actual_heading_counts)}, expected={dict(expected_heading_counts)}"
        )
    if spec.trailing_body_line_spacing is not None:
        compaction = converted["trailing_body_compaction"]
        if compaction is None:
            errors.append("Configured trailing-body compaction was not applied")
        else:
            paragraph = document.paragraphs[compaction["paragraph_index"] - 1]
            if paragraph.paragraph_format.line_spacing != spec.trailing_body_line_spacing:
                errors.append("Trailing-body line spacing does not match the configured value")

    errors.extend(sim_export.validate_caption_field_contracts(package["caption_field_contracts"], captions))
    if formula_records:
        errors.extend(
            sim_export.validate_formula_table_contracts(
                package["formula_table_contracts"], formula_records, formula_layout
            )
        )
    errors.extend(
        sim_export.validate_three_line_table_contracts(
            package["three_line_table_contracts"], caption_counts["table"], table_cell_layout
        )
    )
    with zipfile.ZipFile(output) as archive:
        if archive.testzip():
            errors.append("DOCX package ZIP integrity check failed")

    if errors:
        raise RuntimeError("\n".join(errors))
    return {
        "package": package,
        "image_embedding": image_embedding,
        "image_presentation": image_presentation,
    }


def build_report(spec: ReportSpec, export_pdf: bool, overwrite: bool) -> Path:
    if not TEMPLATE.is_file() or not LAYOUT_REFERENCE.is_file() or not spec.source.is_file():
        raise FileNotFoundError("Missing source, national template, or simulation-report reference")
    if spec.output.exists() and not overwrite:
        raise FileExistsError(f"Refusing to overwrite existing output: {spec.output}")

    spec.work_dir.mkdir(parents=True, exist_ok=True)
    title, source_markdown = sim_export.read_source(spec.source)
    source_images = sim_export.resolve_report_images(
        spec.source, sim_export.image_paths(source_markdown)
    )
    captions = course_caption_records(source_markdown)
    chapters = sim_export.source_primary_chapter_titles(source_markdown)
    expected_items = item_texts(source_markdown)
    expected_headings = markdown_headings(source_markdown)
    expected_code_fences = code_fence_count(source_markdown)
    if not captions or not chapters or not expected_items or not expected_headings:
        raise ValueError("Course source is missing headings, captions, or parenthesized body items")

    formula_manifest = None
    formula_records: list[dict[str, Any]] = []
    formula_layout = None
    if sim_export.LATEX_FENCE_RE.search(source_markdown):
        formula_manifest, formula_records = prepare_course_formula_manifest(spec.source, spec.work_dir)
        formula_layout = sim_export.reference_formula_table_layout(LAYOUT_REFERENCE)

    table_cell_layout = sim_export.reference_table_cell_text_layout(LAYOUT_REFERENCE)
    title_presentation, _ = sim_export.reference_heading_presentations(LAYOUT_REFERENCE)
    image_presentations = reference_image_presentations(LAYOUT_REFERENCE)
    normalized_markdown = spec.work_dir / "normalized_report.md"
    content_docx = spec.work_dir / "pandoc_content.docx"
    pdf_output = spec.work_dir / (spec.output.stem + ".pdf") if export_pdf else None
    normalized_markdown.write_text(course_pandoc_markdown(source_markdown), encoding="utf-8")

    sim_export.run_pandoc(
        sim_export.find_pandoc(),
        normalized_markdown,
        TEMPLATE,
        (spec.source.parent, REPORT_DIR),
        content_docx,
        title,
    )
    converted = postprocess_course_docx(
        content_docx,
        title,
        chapters,
        expected_headings,
        captions,
        formula_records,
        formula_layout,
        table_cell_layout,
        title_presentation,
        image_presentations,
        spec.trailing_body_line_spacing,
    )
    word_counts = sim_export.finalize_with_word(content_docx, spec.output, pdf_output)
    author = (sim_export.template_author(TEMPLATE) or sim_export.template_author(LAYOUT_REFERENCE)).strip()
    sim_export.force_document_settings(spec.output)
    sim_export.force_document_core_properties(spec.output, title, author)
    checked = validate_output(
        spec,
        spec.output,
        source_markdown,
        captions,
        formula_records,
        formula_layout,
        table_cell_layout,
        converted,
        author,
        source_images,
        image_presentations,
    )
    manifest = {
        "report": spec.key,
        "source": str(spec.source),
        "template": str(TEMPLATE),
        "layout_reference": str(LAYOUT_REFERENCE),
        "output": str(spec.output),
        "output_bytes": spec.output.stat().st_size,
        "pdf_output": str(pdf_output) if pdf_output else None,
        "source_images": len(source_images),
        "source_captions": dict(_caption_counts(captions)),
        "source_display_formulae": len(formula_records),
        "source_markdown_headings": len(expected_headings) + 1,
        "source_parenthesized_body_items": len(expected_items),
        "source_code_fences": expected_code_fences,
        "converted": converted,
        "word": word_counts,
        **checked,
        "formula_manifest": str(formula_manifest) if formula_manifest else None,
    }
    (spec.work_dir / "build_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    # Keep the UTF-8 manifest intact while remaining safe in a GBK console.
    print(json.dumps(manifest, ensure_ascii=True, indent=2))
    return spec.output


def main() -> int:
    args = parse_args()
    selected = REPORTS.values() if args.report == "all" else (REPORTS[args.report],)
    for spec in selected:
        build_report(spec, args.export_pdf, args.overwrite)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
