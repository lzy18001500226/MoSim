from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from technical_material_import import (
    ImportStats,
    apply_report_layout,
    import_paragraph_range,
    retain_front_matter,
    validate_report_layout,
)


ROOT = Path(__file__).resolve().parents[3]
COURSE_DIR = ROOT / "Docs" / "课设"
TEMPLATE = COURSE_DIR / "项目综合实践III 封面  任务书  答辩表  目录  正文报告模板.docx"
SOURCE = COURSE_DIR / "项目综合实践III_正文.md"
TECHNICAL_SOURCE = ROOT / "Docs" / "报告" / "MoSim_仿真分析报告.docx"
OUTPUT = COURSE_DIR / "项目综合实践III项目成果报告_MoSim四旋翼仿真平台.docx"


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    run._r.append(instruction_text)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    if placeholder:
        result = paragraph.add_run(placeholder)
        set_east_asia_font(result, "宋体")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia_font(run, "宋体")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cover_field(paragraph, label: str, value: str, size: float) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(f"{label}{value}")
    run.font.size = Pt(size)
    set_east_asia_font(run, "宋体")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "808080")


def remove_body_after(document: Document, paragraph_index: int) -> None:
    """Keep the template's first body heading as the insertion anchor."""
    body = document._element.body
    anchor = document.paragraphs[paragraph_index]._element
    seen_anchor = False
    for child in list(body):
        if child is anchor:
            seen_anchor = True
            continue
        if seen_anchor and child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(3)

    for level, size in ((1, 16), (2, 14), (3, 12)):
        name = f"MoSim Heading {level}"
        if name in document.styles:
            style = document.styles[name]
        else:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)

    if "MoSim Code" not in document.styles:
        code = document.styles.add_style("MoSim Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(8.5)
        code.paragraph_format.first_line_indent = Cm(0)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(3)


def append_markdown_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_east_asia_font(run, "宋体")
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            set_east_asia_font(run, "宋体")


def new_body_paragraph(document: Document, text: str = ""):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    if text:
        append_markdown_runs(paragraph, text)
    return paragraph


def set_heading(paragraph, text: str, level: int) -> None:
    # The source heading text already has its required Chinese chapter number.
    # Custom styles keep that text intact while making it available to the TOC.
    clear_paragraph(paragraph)
    paragraph.style = paragraph.part.document.styles[f"MoSim Heading {level}"]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 11))
    set_east_asia_font(run, "黑体")


def add_heading(document: Document, text: str, level: int):
    paragraph = new_body_paragraph(document)
    set_heading(paragraph, text, level)
    return paragraph


def add_caption(document: Document, text: str, table: bool = False) -> None:
    paragraph = new_body_paragraph(document)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = table
    run.font.size = Pt(9)
    set_east_asia_font(run, "宋体")


def add_table(document: Document, rows: list[list[str]], caption: str | None = None) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalised = [row + [""] * (column_count - len(row)) for row in rows]
    if caption:
        add_caption(document, caption, table=True)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Normal Table"
    set_table_borders(table)
    table.autofit = False
    usable_width = 16.1
    for row_index, values in enumerate(normalised):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            cell.width = Cm(usable_width / column_count)
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code(document: Document, language: str, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(style="MoSim Code")
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.right_indent = Cm(0.55)
        run = paragraph.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    if language:
        paragraph = new_body_paragraph(document, f"代码语言：{language}")
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.runs[0].italic = True
        paragraph.runs[0].font.size = Pt(8.5)


def add_image(document: Document, source: Path, caption: str | None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(source), width=Inches(5.55))
    if caption:
        add_caption(document, caption)


def parse_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)*\|?", line))


def replace_cover_and_task_form(document: Document) -> None:
    title = "基于 MWORKS 的四旋翼位姿控制全链路仿真平台"
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "题目：":
            set_cover_field(paragraph, "题目：", title, 12)
        elif text == "班级：":
            set_cover_field(paragraph, "班级：", "软件2301", 14)
        elif text.startswith("学号："):
            set_cover_field(paragraph, "学号：", "231304113", 8)
        elif text == "姓名：":
            set_cover_field(paragraph, "姓名：", "刘致远", 8.5)
        elif "2026" in text and "06" in text and "月" in text:
            set_cover_field(paragraph, "", "2026年7月20日", 12)

    task_text = (
        "课题名称：基于 MWORKS 的四旋翼位姿控制全链路仿真平台。"
        "场景及意义：面向 A8 四旋翼位姿控制与仿真赛题，构建从物理建模、控制器接入、"
        "实验配置、代码生成到 ROS1/PX4/Gazebo 运行记录的可追溯工程链。"
        "具体要求：（1）建立云纵150参照虚拟机体与参数 Profile；"
        "（2）实现 48 条控制器路线的统一 Adapter 与 FormalRunner；"
        "（3）完成七场景、三机 Figure8 和 OpenBlocks 任务组织；"
        "（4）完成 px4ctrl 图形模型、C99 生成、构建和 MWORKS 内 SIL；"
        "（5）保留原始结果、指标、运行清单和失败记录，并明确证据边界。"
        "进程安排：2026年7月20日完成项目综合实践报告提交与归档。"
    )
    for table in document.tables[3:6]:
        set_cell_text(table.cell(1, 1), "231304113等5人", size=8.5)
        set_cell_text(table.cell(1, 3), "软件工程", size=9)
        set_cell_text(table.cell(2, 1), "刘致远等5人", size=8.5)
        set_cell_text(table.cell(2, 3), "软件2301", size=9)
        set_cell_text(table.cell(4, 0), task_text, size=9)


def _new_xml_paragraph(text: str, size: int, bold: bool = False, centered: bool = False):
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    if centered:
        alignment = OxmlElement("w:jc")
        alignment.set(qn("w:val"), "center")
        properties.append(alignment)
    paragraph.append(properties)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")
    run_properties.append(fonts)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(size * 2))
    run_properties.append(size_element)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    paragraph.append(run)
    return paragraph


def _new_toc_field_paragraph():
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    properties.append(spacing)
    paragraph.append(properties)
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run.append(begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\t "MoSim Heading 1,1,MoSim Heading 2,2,MoSim Heading 3,3" \\h '
    run.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run.append(separate)
    paragraph.append(run)
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "目录将在 Word 中更新"
    placeholder.append(placeholder_text)
    paragraph.append(placeholder)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph.append(end_run)
    return paragraph


def replace_static_toc(document: Document) -> None:
    body = document._element.body
    for control in body.findall(qn("w:sdt")):
        text = "".join(node.text or "" for node in control.iter(qn("w:t")))
        gallery = control.find(
            f"./{qn('w:sdtPr')}/{qn('w:docPartObj')}/{qn('w:docPartGallery')}"
        )
        is_toc = gallery is not None and gallery.get(qn("w:val")) == "Table of Contents"
        if not is_toc and "目录" not in text and "目 录" not in text:
            continue
        content = control.find(qn("w:sdtContent"))
        if content is None:
            raise RuntimeError("Template TOC content control is incomplete")
        for child in list(content):
            content.remove(child)
        content.append(_new_xml_paragraph("目录", 18, bold=True, centered=True))
        content.append(_new_toc_field_paragraph())
        return
    raise RuntimeError("Template TOC content control was not found")


def set_dynamic_page_footers(document: Document, first_section: int) -> None:
    for section in document.sections[first_section:]:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        for child in list(footer._element):
            footer._element.remove(child)
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        add_field(paragraph, " PAGE \\* MERGEFORMAT ", "1")


def render_markdown(document: Document) -> tuple[int, list[Path]]:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_heading_used = False
    image_paths: list[Path] = []
    image_count = 0
    pending_caption: str | None = None
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal pending_caption
        if not paragraph_lines:
            return
        text = " ".join(part.strip() for part in paragraph_lines).strip()
        paragraph_lines.clear()
        if text:
            new_body_paragraph(document, text)
        pending_caption = None

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if line.startswith(">"):
            line = line[1:].strip()

        heading = re.fullmatch(r"(#{1,4})\s+(.+)", line)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            if len(heading.group(1)) == 1:
                # The document title belongs on the template cover, not again
                # at the start of the report body.
                index += 1
                continue
            if not first_heading_used:
                set_heading(document.add_paragraph(), text, level)
                first_heading_used = True
            else:
                add_heading(document, text, level)
            index += 1
            continue

        if line.startswith("~~~"):
            flush_paragraph()
            language = line[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("~~~"):
                code_lines.append(lines[index])
                index += 1
            add_code(document, language, code_lines)
            index += 1
            continue

        image = re.fullmatch(r"!\[[^]]*\]\(([^)]+)\)", line)
        if image:
            flush_paragraph()
            image_path = (COURSE_DIR / image.group(1)).resolve()
            if not image_path.is_file() or COURSE_DIR not in image_path.parents:
                raise FileNotFoundError(f"Invalid report image reference: {image.group(1)}")
            caption = None
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines):
                candidate = lines[lookahead].strip().lstrip(">").strip()
                match = re.fullmatch(r"\*\*(图[^*]+)\*\*(.*)", candidate)
                if match:
                    caption = (match.group(1) + match.group(2)).strip()
                    index = lookahead
            add_image(document, image_path, caption)
            image_paths.append(image_path)
            image_count += 1
            index += 1
            continue

        caption_match = re.fullmatch(r"\*\*((?:图|表)[^*]+)\*\*(.*)", line)
        if caption_match:
            flush_paragraph()
            caption = (caption_match.group(1) + caption_match.group(2)).strip()
            if caption.startswith("表"):
                pending_caption = caption
            else:
                add_caption(document, caption)
            index += 1
            continue

        if line.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                candidate = lines[index].strip()
                if not is_table_separator(candidate):
                    table_lines.append(candidate)
                index += 1
            add_table(document, [parse_table_row(row) for row in table_lines], pending_caption)
            pending_caption = None
            continue

        if re.fullmatch(r"(?:[-*]|\d+\.)\s+.+", line):
            flush_paragraph()
            list_match = re.fullmatch(r"(?:(\d+)\.|[-*])\s+(.+)", line)
            assert list_match
            style_name = "List Number" if list_match.group(1) else "List Bullet"
            style = document.styles[style_name] if style_name in document.styles else document.styles["Normal"]
            paragraph = document.add_paragraph(style=style)
            append_markdown_runs(paragraph, list_match.group(2))
            index += 1
            continue

        if not line:
            flush_paragraph()
            index += 1
            continue

        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    return image_count, image_paths


def enable_field_update(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def validate(output: Path, image_count: int, imported: ImportStats) -> None:
    if imported.image_references != 65:
        raise RuntimeError(f"Project report must receive 65 technical figure references, got {imported.image_references}")
    document = Document(output)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    text = "\n".join(text_parts)
    required = [
        "1. 实践目的与要求", "2. 项目概述", "3. 软件需求分析", "4. 软件设计",
        "5. 软件编码与实现", "6. 软件测试", "7. 软件项目管理", "8. 软件推广与维护",
        "9. 实践总结", "10. 系统建模与控制器实现",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Generated report is missing required content: {missing}")
    if "（简述项目综合实践的目的和要求）" in text:
        raise RuntimeError("Template placeholder text survived in the report body")
    if len(document.inline_shapes) != image_count + imported.image_references:
        raise RuntimeError("Generated image count does not match direct and imported material")
    with zipfile.ZipFile(output) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"DOCX ZIP integrity failed at {bad}")
    print(json.dumps(imported.__dict__, ensure_ascii=False))


def main() -> None:
    for required in (TEMPLATE, SOURCE, TECHNICAL_SOURCE):
        if not required.is_file():
            raise FileNotFoundError(required)
    document = Document(TEMPLATE)
    replace_cover_and_task_form(document)
    retain_front_matter(document, ((0, 31), (82, 87), (92, 97)))
    configure_styles(document)
    replace_static_toc(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    set_dynamic_page_footers(document, len(document.sections) - 1)
    image_count, _ = render_markdown(document)
    imported = import_paragraph_range(
        document,
        TECHNICAL_SOURCE,
        start_paragraph=1,
        end_paragraph=498,
        section_heading="10. 系统建模与控制器实现",
        preserve_table_xml=False,
    )
    layout = apply_report_layout(document)
    enable_field_update(document)
    document.core_properties.title = "项目综合实践III报告：MoSim 四旋翼仿真平台"
    document.core_properties.subject = "基于 MWORKS 的四旋翼位姿控制全链路仿真平台"
    document.core_properties.comments = "项目综合实践 III 独立课程报告。"
    document.save(OUTPUT)
    validate(OUTPUT, image_count, imported)
    validate_report_layout(OUTPUT)
    print(OUTPUT)
    print(json.dumps({"direct_images": image_count, **imported.__dict__, **layout}, ensure_ascii=False))


if __name__ == "__main__":
    main()
