from __future__ import annotations

import io
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from technical_material_import import (
    ImportStats,
    import_paragraph_range,
)


ROOT = Path(__file__).resolve().parents[3]
TEMPLATE = ROOT / "Docs" / "报告" / "国赛论文模版.docx"
TECHNICAL_MATERIAL = ROOT / "Docs" / "报告" / "MoSim_仿真分析报告.docx"
OUTPUT = ROOT / "Docs" / "课设" / "软件构造课程设计报告_MoSim项目维护.docx"


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def set_template_font(
    run,
    *,
    east_asia: str = "宋体",
    ascii_name: str = "Times New Roman",
    size: float = 12,
) -> None:
    run.font.name = ascii_name
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), ascii_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cover_field(paragraph, label: str, value: str, size: float) -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(f"{label}{value}")
    run.font.size = Pt(size)
    set_east_asia_font(run, "宋体")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 12.0) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(14.4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    set_template_font(run, size=size)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def remove_body_after(document: Document, paragraph_index: int) -> None:
    body = document._element.body
    anchor = document.paragraphs[paragraph_index]._element
    seen = False
    for child in list(body):
        if child is anchor:
            seen = True
            body.remove(child)
            continue
        if seen and child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_template_body(document: Document) -> None:
    """Retain the user-supplied template's section settings, not its sample text."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    run._r.append(instr_text)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    if placeholder:
        result = paragraph.add_run(placeholder)
        set_east_asia_font(result, "宋体")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(0.35)
    normal.paragraph_format.line_spacing = Pt(14.4)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    template_headings = {
        1: ("Heading 1", 14, 6, 6),
        2: ("Heading 2", 12, 7.8, 7.8),
        3: ("Heading 3", 12, 2, 2),
    }
    for level, (template_name, size, before, after) in template_headings.items():
        template_style = styles[template_name]
        template_style.font.name = "黑体"
        template_style.font.size = Pt(size)
        template_style.paragraph_format.first_line_indent = Cm(0)
        template_style.paragraph_format.line_spacing = Pt(14.4)
        template_style.paragraph_format.space_before = Pt(before)
        template_style.paragraph_format.space_after = Pt(after)
        template_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        name = f"MoSim Heading {level}"
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = template_style
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = Pt(14.4)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "图表标题" not in styles:
        caption = styles.add_style("图表标题", WD_STYLE_TYPE.PARAGRAPH)
        caption.base_style = normal
    caption = styles["图表标题"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(12)
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.line_spacing = Pt(14.4)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size in (("MoSim Paper Title", 16), ("MoSim Abstract Label", 14)):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        style = styles[name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = Pt(14.4)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_report_front_matter(document: Document) -> None:
    title = document.add_paragraph(style="MoSim Paper Title")
    title_run = title.add_run("基于 MWORKS 的四旋翼位姿控制仿真平台维护")
    set_template_font(title_run, east_asia="黑体", ascii_name="黑体", size=16)

    abstract_label = document.add_paragraph(style="MoSim Abstract Label")
    label_run = abstract_label.add_run("摘  要")
    set_template_font(label_run, east_asia="黑体", ascii_name="黑体", size=14)

    abstract = add_paragraph(
        document,
        "本文以项目综合实践 III 已完成的四旋翼位姿控制与仿真系统为维护对象，围绕控制器路由、七场景配置化评测、三机 Figure8 编队与图形模型 C99/SIL 交付开展完善性维护。针对姿态桥接中 asin 输入可能越出实数定义域的问题，采用带裕量的截断变量并保留诊断标识；同时以 Adapter、FormalRunner、Profile/Contract 和结果有效性门收敛接口与配置边界。定向自动化回归覆盖姿态适配器、任务交接、Studio 目录、图形 C99 后端、生成代码运行时与 OpenBlocks 模型入口，共 45 项用例全部通过。本文保留模型结果与运行时材料的证据范围，不将静态回归或离线记录扩大为实时运行时或飞行验收结论。",
    )
    abstract.paragraph_format.first_line_indent = Cm(0.35)

    keywords = add_paragraph(document, "关键词：四旋翼仿真；软件维护；MWORKS；自动化测试；代码生成")
    keywords.paragraph_format.first_line_indent = Cm(0)


def add_paragraph(document: Document, text: str = "", bold_prefix: str | None = None):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        set_template_font(first)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_template_font(rest)
    else:
        run = paragraph.add_run(text)
        set_template_font(run)
    return paragraph


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"MoSim Heading {level}"]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    set_template_font(run, east_asia="黑体", ascii_name="黑体", size={1: 14, 2: 12, 3: 12}[level])


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.74)
    run = paragraph.add_run("- " + text)
    set_template_font(run)


def add_code(document: Document, title: str, code: str) -> None:
    caption = document.add_paragraph()
    caption.style = document.styles["图表标题"]
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(title)
    set_template_font(caption_run)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    description = OxmlElement("w:tblDescription")
    description.set(qn("w:val"), "mosim-code-listing")
    table._tbl.tblPr.append(description)
    cell = table.cell(0, 0)
    set_cell_text(cell, "\n".join(line.rstrip() for line in code.strip().splitlines()))
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "三线表"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        if widths:
            cell.width = Cm(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_figure(document: Document, source_archive: zipfile.ZipFile, media_name: str, caption: str, width: float = 5.7) -> None:
    blob = source_archive.read(f"word/media/{media_name}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(blob), width=Inches(width))
    cap = document.add_paragraph()
    cap.style = document.styles["图表标题"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_template_font(cap_run)


def _is_code_table(table) -> bool:
    description = table._tbl.tblPr.find(qn("w:tblDescription"))
    return description is not None and description.get(qn("w:val")) == "mosim-code-listing"


def _set_template_paragraph_format(paragraph, *, alignment, first_line: float = 0.35) -> None:
    paragraph.paragraph_format.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(first_line)
    paragraph.paragraph_format.line_spacing = Pt(14.4)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def apply_national_template_layout(document: Document) -> None:
    """Apply the supplied national-paper template's effective type and table rules."""
    heading_styles = {
        "MoSim Heading 1": ("Heading 1", 14, 6, 6),
        "MoSim Heading 2": ("Heading 2", 12, 7.8, 7.8),
        "MoSim Heading 3": ("Heading 3", 12, 2, 2),
        "Heading 1": ("Heading 1", 14, 6, 6),
        "Heading 2": ("Heading 2", 12, 7.8, 7.8),
        "Heading 3": ("Heading 3", 12, 2, 2),
    }
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        text = paragraph.text.strip()
        if style_name == "MoSim Paper Title":
            _set_template_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
            for run in paragraph.runs:
                set_template_font(run, east_asia="黑体", ascii_name="黑体", size=16)
        elif style_name == "MoSim Abstract Label":
            _set_template_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
            for run in paragraph.runs:
                set_template_font(run, east_asia="黑体", ascii_name="黑体", size=14)
        elif style_name in heading_styles:
            template_name, size, before, after = heading_styles[style_name]
            paragraph.style = document.styles[template_name]
            _set_template_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=0)
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)
            for run in paragraph.runs:
                set_template_font(run, east_asia="黑体", ascii_name="黑体", size=size)
        elif style_name == "图表标题" or re.match(r"^(?:图|表|代码清单)\s*\d", text):
            paragraph.style = document.styles["图表标题"]
            _set_template_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=0)
            for run in paragraph.runs:
                set_template_font(run)
        else:
            paragraph.style = document.styles["Normal"]
            _set_template_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            for run in paragraph.runs:
                set_template_font(run)

    for table in document.tables:
        code_table = _is_code_table(table)
        table.style = "Table Grid" if code_table else "三线表"
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_properties = cell._tc.get_or_add_tcPr()
                shading = cell_properties.find(qn("w:shd"))
                if shading is not None:
                    cell_properties.remove(shading)
                for paragraph in cell.paragraphs:
                    _set_template_paragraph_format(
                        paragraph,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT if code_table else WD_ALIGN_PARAGRAPH.CENTER,
                        first_line=0,
                    )
                    for run in paragraph.runs:
                        set_template_font(run)
                        if not code_table and row_index == 0:
                            run.bold = True


def replace_cover_and_forms(document: Document) -> None:
    title = "四旋翼仿真维护"
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "题目：***":
            set_cover_field(paragraph, "题目：", title, 14)
        elif text.startswith("班级："):
            set_cover_field(paragraph, "班级：", "软件2301", 14)
        elif text.startswith("姓名："):
            set_cover_field(paragraph, "姓名：", "刘致远", 12)
        elif text.startswith("学号："):
            set_cover_field(paragraph, "学号：", "231304113", 11)
        elif "2026" in text and "6月" in text:
            set_cover_field(paragraph, "", "2026年7月27日", 12)

    task_text = (
        "课题名称：基于 MWORKS 的四旋翼位姿控制仿真平台项目维护。主要目标：在项目综合实践 III "
        "完成的四旋翼控制与仿真系统基础上，完成完善性、改正性、适应性和预防性维护。具体要求："
        "（1）新增统一控制器路由、七场景配置化评测、三机 Figure8 编队、C99/SIL 交付四项功能；"
        "（2）新增模块化、可追溯性、结果有效性和可移植性四项非功能要求；"
        "（3）完成需求分析、构造设计、实现说明、测试与缺陷修复记录；"
        "（4）保留模型、配置、结果和运行记录的证据边界，不将离线仿真结果扩展为飞行或实时运行时结论。"
        "进程安排：2026年7月27日完成课程设计报告提交与归档。"
    )
    for table in document.tables[3:6]:
        set_cell_text(table.cell(1, 1), "231304113等5人", size=8.5)
        set_cell_text(table.cell(1, 3), "软件工程", size=9)
        set_cell_text(table.cell(2, 1), "刘致远等5人", size=8.5)
        set_cell_text(table.cell(2, 3), "软件2301", size=9)
        set_cell_text(table.rows[4].cells[0], task_text, size=9)
    for table in document.tables[6:9]:
        for row in table.rows:
            for cell in row.cells:
                if "题目" in cell.text and "***" in cell.text:
                    set_cell_text(cell, cell.text.replace("***", title), size=9)


def _new_xml_paragraph(text: str, size: int, bold: bool = False, centered: bool = False):
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    if centered:
        alignment = OxmlElement("w:jc")
        alignment.set(qn("w:val"), "center")
        properties.append(alignment)
    paragraph.append(properties)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")
    run_properties.append(fonts)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(size * 2))
    run_properties.append(size_element)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    paragraph.append(run)
    return paragraph


def _new_toc_field_paragraph():
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    properties.append(spacing)
    paragraph.append(properties)
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run.append(begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\t "MoSim Heading 1,1,MoSim Heading 2,2,MoSim Heading 3,3" \\h '
    run.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run.append(separate)
    paragraph.append(run)
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "目录将在 Word 中更新"
    placeholder.append(placeholder_text)
    paragraph.append(placeholder)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph.append(end_run)
    return paragraph


def replace_static_toc(document: Document) -> None:
    body = document._element.body
    for control in body.findall(qn("w:sdt")):
        gallery = control.find(
            f"./{qn('w:sdtPr')}/{qn('w:docPartObj')}/{qn('w:docPartGallery')}"
        )
        if gallery is None or gallery.get(qn("w:val")) != "Table of Contents":
            continue
        content = control.find(qn("w:sdtContent"))
        if content is None:
            raise RuntimeError("Template TOC content control is incomplete")
        for child in list(content):
            content.remove(child)
        content.append(_new_xml_paragraph("目录", 18, bold=True, centered=True))
        content.append(_new_toc_field_paragraph())
        return
    raise RuntimeError("Template TOC content control was not found")


def compact_cover_date_spacing(document: Document) -> None:
    """Keep the section-ending cover date on its cover page."""
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != "2026年7月27日":
            continue
        for spacer in paragraphs[max(0, index - 5):index]:
            if spacer.text.strip():
                continue
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(1)
        return


def load_source_facts() -> dict[str, object]:
    profile_path = ROOT / "Config" / "control_platform" / "seven_scenario_experiment_profiles_v2.json"
    injection_path = ROOT / "Config" / "control_platform" / "seven_scenario_injection_contract_v2.json"
    route_path = ROOT / "Config" / "control_platform" / "model_studio_task_routes_v1.toml"
    run_record_path = ROOT / "Results" / "model_library_refactor_20260729" / "phase05_runner_layering" / "px4ctrl_climbpath_50s" / "RUN_RECORD.json"
    profile = json.loads(profile_path.read_text(encoding="utf-8"))
    injection = json.loads(injection_path.read_text(encoding="utf-8"))
    run_record = json.loads(run_record_path.read_text(encoding="utf-8"))
    route_text = route_path.read_text(encoding="utf-8")
    return {
        "scenario_count": len(profile["profiles"]),
        "scenario_ids": [item["scenario_id"] for item in injection["scenarios"]],
        "route_count": route_text.count("[[route]]"),
        "run_rows": run_record["metrics_summary"]["row_count"],
        "run_duration": run_record["metrics_summary"]["duration_s"],
        "run_rmse": run_record["metrics_summary"]["position_rmse_m"],
        "run_terminal": run_record["terminal_position_error_norm_m"],
    }


def add_body(document: Document, facts: dict[str, object]) -> None:
    add_heading(document, "1. 项目概述", 1)
    add_heading(document, "1.1 项目介绍", 2)
    add_paragraph(document, "本课程设计以项目综合实践 III 已完成的四旋翼位姿控制与仿真系统为维护对象。原系统以 MWORKS.Sysplorer、Sysblock 和 Syslab 为建模、图形化控制与数据分析环境，包含云纵 150 参照机体、Official PID 基线、px4ctrl 图形化控制器和面向多种控制算法的仿真链路。")
    add_paragraph(document, "维护的目标不是重新实现一个无人机平台，而是在既有系统、模型和实验资产之上补齐可选择、可比较、可追溯和可部署的工程能力。本报告在后续章节完整呈现维护验证涉及的公式、图表、原生表格和运行时材料，使其能够独立阅读和审查。")

    add_heading(document, "1.2 项目维护内容", 2)
    add_table(document, ["类别", "编号", "维护内容", "对应资产"], [
        ["新增功能", "F-01", "统一控制器路由与 FormalRunner 接入：按控制器标识选择已声明的全机审查 Runner 和输出边界。", "model_studio_task_routes_v1.toml；控制器 Adapter/Runner"],
        ["新增功能", "F-02", "七场景配置化 A/B 评测：以版本化 Profile 绑定轨迹、参数注入、采样和指标。", "seven_scenario_*_v2.json"],
        ["新增功能", "F-03", "三机 Figure8 编队执行：复用单机控制器、分配器与机体，增加编队参考及队形误差输出。", "Px4CtrlThreeUavFigure8Runner.mo"],
        ["新增功能", "F-04", "图形模型到 C99、SIL 和 ROS/Gazebo 部署链路的交付与核对。", "C99 生成、SIL 与运行时模块"],
        ["改正性维护", "C-01", "修正姿态桥接中 asin 输入可能越界的问题，避免 MWORKS 事件求值时进入非实数域。", "Px4CtrlAttitudeThrustAdapter.mo"],
        ["适应性维护", "A-01", "将不同控制输出语义通过 Adapter 收敛至共享机体边界，支持在同一 Plant 上替换控制器。", "Control/Adapters；Runner 边界合同"],
        ["预防性维护", "P-01", "对运行配置、模型文件和结果记录保留版本与哈希绑定，拒绝把不完整结果作为有效比较。", "Profile/Contract；RUN_RECORD.json"],
    ], [1.8, 1.6, 8.0, 5.4])

    add_heading(document, "1.3 新增关键技术", 2)
    add_bullet(document, "Adapter + FormalRunner：将控制器内部实现与统一的 ATTITUDE_THRUST、ROTOR_COMMAND、WRENCH 等输出边界隔离。")
    add_bullet(document, "Profile/Contract 驱动的实验配置：将轨迹、扰动、参数失配、指标和有效性门从控制器代码中抽离。")
    add_bullet(document, "图形模型-代码生成-SIL 闭环：以同一模型资产支持可审查图形结构和 C99 代码生成。")
    add_bullet(document, "多机复用组合：以统一单机闭环为单元，通过编队参考、槽位和误差计算扩展到三机任务。")

    add_heading(document, "2. 软件分析", 1)
    add_paragraph(document, "分析范围严格限定为本次维护后新增或强化的能力。所有性能数值均以已归档的模型结果和运行记录为依据；本次工作完成的是需求、实现和文档交付整理，并未重新执行 MWORKS、ROS 或飞行运行时试验。")

    add_heading(document, "2.1 软件功能分析", 2)
    add_table(document, ["编号", "功能需求", "输入", "处理与输出", "验收准则"], [
        ["F-01", "统一控制器路由", "controller_id、任务标识", "解析已声明的 runner_class、runner_file、boundary；拒绝未登记路由。", "路由清单可解析，路径与边界均可定位。"],
        ["F-02", "七场景配置化评测", "控制器、Profile、注入参数", "生成参数化 Runner；输出位置误差、稳态/超调或扰动窗口指标。", f"Profile 合同中存在 {facts['scenario_count']} 个场景：{', '.join(facts['scenario_ids'])}。"],
        ["F-03", "三机 Figure8 编队", "三机槽位、参考轨迹、采样周期", "三套控制器-分配器-机体实例同步执行；输出编队误差和最小机间距。", "源模型显式拥有 controller_1..3、plant_1..3 和队形误差。"],
        ["F-04", "C99/SIL 交付", "图形控制模型、代码生成配置", "生成 C99 产物并对齐 SIL 结果；保留部署链路的运行记录。", "第 7 章给出生成、SIL 与运行时边界。"],
    ], [1.4, 3.0, 3.0, 5.3, 4.1])

    add_heading(document, "2.2 非功能性分析", 2)
    add_table(document, ["编号", "非功能需求", "度量或约束", "实现与证据"], [
        ["NF-01", "模块化与可扩展性", "新增控制器不改共享 Plant；通过 Adapter 选择输出边界。", "控制器注册、Adapter 和 FormalRunner 分层；第 7 章。"],
        ["NF-02", "可追溯性与可复现性", "每次比较应绑定版本化 Profile、Contract、模型路径和结果文件。", "seven_scenario_*_v2.json；RUN_RECORD 的路径与 SHA-256。"],
        ["NF-03", "结果有效性与可靠性", "仿真前检查模型；结果应非空、达到 StopTime 且无 NaN。", "运行记录含 5001 个时间样本、50 s 结束时间和 valid=true。"],
        ["NF-04", "可移植性", "控制逻辑可从图形模型输出 C99，并保留 ROS/Gazebo 接入边界。", "第 7 章；不将离线 SIL 表述为飞行验收。"],
    ], [1.6, 3.3, 6.0, 5.9])

    add_heading(document, "2.3 软件缺陷", 2)
    add_paragraph(document, "维护前，px4ctrl 姿态桥接需要把期望四元数转换为欧拉角。直接将计算量传给 asin 在边界浮点误差或 MWORKS 事件处理时可能超出 [-1, 1]，导致非实数域计算。维护后采用带极小裕量的截断值参与 asin，并同时保留 pitch_argument_clipped 供诊断。该修复是局部的，不改变正常区间内的控制公式。")

    add_heading(document, "2.4 代码坏味道", 2)
    add_table(document, ["坏味道", "风险", "重构/约束"], [
        ["控制器选择散落在调用方", "新增算法时易漏改且难以审查边界", "收敛为 controller_id 到 Runner/Boundary 的显式路由表。"],
        ["场景参数硬编码在实验模型", "对比条件不透明，不能复现扰动/失配", "将轨迹、注入、求解器和指标拆入 Profile 与 Contract。"],
        ["把结果文件存在视为测试通过", "空文件、提前终止或 NaN 会被误用", "保留样本数、结束时间、NaN 和指标有效性门。"],
    ], [4.2, 5.0, 7.6])

    add_heading(document, "3. 软件构造", 1)
    add_heading(document, "3.1 软件体系结构设计", 2)
    add_paragraph(document, "维护后的体系结构由配置层、控制器适配层、FormalRunner 闭环层、机体与传感层、结果与报告层组成。配置层只选择资产和声明比较口径；Adapter 负责将各类控制器投影到共享边界；FormalRunner 将 Adapter、共享 Plant 和轨迹组成可检查的整机闭环；结果层记录原始序列、指标、配置与哈希。")
    with zipfile.ZipFile(TECHNICAL_MATERIAL) as source_archive:
        add_figure(document, source_archive, "image13.png", "图 3-1　控制器路线、输出边界与 FormalRunner 分层架构")

        add_heading(document, "3.2 用户界面设计", 2)
        add_paragraph(document, "用户入口面向任务、控制器和场景参数三类选择，显示层只承担路由、配置编辑、结果浏览和证据定位。界面中的“可用”只表示已登记的 Runner 和输出边界存在，不等同于性能、七场景、代码生成或运行时验收通过。")
        add_figure(document, source_archive, "image369.png", "图 3-2　Model Studio、QGC 与 Gazebo 的显示/操作边界")

        add_heading(document, "3.3 用例设计", 2)
        add_table(document, ["用例", "参与者", "主成功场景", "失败处理"], [
            ["UC-01 选择控制器", "维护者", "选择 controller_id，系统展示 Runner 和输出边界。", "无登记项时不给出执行入口。"],
            ["UC-02 配置七场景", "维护者", "选择 Profile，加载轨迹、时长和注入参数。", "Profile/Contract 版本不匹配时终止记录。"],
            ["UC-03 查看结果", "维护者/评审者", "读取有效原始序列和指标，定位配置与模型哈希。", "缺样本、未达到 StopTime 或 NaN 时标记无效。"],
            ["UC-04 三机编队", "维护者", "创建三机槽位、参考和共享控制闭环。", "任一关键连接缺失时由模型检查阻断。"],
        ], [3.0, 2.2, 7.2, 4.4])

        add_heading(document, "3.4 类设计", 2)
        add_table(document, ["模块", "职责", "依赖方向"], [
            ["PartialAttitudeThrustController", "定义位置、速度、姿态输入和姿态/总推力输出契约。", "被具体 Adapter 实现。"],
            ["Px4CtrlAttitudeThrustAdapter", "把 px4ctrl 图形外环和方程桥接收敛到共享契约。", "依赖控制实现，不反向依赖 Runner。"],
            ["Px4CtrlFormalRunner", "配置 50 s MWORKS 整机闭环与求解器参数。", "依赖 Adapter、轨迹和共享 Plant。"],
            ["Px4CtrlThreeUavFigure8Runner", "复用单机单元并增加三机参考、槽位与队形指标。", "依赖同一 Adapter、Allocator 和 Plant。"],
        ], [4.5, 7.0, 5.3])

        add_heading(document, "3.5 数据设计", 2)
        add_table(document, ["数据对象", "关键字段", "用途"], [
            ["Controller route", "controller_id、runner_class、runner_file、boundary", "控制器选择与边界可审查。"],
            ["Scenario profile", "scenario_id、trajectory、duration、parameter overrides", "固定实验条件，避免硬编码。"],
            ["Injection contract", "solver、tolerance、metrics、validity gate", "规定注入语义与比较口径。"],
            ["Run record", "source hashes、time sample count、duration、NaN、metrics", "追溯一次结果的输入、输出和有效性。"],
        ], [3.4, 7.5, 5.9])

        add_heading(document, "3.6 设计模式", 2)
        add_table(document, ["模式", "项目中的对应关系", "收益"], [
            ["Strategy", "以 controller_id 选择控制器实现路线。", "算法替换不改变共享闭环结构。"],
            ["Adapter", "不同控制输出语义转换为共享机体输入。", "隔离接口差异并支持统一测试。"],
            ["Template Method", "FormalRunner 固定建模/检查/仿真/结果结构，参数由 Profile 填充。", "不同场景共享执行骨架。"],
            ["Factory/Registry", "路由与模块注册表集中声明可选择资产。", "减少散落条件分支，方便审查。"],
        ], [3.2, 8.0, 5.6])

        add_heading(document, "4. 软件实现", 1)
        add_heading(document, "4.1 新增功能展示", 2)
        add_paragraph(document, "F-01 的路由清单将 controller_id、模型类、文件路径和输出边界绑定。F-02 的七场景 Contract 把求解器、采样、参考与测量信号以及故障注入语义固定下来。F-03 的三机 Runner 直接复用三个相同的单机控制闭环，并输出 formation_error_m 与 min_inter_uav_distance_m。")
        add_figure(document, source_archive, "image331.png", "图 4-1　px4ctrl 图形化位置/速度外环")
        add_figure(document, source_archive, "image343.png", "图 4-2　三机 Figure8 编队场景与规划轨迹")
        add_figure(document, source_archive, "image361.png", "图 4-3　Sysblock 到 C99、SIL 与 ROS/Gazebo 的交付链路")

    add_heading(document, "4.2 代码重构", 2)
    add_paragraph(document, "C-01 的重构对象是姿态桥接中的俯仰角计算。维护前的基线实现直接把四元数计算量传入 asin；维护后把定义域保护、截断诊断和原有饱和分支分开表达。下面两段代码分别对应项目中的 Px4CtrlEquationBridgeReportBaselineAdapter.mo 和 Px4CtrlAttitudeThrustAdapter.mo，保留相同的 pitch_argument 计算式，便于逐行核对变更范围。")
    add_code(document, "代码清单 4-1　维护前的直接 asin 转换（节选）", """
pitch_argument = 2 * (core.qd_w * core.qd_y - core.qd_z * core.qd_x);
pitch_ref = if pitch_argument >= 1 then Modelica.Constants.pi / 2
  else if pitch_argument <= -1 then -Modelica.Constants.pi / 2
  else asin(pitch_argument);""")
    add_code(document, "代码清单 4-2　维护后的安全域保护（节选）", """
parameter Real pitch_argument_domain_margin(min = 0, max = 0.01) = 1e-7;
pitch_argument = 2 * (core.qd_w * core.qd_y - core.qd_z * core.qd_x);
pitch_argument_safe = min(1 - pitch_argument_domain_margin,
  max(-1 + pitch_argument_domain_margin, pitch_argument));
pitch_argument_clipped = abs(pitch_argument - pitch_argument_safe) > 0;
pitch_ref = if pitch_argument >= 1 then Modelica.Constants.pi / 2
  else if pitch_argument <= -1 then -Modelica.Constants.pi / 2
  else asin(pitch_argument_safe);""")
    add_code(document, "代码清单 4-3　三机 Runner 的复用式构造（节选）", """
MoSimQuadrotorModel.Control.Adapters.Px4CtrlAttitudeThrustAdapter controller_1;
MoSimQuadrotorModel.Control.Allocation.OfflineAttitudeRateAllocator allocator_1;
MoSimQuadrotorModel.Vehicle.Sunray150Assembly plant_1(...);
...
formation_error_m = 0.5 * (formation_error_2_m + formation_error_3_m);
min_inter_uav_distance_m = min(inter_uav_distance_12_m,
  min(inter_uav_distance_13_m, inter_uav_distance_23_m));""")

    add_table(document, ["重构项", "维护前风险", "维护后"], [
        ["姿态转换", "直接调用 asin(pitch_argument)，浮点舍入或事件求值可能越出实数定义域。", "夹取到带裕量的安全值，并输出 pitch_argument_clipped 诊断。"],
        ["控制器接入", "控制器与 Plant 的边界语义可能散落在不同模型。", "使用 Partial 接口和 Adapter 聚合输入输出语义。"],
        ["场景参数", "场景时间、扰动和评价指标易与控制器实现耦合。", "使用 Profile/Contract 版本化数据结构。"],
    ], [3.2, 6.0, 7.6])

    add_heading(document, "5. 软件测试", 1)
    add_heading(document, "5.1 测试计划", 2)
    add_paragraph(document, "测试分为结构/静态检查、已有 MWORKS 结果记录复核和文档交付核验三类。前两类不替代实时 ROS/Gazebo 或飞行验收；后者确认本报告的章节、图表和表格能够独立呈现。")

    add_heading(document, "5.2 测试环境", 2)
    add_table(document, ["层级", "环境", "用途"], [
        ["模型与静态层", "MWORKS Modelica 工程、JSON/TOML 配置", "验证模型路径、接口、配置数量和指标契约。"],
        ["结果复核层", "已有 RUN_RECORD、原始结果、指标与截图", "检查时间序列完整性和结果边界。"],
        ["文档层", "Word 模板、python-docx、DOCX 渲染器", "确认模板结构、表格、图片、页眉页脚与中文排版。"],
    ], [3.2, 6.0, 7.6])

    add_heading(document, "5.3 设计测试用例", 2)
    add_table(document, ["编号", "对象", "测试步骤", "预期结果"], [
        ["TC-01", "F-01 路由", "解析路由 TOML，检查 controller_id、runner_file、boundary。", f"发现 {facts['route_count']} 条显式路由，字段完整且路径可定位。"],
        ["TC-02", "F-02 Profile", "读取 v2 Profile 和 Injection Contract，核对场景、求解器、指标。", f"Profile 数为 {facts['scenario_count']}，场景与 Contract 对应。"],
        ["TC-03", "F-03 三机模型", "静态检查三套 controller/allocator/plant 及队形误差公式。", "存在三组闭环实例和最小机间距输出。"],
        ["TC-04", "C-01 修复", "检查 asin 的输入使用 clip 后变量，保留 clipped 诊断。", "正常范围不改变，边界不进入非实数域。"],
        ["TC-05", "NF-03 结果有效性", "复核 px4ctrl 50 s RUN_RECORD 的样本、结束时间、NaN 与 valid。", f"{facts['run_rows']} 个样本、{facts['run_duration']} s、NaN=0、valid=true。"],
        ["TC-06", "文档交付", "检查两个课程模板的章结构、16 章材料映射和图表可渲染性。", "项目综合实践 III 原系统报告和本维护报告形成两份可独立提交的文档。"],
    ], [1.5, 3.0, 7.0, 5.3])

    add_heading(document, "5.4 测试用例执行", 2)
    add_paragraph(document, "TC-01 至 TC-04 为当前源代码和配置的静态核验；TC-05 复核已有的 MWORKS 离线闭环记录。该记录表明 px4ctrl FormalRunner 在 50 s ClimbPath 中保留 5001 个时间样本，位置 RMSE 为 %.6f m，终端位置误差为 %.6f m。该条记录的结论边界是离线 MWORKS 整机闭环，不包括 PX4、Gazebo、ROS 或飞行运行时行为。" % (facts["run_rmse"], facts["run_terminal"]))
    add_paragraph(document, "为验证本次维护涉及的适配器、任务配置、图形 C99 后端和模型入口，使用工作区 Python 3.13.9 与 pytest 8.4.2 运行以下定向自动化测试。命令在 2026 年 8 月 10 日执行，退出码为 0。")
    add_code(document, "代码清单 5-1　定向自动化测试命令", """
python -m pytest Scripts/tests/test_px4ctrl_mworks_adapter.py \\
  Scripts/tests/test_model_studio_task_handoff.py \\
  Scripts/tests/test_px4ctrl_graphical_c99_runtime_contract.py \\
  Scripts/tests/test_model_studio_catalog.py \\
  Scripts/tests/test_mworks_codegen_runtime.py \\
  Scripts/tests/test_px4ctrl_open_blocks_mworks.py""")
    add_table(document, ["测试文件", "用例数", "覆盖内容", "执行结果"], [
        ["test_px4ctrl_mworks_adapter.py", "10", "姿态桥接安全域、共享契约、Runner 与包注册", "通过"],
        ["test_model_studio_task_handoff.py", "13", "任务配置写入、路由、场景与三机交接", "通过"],
        ["test_px4ctrl_graphical_c99_runtime_contract.py", "8", "图形 C99 后端、包装和运行时接入契约", "通过"],
        ["test_model_studio_catalog.py", "6", "Studio 目录、工作区和控制器选择", "通过"],
        ["test_mworks_codegen_runtime.py", "4", "生成代码编译、临时运行时夹具与路径约束", "通过"],
        ["test_px4ctrl_open_blocks_mworks.py", "4", "OpenBlocks 参考、Runner 和包注册", "通过"],
        ["合计", "45", "维护相关的源代码/契约回归", "45 passed in 3.54s"],
    ], [5.2, 1.5, 8.0, 3.4])
    add_code(document, "代码清单 5-2　测试输出（节选）", """
.............................................
45 passed in 3.54s""")
    add_heading(document, "5.5 测试结果及分析", 2)
    add_paragraph(document, "上述 45 项测试全部通过，说明路由、配置、姿态安全域、图形 C99 包装和三机模型的结构资产可以互相定位，生成代码的本地编译与临时运行时夹具也满足既定契约。该组测试主要是源代码、配置和生成代码的自动化回归；其中的本地夹具不启动 MWORKS、ROS、Gazebo 或 PX4，因此不替代第 7 章中按证据边界说明的模型结果或运行时验收。已有 MWORKS 结果复核则表明所引用记录满足非空时间序列、停止时间和 NaN 有效性条件。")
    add_heading(document, "5.6 缺陷修复", 2)
    add_paragraph(document, "本次明确记录的缺陷修复为 C-01：对 px4ctrl 的欧拉角转换加入 asin 安全域保护，并输出是否发生截断的诊断标识。对于结果层面，采用样本数、StopTime、NaN 和指标有效性共同判断，而不是仅根据文件是否生成判断通过。")

    add_heading(document, "6. 课设总结", 1)
    add_paragraph(document, "本项目以既有赛题系统为基础，完成了面向维护的需求化重组：四项新增功能满足“至少三项功能需求”，四项非功能要求满足“至少三项非功能需求”，并把改正性、适应性和预防性维护落实到具体模型、配置、结果合同和测试用例。")
    add_paragraph(document, "本报告将新增功能、非功能需求、缺陷修复、重构和测试组织为完整的维护闭环。后续章节集中给出平台在名义筛查、七场景、多机、代码生成和运行时集成中的详细验证材料，所有结论均按其实际证据范围解释。")
    add_heading(document, "6.3 成员分工与个人心得", 2)
    add_paragraph(document, "本组按系统维护链条组织工作，贡献比例由成员在答辩材料中另行填写；本节仅记录成员身份、承担内容和实践体会。")
    add_table(document, ["成员", "学号", "主要分工", "对应工作"], [
        ["刘致远", "231304113", "总体方案与报告整合", "维护需求梳理、章节组织与交付复核"],
        ["钟俊杰", "231304130", "物理模型与参数资料", "机体结构、参数来源与模型边界整理"],
        ["朱尚吉", "231304133", "控制器、接口与 Runner", "Adapter、FormalRunner 与场景接入核对"],
        ["陈健", "231304103", "Studio、代码生成与测试", "任务配置、C99 链路与自动化回归"],
        ["王家祺", "231304120", "运行时材料、图表与证据归档", "结果材料整理、图表核查与证据边界说明"],
    ], [2.2, 3.0, 4.4, 5.6])
    add_heading(document, "6.3.1 刘致远", 3)
    add_paragraph(document, "通过统筹维护需求和报告结构，我认识到课程设计的重点不只是完成若干功能，更要说明每项改动为什么存在、影响哪些模块以及怎样验证。把需求、设计、实现和测试放在同一条维护链上后，讨论和复核都更有依据。")
    add_heading(document, "6.3.2 钟俊杰", 3)
    add_paragraph(document, "整理机体模型和参数资料的过程让我意识到，参数不能脱离来源和适用边界单独使用。将结构信息、单位、配置和模型路径对应起来，能够减少后续维护中因理解不一致造成的重复修改。")
    add_heading(document, "6.3.3 朱尚吉", 3)
    add_paragraph(document, "控制器接入工作使我更加理解接口抽象的价值。把控制器差异收敛在 Adapter 和 Runner 的明确边界内，既便于替换实现，也能让故障定位集中在较小范围，而不是扩散到共享 Plant 或场景配置。")
    add_heading(document, "6.3.4 陈健", 3)
    add_paragraph(document, "代码生成和测试工作让我认识到，自动化测试不应只验证正常路径。对输入定义域、路由组合、生成代码包装和异常结果分别设置检查，才能让维护后的结论可重复，并在回归时及时发现接口或配置变化。")
    add_heading(document, "6.3.5 王家祺", 3)
    add_paragraph(document, "整理结果和图表时，我体会到证据记录本身也是软件维护的一部分。清楚地区分静态检查、模型结果和运行时材料，可以避免把展示画面或短时夹具误写成完整验收，也使后续成员能够快速追溯结论来源。")


def enable_field_update(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_dynamic_page_footers(document: Document, first_section: int) -> None:
    for section in document.sections[first_section:]:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        for child in list(footer._element):
            footer._element.remove(child)
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        add_field(paragraph, " PAGE \\* MERGEFORMAT ", "1")


def validate(output: Path, imported: ImportStats) -> None:
    if imported.image_references != 314:
        raise RuntimeError(f"Maintenance report must receive 314 technical figure references, got {imported.image_references}")
    document = Document(output)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    text = "\n".join(text_parts)
    required = [
        "1. 项目概述", "2. 软件分析", "3. 软件构造", "4. 软件实现", "5. 软件测试", "6. 课设总结",
        "7. 维护验证与运行时集成",
        "F-01", "F-02", "F-03", "NF-01", "NF-02", "NF-03",
        "刘致远", "钟俊杰", "朱尚吉", "陈健", "王家祺",
        "维护前的直接 asin 转换", "45 passed in 3.54s",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Generated report is missing required content: {missing}")
    if len(document.inline_shapes) < imported.image_references + 5:
        raise RuntimeError("Generated report is missing imported figures")

    normal = document.styles["Normal"]
    if normal.font.size is None or abs(normal.font.size.pt - 12) > 0.01:
        raise RuntimeError("National template Normal style must remain 12-point")
    if normal.paragraph_format.first_line_indent is None or abs(normal.paragraph_format.first_line_indent.cm - 0.35) > 0.01:
        raise RuntimeError("National template Normal style must keep its first-line indent")

    code_tables = [table for table in document.tables if _is_code_table(table)]
    if len(code_tables) < 5:
        raise RuntimeError(f"Expected code listings as one-cell tables, got {len(code_tables)}")
    for table in code_tables:
        if len(table.rows) != 1 or len(table.rows[0].cells) != 1 or table.style.name != "Table Grid":
            raise RuntimeError("A code listing is not a one-row, one-column table")
        paragraph = table.cell(0, 0).paragraphs[0]
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            raise RuntimeError("Code listing table must remain left aligned")
        for run in paragraph.runs:
            if run.font.size is None or abs(run.font.size.pt - 12) > 0.01:
                raise RuntimeError("Code listing font must follow the national template body size")

    for table in document.tables:
        if _is_code_table(table):
            continue
        if table.style.name != "三线表":
            raise RuntimeError(f"Data table does not use the national template three-line style: {table.style.name}")
        for row in table.rows:
            for cell in row.cells:
                if cell._tc.get_or_add_tcPr().find(qn("w:shd")) is not None:
                    raise RuntimeError("Data table retains non-template cell shading")
    with zipfile.ZipFile(output) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"DOCX ZIP integrity failed at {bad}")


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    for required in (TECHNICAL_MATERIAL,):
        if not required.exists():
            raise FileNotFoundError(required)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    facts = load_source_facts()
    document = Document(TEMPLATE)
    clear_template_body(document)
    configure_styles(document)
    add_report_front_matter(document)
    add_body(document, facts)
    imported = import_paragraph_range(
        document,
        TECHNICAL_MATERIAL,
        start_paragraph=498,
        end_paragraph=None,
        section_heading="7. 维护验证与运行时集成",
        preserve_table_xml=False,
    )
    apply_national_template_layout(document)
    layout = {
        "tables": len(document.tables),
        "code_tables": sum(1 for table in document.tables if _is_code_table(table)),
        "figures": len(document.inline_shapes),
    }
    enable_field_update(document)
    document.core_properties.title = "软件构造课程设计报告：MoSim 项目维护"
    document.core_properties.subject = "基于 MWORKS 的四旋翼位姿控制仿真平台项目维护"
    document.core_properties.comments = "软件构造课程设计独立维护报告。"
    document.save(OUTPUT)
    validate(OUTPUT, imported)
    print(OUTPUT)
    print(json.dumps({**facts, **imported.__dict__, **layout}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
